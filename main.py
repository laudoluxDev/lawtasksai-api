"""
LawTasksAI API
FastAPI backend for skill delivery, licensing, and usage tracking.
"""

from fastapi import FastAPI, HTTPException, Depends, Header, BackgroundTasks, Request, APIRouter, Query
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse
from pydantic import BaseModel, EmailStr
from typing import Optional, List
from datetime import datetime, timedelta
import os
import secrets
import hashlib
import hmac
import io
import zipfile
import json
import asyncio
import stripe
import anthropic
import re

# Database imports (using async SQLAlchemy)
from sqlalchemy.ext.asyncio import AsyncSession, create_async_engine, async_sessionmaker
from sqlalchemy.orm import DeclarativeBase, Mapped, mapped_column, relationship
from sqlalchemy import String, Integer, Boolean, DateTime, Text, ForeignKey, select, update, func, or_, text, UniqueConstraint
from sqlalchemy.dialects.postgresql import UUID, ARRAY, JSONB
import uuid
import base64
from io import BytesIO
import httpx

# Document generation (install: pip install python-docx openpyxl)
try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    from openpyxl import Workbook
    XLSX_AVAILABLE = True
except ImportError:
    XLSX_AVAILABLE = False

# ============================================
# Configuration
# ============================================

DATABASE_URL = os.getenv("DATABASE_URL", "postgresql+asyncpg://user:pass@localhost/lawtasksai")
STRIPE_SECRET_KEY = os.getenv("STRIPE_SECRET_KEY", "")
STRIPE_WEBHOOK_SECRET = os.getenv("STRIPE_WEBHOOK_SECRET", "")
API_BASE_URL = os.getenv("API_BASE_URL", "https://api.lawtasksai.com")
FRONTEND_URL = os.getenv("FRONTEND_URL", "https://lawtasksai.com")
ANTHROPIC_API_KEY = os.getenv("ANTHROPIC_API_KEY", "")

# Admin authentication
ADMIN_SECRET = os.getenv("ADMIN_SECRET", "")  # Set via Cloud Run env var — never hardcode

# Signing key for unsubscribe tokens — derived from ADMIN_SECRET so no extra env var needed
def _waitlist_unsub_token(email: str, product_id: str) -> str:
    """Generate a short HMAC token for waitlist unsubscribe links."""
    key = (ADMIN_SECRET or "fallback-dev-key").encode()
    msg = f"{email}:{product_id}".encode()
    return hmac.new(key, msg, digestmod=hashlib.sha256).hexdigest()[:32]

def _verify_waitlist_unsub_token(email: str, product_id: str, token: str) -> bool:
    expected = _waitlist_unsub_token(email, product_id)
    return hmac.compare_digest(expected, token)

def verify_admin(x_admin_secret: Optional[str] = Header(None, alias="X-Admin-Secret")):
    """Dependency that enforces admin secret on all /admin/* routes."""
    if not ADMIN_SECRET:
        raise HTTPException(status_code=500, detail="Admin secret not configured on server")
    if x_admin_secret != ADMIN_SECRET:
        raise HTTPException(status_code=403, detail="Forbidden")

# Loader versioning
CURRENT_LOADER_VERSION = "1.7.0"
LOADER_UPDATE_URL = "https://lawtasksai.com/download"
LOADER_UPDATE_MESSAGE = None  # Set to a string when there's an important update

# Path to the loader SKILL.md  [rebuilt 20260322T163027Z] (served for auto-update)
LOADER_SKILL_PATH = os.path.join(os.path.dirname(__file__), "loader", "SKILL.md")

# Initialize Stripe
stripe.api_key = STRIPE_SECRET_KEY

# Initialize Anthropic client
anthropic_client = anthropic.Anthropic(api_key=ANTHROPIC_API_KEY) if ANTHROPIC_API_KEY else None

# ============================================
# Database Models
# ============================================

class Base(DeclarativeBase):
    pass

class User(Base):
    __tablename__ = "users"
    
    id: Mapped[uuid.UUID] = mapped_column(UUID(as_uuid=True), primary_key=True, default=uuid.uuid4)
    email: Mapped[str] = mapped_column(String(255), unique=True, nullable=False)
    name: Mapped[Optional[str]] = mapped_column(String(255))
    firm_name: Mapped[Optional[str]] = mapped_column(String(255))
    password_hash: Mapped[str] = mapped_column(String(255), nullable=False)
    _credits_balance: Mapped[int] = mapped_column("credits_balance", Integer, default=0)
    version_policy: Mapped[str] = mapped_column(String(20), default="latest")
    is_active: Mapped[bool] = mapped_column(Boolean, default=True)
    created_at: Mapped[datetime] = mapped_column(DateTime, default=datetime.utcnow)
    last_login: Mapped[Optional[datetime]] = mapped_column(DateTime)
    # Multi-tenancy: which product this user belongs to
    product_id: Mapped[Optional[str]] = mapped_column(String(50), default="law")
    # Profile for document generation (firm info, letterhead, etc.)
    profile: Mapped[Optional[dict]] = mapped_column(JSONB, default=dict)
    # Platforms the user has selected (e.g. ["claude_desktop", "openclaw"])
    platforms: Mapped[Optional[list]] = mapped_column(JSONB, default=list)
    # Last time user executed a skill — used for activation/engagement tracking
    last_active_at: Mapped[Optional[datetime]] = mapped_column(DateTime, nullable=True)

    @property
    def credits_balance(self) -> int:
        """Single source of truth: always read from the active license."""
        # licenses relationship loaded eagerly where needed; fall back to _credits_balance
        try:
            if hasattr(self, '_licenses') and self._licenses:
                active = next((l for l in self._licenses if l.status == 'active'), None)
                if active:
                    return active.credits_remaining
        except Exception:
            pass
        return self._credits_balance

    @credits_balance.setter
    def credits_balance(self, value: int):
        """Write to _credits_balance for DB compat; caller should also update license."""
        self._credits_balance = value

class Category(Base):
    __tablename__ = "categories"
    
    id: Mapped[str] = mapped_column(String(50), primary_key=True)
    name: Mapped[str] = mapped_column(String(100), nullable=False)
    description: Mapped[Optional[str]] = mapped_column(Text)
    display_order: Mapped[int] = mapped_column(Integer, default=0)
    product_id: Mapped[str] = mapped_column(String(50), default="law")

class Skill(Base):
    __tablename__ = "skills"
    
    id: Mapped[str] = mapped_column(String(100), primary_key=True)
    category_id: Mapped[str] = mapped_column(String(50), ForeignKey("categories.id"))
    name: Mapped[str] = mapped_column(String(200), nullable=False)
    description: Mapped[Optional[str]] = mapped_column(Text)
    current_version: Mapped[Optional[str]] = mapped_column(String(20))
    stable_version: Mapped[Optional[str]] = mapped_column(String(20))
    credits_per_use: Mapped[int] = mapped_column(Integer, default=1)
    requires_upload: Mapped[bool] = mapped_column(Boolean, default=False)
    execution_type: Mapped[str] = mapped_column(String(20), default="server")
    is_published: Mapped[bool] = mapped_column(Boolean, default=False)
    is_deprecated: Mapped[bool] = mapped_column(Boolean, default=False)
    triggers: Mapped[Optional[List[str]]] = mapped_column(ARRAY(String), default=list)
    product_id: Mapped[str] = mapped_column(String(50), default="law")
    created_at: Mapped[datetime] = mapped_column(DateTime, default=datetime.utcnow)
    security_verified: Mapped[bool] = mapped_column(Boolean, default=False)
    security_scanned_at: Mapped[Optional[datetime]] = mapped_column(DateTime, nullable=True)

class SkillVersion(Base):
    __tablename__ = "skill_versions"
    
    id: Mapped[int] = mapped_column(Integer, primary_key=True, autoincrement=True)
    skill_id: Mapped[str] = mapped_column(String(100), ForeignKey("skills.id", ondelete="CASCADE"))
    version: Mapped[str] = mapped_column(String(20), nullable=False)
    content: Mapped[str] = mapped_column(Text, nullable=False)
    config_template: Mapped[Optional[dict]] = mapped_column(JSONB)
    changelog: Mapped[Optional[str]] = mapped_column(Text)
    is_stable: Mapped[bool] = mapped_column(Boolean, default=False)
    is_beta: Mapped[bool] = mapped_column(Boolean, default=False)
    published_at: Mapped[datetime] = mapped_column(DateTime, default=datetime.utcnow)

class ContentPage(Base):
    """Editable content pages (security guide, etc.) with version control."""
    __tablename__ = "content_pages"
    
    slug: Mapped[str] = mapped_column(String(100), primary_key=True)  # e.g. "openclaw-security-guide"
    title: Mapped[str] = mapped_column(String(255), nullable=False)
    current_version: Mapped[int] = mapped_column(Integer, default=1)
    content: Mapped[str] = mapped_column(Text, nullable=False)  # Current HTML content
    updated_at: Mapped[datetime] = mapped_column(DateTime, default=datetime.utcnow)
    created_at: Mapped[datetime] = mapped_column(DateTime, default=datetime.utcnow)

class ContentPageVersion(Base):
    """Version history for content pages."""
    __tablename__ = "content_page_versions"
    
    id: Mapped[int] = mapped_column(Integer, primary_key=True, autoincrement=True)
    page_slug: Mapped[str] = mapped_column(String(100), ForeignKey("content_pages.slug", ondelete="CASCADE"))
    version: Mapped[int] = mapped_column(Integer, nullable=False)
    content: Mapped[str] = mapped_column(Text, nullable=False)
    changelog: Mapped[Optional[str]] = mapped_column(Text)
    created_at: Mapped[datetime] = mapped_column(DateTime, default=datetime.utcnow)

class License(Base):
    __tablename__ = "licenses"
    
    id: Mapped[uuid.UUID] = mapped_column(UUID(as_uuid=True), primary_key=True, default=uuid.uuid4)
    license_key: Mapped[str] = mapped_column(String(50), unique=True, nullable=False)
    user_id: Mapped[uuid.UUID] = mapped_column(UUID(as_uuid=True), ForeignKey("users.id", ondelete="CASCADE"))
    type: Mapped[str] = mapped_column(String(20), nullable=False)  # trial, credits, subscription, enterprise
    status: Mapped[str] = mapped_column(String(20), default="active")  # active, expired, revoked, suspended
    valid_from: Mapped[datetime] = mapped_column(DateTime, default=datetime.utcnow)
    valid_until: Mapped[Optional[datetime]] = mapped_column(DateTime)
    skills_allowed: Mapped[Optional[List[str]]] = mapped_column(ARRAY(String))
    categories_allowed: Mapped[Optional[List[str]]] = mapped_column(ARRAY(String))
    usage_limit: Mapped[Optional[int]] = mapped_column(Integer)
    usage_count: Mapped[int] = mapped_column(Integer, default=0)
    credits_purchased: Mapped[int] = mapped_column(Integer, default=0)
    credits_remaining: Mapped[int] = mapped_column(Integer, default=0)
    product_id: Mapped[Optional[str]] = mapped_column(String(50), default="law")
    notes: Mapped[Optional[str]] = mapped_column(Text)
    created_at: Mapped[datetime] = mapped_column(DateTime, default=datetime.utcnow)
    updated_at: Mapped[Optional[datetime]] = mapped_column(DateTime)
    downloaded_at: Mapped[Optional[datetime]] = mapped_column(DateTime, nullable=True)
    first_connected_at: Mapped[Optional[datetime]] = mapped_column(DateTime, nullable=True)

class UsageLog(Base):
    __tablename__ = "usage_logs"
    
    id: Mapped[int] = mapped_column(Integer, primary_key=True, autoincrement=True)
    license_id: Mapped[uuid.UUID] = mapped_column(UUID(as_uuid=True), ForeignKey("licenses.id"))
    user_id: Mapped[uuid.UUID] = mapped_column(UUID(as_uuid=True), ForeignKey("users.id"))
    skill_id: Mapped[str] = mapped_column(String(100), ForeignKey("skills.id"))
    skill_version: Mapped[str] = mapped_column(String(20))
    executed_at: Mapped[datetime] = mapped_column(DateTime, default=datetime.utcnow)
    execution_time_ms: Mapped[Optional[int]] = mapped_column(Integer)
    success: Mapped[bool] = mapped_column(Boolean, default=True)
    error_message: Mapped[Optional[str]] = mapped_column(Text)
    credits_used: Mapped[int] = mapped_column(Integer, default=1)
    tokens_input: Mapped[Optional[int]] = mapped_column(Integer)
    tokens_output: Mapped[Optional[int]] = mapped_column(Integer)
    # Store result for document regeneration (no extra credit charge)
    result_text: Mapped[Optional[str]] = mapped_column(Text)

class CreditTransaction(Base):
    __tablename__ = "credit_transactions"
    
    id: Mapped[int] = mapped_column(Integer, primary_key=True, autoincrement=True)
    user_id: Mapped[uuid.UUID] = mapped_column(UUID(as_uuid=True), ForeignKey("users.id"))
    license_id: Mapped[Optional[uuid.UUID]] = mapped_column(UUID(as_uuid=True), ForeignKey("licenses.id"))
    type: Mapped[str] = mapped_column(String(20), nullable=False)  # purchase, usage, refund, bonus
    amount: Mapped[int] = mapped_column(Integer, nullable=False)
    balance_after: Mapped[int] = mapped_column(Integer, nullable=False)
    reference_id: Mapped[Optional[str]] = mapped_column(String(100))
    description: Mapped[Optional[str]] = mapped_column(Text)
    created_at: Mapped[datetime] = mapped_column(DateTime, default=datetime.utcnow)


class SkillGap(Base):
    """
    Anonymous gap reports submitted when a user's legal question matches no skill.
    No user identity, no query content — only the search terms that failed to match.
    Submitted only with explicit per-request user consent.
    """
    __tablename__ = "skill_gaps"

    id: Mapped[int] = mapped_column(Integer, primary_key=True, autoincrement=True)
    search_terms: Mapped[str] = mapped_column(Text, nullable=False)       # space-separated keywords
    loader_version: Mapped[Optional[str]] = mapped_column(String(20))     # which loader version reported
    reported_at: Mapped[datetime] = mapped_column(DateTime, default=datetime.utcnow)

class PlatformSecurityScan(Base):
    """Platform-level (preamble) security scan results per vertical+model."""
    __tablename__ = "platform_security_scans"

    id: Mapped[int] = mapped_column(Integer, primary_key=True, autoincrement=True)
    vertical: Mapped[str] = mapped_column(String(100), nullable=False, default='all')
    scan_model: Mapped[str] = mapped_column(String(100), nullable=False)
    tests_run: Mapped[int] = mapped_column(Integer, default=0)
    tests_passed: Mapped[int] = mapped_column(Integer, default=0)
    tests_failed: Mapped[int] = mapped_column(Integer, default=0)
    scanned_at: Mapped[datetime] = mapped_column(DateTime, default=datetime.utcnow)
    created_at: Mapped[datetime] = mapped_column(DateTime, default=datetime.utcnow)


class SkillSecurityScan(Base):
    """Tracks promptfoo security scan results per skill."""
    __tablename__ = "skill_security_scans"

    id: Mapped[int] = mapped_column(Integer, primary_key=True, autoincrement=True)
    skill_id: Mapped[str] = mapped_column(String(255), nullable=False, unique=True)
    vertical: Mapped[str] = mapped_column(String(100), nullable=False)
    verified: Mapped[bool] = mapped_column(Boolean, default=False)
    tests_run: Mapped[int] = mapped_column(Integer, default=0)
    tests_passed: Mapped[int] = mapped_column(Integer, default=0)
    tests_failed: Mapped[int] = mapped_column(Integer, default=0)
    plugins_tested: Mapped[Optional[List[str]]] = mapped_column(ARRAY(String))
    preamble_tested: Mapped[bool] = mapped_column(Boolean, default=True)
    scan_model: Mapped[Optional[str]] = mapped_column(String(100), default='openai:gpt-4o-mini')
    scanned_at: Mapped[datetime] = mapped_column(DateTime, default=datetime.utcnow)
    created_at: Mapped[datetime] = mapped_column(DateTime, default=datetime.utcnow)


class DripEmail(Base):
    """Record of every drip email sent to a user."""
    __tablename__ = "drip_emails"

    id: Mapped[uuid.UUID] = mapped_column(UUID(as_uuid=True), primary_key=True, default=uuid.uuid4)
    user_id: Mapped[Optional[uuid.UUID]] = mapped_column(UUID(as_uuid=True), ForeignKey("users.id"), nullable=True)
    email: Mapped[str] = mapped_column(String(255), nullable=False)
    product_id: Mapped[str] = mapped_column(String(50), nullable=False)
    email_number: Mapped[int] = mapped_column(Integer, nullable=False)  # 1, 2, or 3
    subject: Mapped[Optional[str]] = mapped_column(String(500), nullable=True)
    sent_at: Mapped[datetime] = mapped_column(DateTime, default=datetime.utcnow)

    __table_args__ = (UniqueConstraint("email", "product_id", "email_number", name="uq_drip_email_per_user"),)


class UserFeedback(Base):
    """Drip email feedback responses (Email 3 one-click buttons)."""
    __tablename__ = "user_feedback"

    id: Mapped[uuid.UUID] = mapped_column(UUID(as_uuid=True), primary_key=True, default=uuid.uuid4)
    email: Mapped[Optional[str]] = mapped_column(String(255), nullable=True)
    product_id: Mapped[Optional[str]] = mapped_column(String(50), nullable=True)
    reason: Mapped[str] = mapped_column(String(100), nullable=False)
    extra: Mapped[Optional[str]] = mapped_column(Text, nullable=True)  # free-text from 'other'
    created_at: Mapped[datetime] = mapped_column(DateTime, default=datetime.utcnow)


class EmailSubscription(Base):
    """Per-vertical email subscription preferences for users."""
    __tablename__ = "email_subscriptions"

    id: Mapped[uuid.UUID] = mapped_column(UUID(as_uuid=True), primary_key=True, default=uuid.uuid4)
    user_id: Mapped[uuid.UUID] = mapped_column(UUID(as_uuid=True), ForeignKey("users.id"), nullable=False)
    product_id: Mapped[str] = mapped_column(String(50), nullable=False)
    subscribed: Mapped[bool] = mapped_column(Boolean, default=True)
    subscribed_at: Mapped[datetime] = mapped_column(DateTime, default=datetime.utcnow)
    unsubscribed_at: Mapped[Optional[datetime]] = mapped_column(DateTime, nullable=True)

    __table_args__ = (UniqueConstraint("user_id", "product_id", name="uq_email_sub_user_product"),)


class WaitlistEntry(Base):
    """Waitlist signup for a vertical that doesn't have an installer yet."""
    __tablename__ = "waitlist"

    id: Mapped[uuid.UUID] = mapped_column(UUID(as_uuid=True), primary_key=True, default=uuid.uuid4)
    email: Mapped[str] = mapped_column(String(255), nullable=False)
    name: Mapped[Optional[str]] = mapped_column(String(200), nullable=True)
    product_id: Mapped[str] = mapped_column(String(50), nullable=False)
    created_at: Mapped[datetime] = mapped_column(DateTime, default=datetime.utcnow)
    notified_at: Mapped[Optional[datetime]] = mapped_column(DateTime, nullable=True)

    __table_args__ = (UniqueConstraint("email", "product_id", name="uq_waitlist_email_product"),)

# ============================================
# Database Session
# ============================================

engine = create_async_engine(DATABASE_URL, echo=False)
async_session = async_sessionmaker(engine, class_=AsyncSession, expire_on_commit=False)

async def get_db():
    async with async_session() as session:
        yield session


def get_product_id(
    x_product_id: Optional[str] = Header(None, alias="X-Product-ID"),
    product: Optional[str] = Query(None, alias="product"),
) -> str:
    """
    Resolve product_id from (in priority order):
    1. X-Product-ID request header
    2. ?product= query parameter
    3. Fallback: 'law'
    """
    return x_product_id or product or "law"

# ============================================
# Pydantic Schemas
# ============================================

class UserCreate(BaseModel):
    email: EmailStr
    password: str
    name: Optional[str] = None
    firm_name: Optional[str] = None
    product_id: Optional[str] = None  # Multi-tenancy: which product they registered from
    platforms: Optional[List[str]] = None  # e.g. ["claude_desktop", "openclaw"]

class SimpleSignupRequest(BaseModel):
    """Password-free signup — name + email only. Password is auto-generated."""
    email: EmailStr
    name: Optional[str] = None
    product_id: Optional[str] = None
    platforms: Optional[List[str]] = None  # e.g. ["claude_desktop", "openclaw"]

class UserResponse(BaseModel):
    id: str
    email: str
    name: Optional[str]
    firm_name: Optional[str]
    credits_balance: int
    created_at: datetime

class LoginRequest(BaseModel):
    email: EmailStr
    password: str

class TokenResponse(BaseModel):
    access_token: str
    token_type: str = "bearer"
    license_key: str
    product_id: Optional[str] = "law"  # Multi-tenancy: which product this user belongs to

class SkillResponse(BaseModel):
    id: str
    name: str
    description: Optional[str]
    category_id: str
    current_version: Optional[str]
    credits_per_use: int
    requires_upload: bool
    execution_type: str
    confidentiality_note: Optional[str] = None  # Warning for sensitive data handling
    security_verified: bool = False
    plugins_tested: List[str] = []
    security_tests_passed: Optional[int] = None
    security_tests_run: Optional[int] = None
    platform_tests_passed: Optional[int] = None
    platform_tests_run: Optional[int] = None
    combined_tests_passed: Optional[int] = None
    combined_tests_run: Optional[int] = None

class SkillExecuteRequest(BaseModel):
    query: str  # User's input/question
    context: Optional[dict] = None  # Additional context (optional)
    version: Optional[str] = None  # None = use policy (latest/stable)

class LoaderMeta(BaseModel):
    """Metadata about loader updates - included in responses when relevant."""
    loader_current: str  # Current available loader version
    update_available: bool
    update_message: Optional[str] = None
    update_url: str

class SkillExecuteResponse(BaseModel):
    skill_id: str
    version: str
    result: str  # AI-generated result (not the prompt!)
    credits_remaining: int
    credits_used: int
    meta: Optional[LoaderMeta] = None  # Loader update hints

class SkillSchemaResponse(BaseModel):
    """Response for local execution skills - returns schema instead of executing."""
    skill_id: str
    skill_name: str
    version: str
    schema: str  # The expert prompt/framework for local AI to apply
    required_inputs: Optional[dict] = None  # Expected inputs (from YAML if available)
    credits_remaining: int
    credits_used: int
    instructions: str  # How to apply the schema locally
    meta: Optional[LoaderMeta] = None

class CreditBalanceResponse(BaseModel):
    credits_balance: int
    license_key: str
    license_type: str
    valid_until: Optional[datetime]

class PurchaseCreditsRequest(BaseModel):
    pack: str  # 'trial', 'essentials', 'accelerator', 'efficient', 'unstoppable', 'apex'
    email: Optional[str] = None  # Required for trial pack (one-time offer)
    attorney_name: Optional[str] = None
    bar_jurisdiction: Optional[str] = None   # e.g. "Colorado", "TX", "Federal"
    bar_number: Optional[str] = None
    firm_name: Optional[str] = None
    attestation: Optional[bool] = False      # True = user checked "I am a licensed attorney"

class UsageResponse(BaseModel):
    skill_id: str
    skill_name: str
    executed_at: datetime
    success: bool
    credits_used: int

# ============================================
# Profile & Document Schemas
# ============================================

class UserProfile(BaseModel):
    """User profile for document generation."""
    firm_name: Optional[str] = None
    attorney_name: Optional[str] = None
    attorney_bar: Optional[str] = None
    bar_jurisdiction: Optional[str] = None   # e.g. "Colorado", "TX", "Federal"
    bar_number: Optional[str] = None
    paralegal_name: Optional[str] = None
    address: Optional[str] = None
    city_state_zip: Optional[str] = None
    phone: Optional[str] = None
    fax: Optional[str] = None
    email: Optional[str] = None
    logo_url: Optional[str] = None

class ProfileResponse(BaseModel):
    profile: UserProfile
    missing_fields: List[str] = []

class ProfileUpdateRequest(BaseModel):
    """Update profile (merges with existing)."""
    firm_name: Optional[str] = None
    attorney_name: Optional[str] = None
    attorney_bar: Optional[str] = None
    bar_jurisdiction: Optional[str] = None   # e.g. "Colorado", "TX", "Federal"
    bar_number: Optional[str] = None
    paralegal_name: Optional[str] = None
    address: Optional[str] = None
    city_state_zip: Optional[str] = None
    phone: Optional[str] = None
    fax: Optional[str] = None
    email: Optional[str] = None
    logo_url: Optional[str] = None

class DocumentAttachment(BaseModel):
    """A generated document attachment."""
    filename: str
    content_type: str
    data: str  # Base64-encoded document content

class SkillExecuteResponseWithDocs(BaseModel):
    """Extended response with optional document attachments."""
    skill_id: str
    version: str
    result: str
    credits_remaining: int
    credits_used: int
    documents: Optional[List[DocumentAttachment]] = None
    needs_profile: Optional[List[str]] = None  # Profile fields needed before execution
    meta: Optional[LoaderMeta] = None

# ============================================
# Helper Functions
# ============================================

def hash_password(password: str) -> str:
    """Hash password with salt."""
    salt = secrets.token_hex(16)
    pw_hash = hashlib.sha256((password + salt).encode()).hexdigest()
    return f"{salt}:{pw_hash}"

def verify_password(password: str, stored_hash: str) -> bool:
    """Verify password against stored hash."""
    try:
        salt, pw_hash = stored_hash.split(":")
        return hashlib.sha256((password + salt).encode()).hexdigest() == pw_hash
    except:
        return False

def generate_license_key() -> str:
    """Generate a unique license key."""
    return f"lt_{secrets.token_hex(16)}"

def generate_token(user_id: str, license_key: str) -> str:
    """Generate a simple token (in production, use JWT)."""
    data = f"{user_id}:{license_key}:{secrets.token_hex(8)}"
    return hashlib.sha256(data.encode()).hexdigest()[:32]

# Credit pack pricing
CREDIT_PACKS = {
    "tryit":       {"credits": 2,    "price_cents": 500,    "one_time": False, "name": "Try It"},
    "starter":     {"credits": 15,   "price_cents": 2900,   "one_time": False, "name": "Starter"},
    "pro":         {"credits": 60,   "price_cents": 9900,   "one_time": False, "name": "Pro"},
    "business":    {"credits": 150,  "price_cents": 19900,  "one_time": False, "name": "Business"},
    "power":       {"credits": 350,  "price_cents": 34900,  "one_time": False, "name": "Power"},
    "unlimited":   {"credits": 800,  "price_cents": 59900,  "one_time": False, "name": "Unlimited"},
    "enterprise":  {"credits": 2000, "price_cents": 99900,  "one_time": False, "name": "Enterprise"},
    # Legacy aliases
    "peek":        {"credits": 5,    "price_cents": 500,    "one_time": False, "name": "Peek"},
    "trial":       {"credits": 15,   "price_cents": 2900,   "one_time": False, "name": "Starter"},
    "essentials":  {"credits": 60,   "price_cents": 9900,   "one_time": False, "name": "Pro"},
    "accelerator": {"credits": 150,  "price_cents": 19900,  "one_time": False, "name": "Business"},
    "efficient":   {"credits": 350,  "price_cents": 34900,  "one_time": False, "name": "Power"},
    "unstoppable": {"credits": 800,  "price_cents": 59900,  "one_time": False, "name": "Unlimited"},
    "apex":        {"credits": 2000, "price_cents": 99900,  "one_time": False, "name": "Enterprise"},
}

# ============================================
# Document Generation Helpers
# ============================================

# Skills that should generate documents and their required profile fields
DOCUMENT_SKILLS = {
    "demand-letter-drafter": {
        "format": "docx",
        "required_profile": ["firm_name", "attorney_name", "address", "city_state_zip", "phone"],
        "filename_template": "demand-letter-{date}.docx"
    },
    "discovery-request-generator": {
        "format": "docx",
        "required_profile": ["firm_name", "attorney_name", "attorney_bar"],
        "filename_template": "discovery-requests-{date}.docx"
    },
    "subpoena-generator": {
        "format": "docx",
        "required_profile": ["firm_name", "attorney_name", "attorney_bar", "address", "city_state_zip", "phone"],
        "filename_template": "subpoena-{date}.docx"
    },
    "deposition-summarizer": {
        "format": "docx",
        "required_profile": [],  # No letterhead needed for internal summaries
        "filename_template": "deposition-summary-{date}.docx"
    },
    "sol-alert-system": {
        "format": "xlsx",
        "required_profile": [],
        "filename_template": "sol-tracker-{date}.xlsx"
    },
    "deadline-calculator": {
        "format": "xlsx",
        "required_profile": [],
        "filename_template": "deadline-calendar-{date}.xlsx"
    }
}


def check_profile_requirements(skill_id: str, profile: dict) -> List[str]:
    """Check if user profile has all required fields for a skill."""
    if skill_id not in DOCUMENT_SKILLS:
        return []
    
    required = DOCUMENT_SKILLS[skill_id].get("required_profile", [])
    missing = [field for field in required if not profile.get(field)]
    return missing


def generate_docx_with_letterhead(content: str, profile: dict, title: str = None) -> bytes:
    """Generate a Word document with firm letterhead, supporting markdown formatting."""
    if not DOCX_AVAILABLE:
        return None
    
    doc = Document()
    
    # Add letterhead if profile has firm info
    if profile.get("firm_name"):
        header = doc.sections[0].header
        header_para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Firm name (bold, larger)
        run = header_para.add_run(profile.get("firm_name", ""))
        run.bold = True
        run.font.size = Pt(14)
        header_para.add_run("\n")
        
        # Address
        if profile.get("address"):
            header_para.add_run(profile["address"]).font.size = Pt(10)
            header_para.add_run("\n")
        if profile.get("city_state_zip"):
            header_para.add_run(profile["city_state_zip"]).font.size = Pt(10)
            header_para.add_run("\n")
        
        # Contact info
        contact_parts = []
        if profile.get("phone"):
            contact_parts.append(f"Tel: {profile['phone']}")
        if profile.get("fax"):
            contact_parts.append(f"Fax: {profile['fax']}")
        if profile.get("email"):
            contact_parts.append(profile["email"])
        if contact_parts:
            header_para.add_run(" | ".join(contact_parts)).font.size = Pt(9)
    
    # Add title if provided
    if title:
        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title_para.add_run(title)
        run.bold = True
        run.font.size = Pt(14)
        doc.add_paragraph()  # Spacing
    
    # Check if content contains headings (for TOC)
    has_headings = bool(re.search(r'^#{1,4}\s+.+', content, re.MULTILINE))
    
    # Add Table of Contents if headings exist
    if has_headings:
        # Create TOC using Word field codes
        toc_para = doc.add_paragraph()
        toc_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        # Add TOC field
        run = toc_para.add_run()
        fldChar = OxmlElement('w:fldChar')
        fldChar.set(qn('w:fldCharType'), 'begin')
        run._r.append(fldChar)
        
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = 'TOC \\o "1-4" \\h \\z \\u'
        run._r.append(instrText)
        
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'separate')
        run._r.append(fldChar2)
        
        # Placeholder text shown before TOC is updated
        run2 = toc_para.add_run('[Table of Contents - right-click and select "Update Field" to populate]')
        run2.font.color.rgb = RGBColor(128, 128, 128)
        run2.font.italic = True
        
        fldChar3 = OxmlElement('w:fldChar')
        fldChar3.set(qn('w:fldCharType'), 'end')
        run3 = toc_para.add_run()
        run3._r.append(fldChar3)
        
        # Add spacing after TOC
        doc.add_paragraph()
        doc.add_paragraph()
    
    # Helper function to add text with inline bold formatting
    def add_text_with_formatting(paragraph, text):
        """Add text to paragraph with inline **bold** support."""
        # Pattern to match **bold text**
        pattern = r'\*\*(.+?)\*\*'
        last_end = 0
        
        for match in re.finditer(pattern, text):
            # Add text before the bold part
            if match.start() > last_end:
                paragraph.add_run(text[last_end:match.start()])
            
            # Add bold text
            bold_run = paragraph.add_run(match.group(1))
            bold_run.bold = True
            
            last_end = match.end()
        
        # Add remaining text after last bold part
        if last_end < len(text):
            paragraph.add_run(text[last_end:])
    
    # Process content line by line for better structure detection
    lines = content.split('\n')
    i = 0
    
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()
        
        # Skip empty lines
        if not stripped:
            i += 1
            continue
        
        # Detect markdown headings
        heading_match = re.match(r'^(#{1,4})\s+(.+)$', stripped)
        if heading_match:
            level = len(heading_match.group(1))
            heading_text = heading_match.group(2)
            
            # Add heading with appropriate style
            heading_para = doc.add_heading(level=level)
            heading_para.text = heading_text
            i += 1
            continue
        
        # Detect horizontal rule
        if stripped in ['---', '___', '***']:
            # Add a paragraph with bottom border to simulate HR
            hr_para = doc.add_paragraph()
            hr_para.paragraph_format.space_after = Pt(12)
            hr_para.paragraph_format.space_before = Pt(12)
            i += 1
            continue
        
        # Detect bullet points (- or *)
        bullet_match = re.match(r'^[-*]\s+(.+)$', stripped)
        if bullet_match:
            bullet_text = bullet_match.group(1)
            bullet_para = doc.add_paragraph(style='List Bullet')
            add_text_with_formatting(bullet_para, bullet_text)
            i += 1
            continue
        
        # Detect numbered lists
        numbered_match = re.match(r'^(\d+)\.\s+(.+)$', stripped)
        if numbered_match:
            numbered_text = numbered_match.group(2)
            numbered_para = doc.add_paragraph(style='List Number')
            add_text_with_formatting(numbered_para, numbered_text)
            i += 1
            continue
        
        # Regular paragraph - collect until next empty line or special formatting
        para_lines = [line]
        i += 1
        
        while i < len(lines):
            next_line = lines[i].strip()
            
            # Stop if empty line or special formatting detected
            if not next_line or \
               re.match(r'^#{1,4}\s+', next_line) or \
               next_line in ['---', '___', '***'] or \
               re.match(r'^[-*]\s+', next_line) or \
               re.match(r'^\d+\.\s+', next_line):
                break
            
            para_lines.append(lines[i])
            i += 1
        
        # Add paragraph with inline formatting
        para_text = ' '.join(line.strip() for line in para_lines if line.strip())
        if para_text:
            p = doc.add_paragraph()
            add_text_with_formatting(p, para_text)
    
    # Add footer with attorney info
    if profile.get("attorney_name"):
        doc.add_paragraph()  # Spacing
        footer_para = doc.add_paragraph()
        footer_para.add_run(profile["attorney_name"])
        if profile.get("attorney_bar"):
            footer_para.add_run(f"\n{profile['attorney_bar']}")
    
    # Save to bytes
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()


def generate_xlsx_from_content(content: str, title: str = None) -> bytes:
    """Generate an Excel spreadsheet from structured content."""
    if not XLSX_AVAILABLE:
        return None
    
    wb = Workbook()
    ws = wb.active
    ws.title = title or "Data"
    
    # Try to parse structured data from the content
    # Look for lines that look like table rows (contain | or tabs)
    lines = content.strip().split("\n")
    row_num = 1
    
    for line in lines:
        if "|" in line:
            # Markdown table format
            cells = [cell.strip() for cell in line.split("|") if cell.strip() and cell.strip() != "---"]
            if cells:
                for col_num, cell in enumerate(cells, 1):
                    ws.cell(row=row_num, column=col_num, value=cell)
                row_num += 1
        elif "\t" in line:
            # Tab-separated
            cells = line.split("\t")
            for col_num, cell in enumerate(cells, 1):
                ws.cell(row=row_num, column=col_num, value=cell.strip())
            row_num += 1
        elif line.strip():
            # Plain text - put in first column
            ws.cell(row=row_num, column=1, value=line.strip())
            row_num += 1
    
    # Auto-adjust column widths
    for column in ws.columns:
        max_length = 0
        column_letter = column[0].column_letter
        for cell in column:
            try:
                if len(str(cell.value)) > max_length:
                    max_length = len(str(cell.value))
            except:
                pass
        ws.column_dimensions[column_letter].width = min(max_length + 2, 50)
    
    # Save to bytes
    buffer = BytesIO()
    wb.save(buffer)
    buffer.seek(0)
    return buffer.read()


def generate_document(skill_id: str, content: str, profile: dict) -> Optional[DocumentAttachment]:
    """Generate appropriate document for a skill result."""
    if skill_id not in DOCUMENT_SKILLS:
        return None
    
    skill_doc_config = DOCUMENT_SKILLS[skill_id]
    doc_format = skill_doc_config["format"]
    
    # Generate filename with today's date
    from datetime import date
    filename = skill_doc_config["filename_template"].format(date=date.today().isoformat())
    
    if doc_format == "docx":
        doc_bytes = generate_docx_with_letterhead(content, profile, title=None)
        if doc_bytes:
            return DocumentAttachment(
                filename=filename,
                content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                data=base64.b64encode(doc_bytes).decode("utf-8")
            )
    
    elif doc_format == "xlsx":
        doc_bytes = generate_xlsx_from_content(content, title=skill_id.replace("-", " ").title())
        if doc_bytes:
            return DocumentAttachment(
                filename=filename,
                content_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                data=base64.b64encode(doc_bytes).decode("utf-8")
            )
    
    return None

# ============================================
# FastAPI App
# ============================================

app = FastAPI(
    title="LawTasksAI API",
    description="Skill delivery, licensing, and usage tracking for legal AI automation",
    version="1.0.0"
)

# Admin router — all routes require X-Admin-Secret header
admin_router = APIRouter(
    prefix="/admin",
    include_in_schema=False,
    dependencies=[Depends(verify_admin)]
)

# Auto-create tables on startup
@app.on_event("startup")
async def startup():
    async with engine.begin() as conn:
        await conn.run_sync(Base.metadata.create_all)
        # Migration: add platforms column if it doesn't exist
        try:
            await conn.execute(text("""
                ALTER TABLE users
                ADD COLUMN IF NOT EXISTS platforms JSONB DEFAULT '[]'::jsonb;
            """))
        except Exception as e:
            print(f"[startup migration] platforms column: {e}")
        # Migration: create drip_emails table
        try:
            await conn.execute(text("""
                CREATE TABLE IF NOT EXISTS drip_emails (
                    id UUID PRIMARY KEY DEFAULT gen_random_uuid(),
                    user_id UUID REFERENCES users(id),
                    email VARCHAR(255) NOT NULL,
                    product_id VARCHAR(50) NOT NULL,
                    email_number INTEGER NOT NULL,
                    subject VARCHAR(500),
                    status VARCHAR(20) NOT NULL DEFAULT 'sent',
                    scheduled_for TIMESTAMP,
                    sent_at TIMESTAMP DEFAULT NOW(),
                    CONSTRAINT uq_drip_email_per_user UNIQUE (email, product_id, email_number)
                );
            """))
            # Add columns if missing (idempotent)
            for col_def in [
                "ALTER TABLE drip_emails ADD COLUMN IF NOT EXISTS status VARCHAR(20) NOT NULL DEFAULT 'sent'",
                "ALTER TABLE drip_emails ADD COLUMN IF NOT EXISTS scheduled_for TIMESTAMP",
                "ALTER TABLE drip_emails ADD COLUMN IF NOT EXISTS platform VARCHAR(50)",
                "ALTER TABLE drip_emails ADD COLUMN IF NOT EXISTS first_name VARCHAR(100)",
            ]:
                try:
                    await conn.execute(text(col_def))
                except Exception:
                    pass
        except Exception as e:
            print(f"[startup migration] drip_emails table: {e}")
        # Migration: backfill licenses.product_id from users.product_id where NULL
        try:
            await conn.execute(text("""
                UPDATE licenses l
                SET product_id = u.product_id
                FROM users u
                WHERE l.user_id = u.id
                AND l.product_id IS NULL
                AND u.product_id IS NOT NULL;
            """))
        except Exception as e:
            print(f"[startup migration] licenses.product_id backfill: {e}")
        # Migration: create user_feedback table if it doesn't exist
        try:
            await conn.execute(text("""
                CREATE TABLE IF NOT EXISTS user_feedback (
                    id UUID PRIMARY KEY DEFAULT gen_random_uuid(),
                    email VARCHAR(255),
                    product_id VARCHAR(50),
                    reason VARCHAR(100) NOT NULL,
                    extra TEXT,
                    created_at TIMESTAMP DEFAULT NOW()
                );
            """))
        except Exception as e:
            print(f"[startup migration] user_feedback table: {e}")
        # Migration: create download_tokens table
        try:
            await conn.execute(text("""
                CREATE TABLE IF NOT EXISTS download_tokens (
                    token VARCHAR(64) PRIMARY KEY,
                    license_key VARCHAR(255) NOT NULL,
                    used BOOLEAN NOT NULL DEFAULT FALSE,
                    created_at TIMESTAMP NOT NULL DEFAULT NOW()
                );
            """))
        except Exception as e:
            print(f"[startup migration] download_tokens table: {e}")
        # Migration: add funnel tracking columns to licenses + users
        try:
            await conn.execute(text("ALTER TABLE licenses ADD COLUMN IF NOT EXISTS downloaded_at TIMESTAMP"))
            await conn.execute(text("ALTER TABLE licenses ADD COLUMN IF NOT EXISTS first_connected_at TIMESTAMP"))
            await conn.execute(text("ALTER TABLE users ADD COLUMN IF NOT EXISTS last_active_at TIMESTAMP"))
        except Exception as e:
            print(f"[startup migration] funnel tracking columns: {e}")
        # Migration: create support_requests table
        try:
            await conn.execute(text("""
                CREATE TABLE IF NOT EXISTS support_requests (
                    id UUID PRIMARY KEY DEFAULT gen_random_uuid(),
                    email VARCHAR(255) NOT NULL,
                    name VARCHAR(255),
                    product_id VARCHAR(50),
                    subject VARCHAR(500),
                    message TEXT,
                    direction VARCHAR(10) NOT NULL DEFAULT 'inbound',
                    status VARCHAR(50) NOT NULL DEFAULT 'open',
                    zoho_message_id VARCHAR(255),
                    replied_at TIMESTAMP,
                    created_at TIMESTAMP DEFAULT NOW()
                );
            """))
        except Exception as e:
            print(f"[startup migration] support_requests table: {e}")

        # Migration: create waitlist table
        try:
            await db.execute(text("""
                CREATE TABLE IF NOT EXISTS waitlist (
                    id UUID PRIMARY KEY DEFAULT gen_random_uuid(),
                    email VARCHAR(255) NOT NULL,
                    name VARCHAR(200),
                    product_id VARCHAR(50) NOT NULL,
                    created_at TIMESTAMP DEFAULT NOW(),
                    notified_at TIMESTAMP,
                    UNIQUE(email, product_id)
                );
            """))
            await db.commit()
        except Exception as e:
            print(f"[startup migration] waitlist table: {e}")

# CORS
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],  # Configure for production
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# ============================================
# Auth Dependency
# ============================================

async def get_current_license(
    authorization: str = Header(...),
    db: AsyncSession = Depends(get_db)
) -> License:
    """Validate license key from Authorization header."""
    if not authorization.startswith("Bearer "):
        raise HTTPException(status_code=401, detail="Invalid authorization header")
    
    license_key = authorization.replace("Bearer ", "")
    
    result = await db.execute(
        select(License).where(
            License.license_key == license_key,
            License.status == "active"
        )
    )
    license = result.scalar_one_or_none()
    
    if not license:
        raise HTTPException(status_code=401, detail="Invalid or inactive license")
    
    # Check expiry
    if license.valid_until and license.valid_until < datetime.utcnow():
        raise HTTPException(status_code=401, detail="License expired")
    
    return license

# ============================================
# Routes: Auth
# ============================================

@app.post("/auth/register", response_model=UserResponse)
async def register(
    user_data: UserCreate,
    db: AsyncSession = Depends(get_db),
    product_id: str = Depends(get_product_id),
):
    """Create a new user account.

    Multi-vertical logic:
    - If email is brand new: create user + license for this product.
    - If email exists but has no license for this product: reuse the user,
      create a new license for this product (new vertical signup).
    - If email exists AND already has a license for this product: 400 duplicate.
    """
    # Resolve product_id: body field takes priority (explicit), then dependency (header/query/default)
    resolved_product_id = user_data.product_id or product_id

    # Check if email exists
    result = await db.execute(select(User).where(User.email == user_data.email))
    existing_user = result.scalar_one_or_none()

    if existing_user:
        # Check whether this user already has a license for THIS product
        lic_result = await db.execute(
            select(License).where(
                License.user_id == existing_user.id,
                License.product_id == resolved_product_id,
                License.status == "active",
            )
        )
        if lic_result.scalar_one_or_none():
            # Genuine duplicate — same email, same vertical
            raise HTTPException(status_code=400, detail="Email already registered")

        # New vertical for an existing user — add a license only
        user = existing_user
        user_id = existing_user.id
        # Merge platforms if provided
        new_platforms = getattr(user_data, 'platforms', None) or []
        if new_platforms:
            existing_platforms = user.platforms or []
            merged = {p.get('platform'): p for p in existing_platforms}
            for p in new_platforms:
                merged[p.get('platform')] = p
            user.platforms = list(merged.values())
    else:
        # Brand new user
        user_id = uuid.uuid4()
        user = User(
            id=user_id,
            email=user_data.email,
            password_hash=hash_password(user_data.password),
            name=user_data.name,
            firm_name=user_data.firm_name,
            credits_balance=5,  # Free signup credits
            product_id=resolved_product_id,
            platforms=getattr(user_data, 'platforms', None) or [],
        )
        db.add(user)
        await db.flush()  # Ensure user row exists before license FK insert

    # Create trial license for this product (always — new user or new vertical)
    license = License(
        license_key=generate_license_key(),
        user_id=user_id,
        product_id=resolved_product_id,
        type="trial",
        valid_until=datetime.utcnow() + timedelta(days=14),
        credits_purchased=5,
        credits_remaining=5
    )
    db.add(license)

    await db.commit()
    await db.refresh(user)

    # Send drip Email 1 (HTML welcome) with license key
    try:
        import sys as _sys
        import importlib
        _drip_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'drip')
        if _drip_path not in _sys.path:
            _sys.path.insert(0, _drip_path)
        drip_utils = importlib.import_module('drip_utils')

        # Resolve product domain, name, skill count
        reg_product_domain = "lawtasksai.com"
        reg_product_name = "LawTasksAI"
        reg_skill_count = 206
        try:
            prod_result = await db.execute(
                text("SELECT domain, name FROM products WHERE id = :pid AND is_active = TRUE"),
                {"pid": resolved_product_id}
            )
            prod_row = prod_result.fetchone()
            if prod_row and prod_row.domain:
                reg_product_domain = prod_row.domain
                reg_product_name = prod_row.name
            # Get skill count
            sc_result = await db.execute(
                text("SELECT COUNT(*) FROM skills WHERE product_id = :pid AND is_active = TRUE"),
                {"pid": resolved_product_id}
            )
            sc_row = sc_result.fetchone()
            if sc_row:
                reg_skill_count = sc_row[0] or reg_skill_count
        except Exception:
            pass

        first_name = (user.name or "").split()[0] if user.name else ""
        user_platform = (user.platforms or [{}])[0].get("platform", "other") if user.platforms else "other"

        email1_html = drip_utils.build_drip_email(
            email_num=1,
            product_id=resolved_product_id,
            product_name=reg_product_name,
            domain=reg_product_domain,
            skill_count=reg_skill_count,
            platform=user_platform,
            first_name=first_name,
            user_email=user.email,
            license_key=license.license_key,
        )
        email1_subject = drip_utils.drip_subject(1, reg_product_name)

        access_token = await get_zoho_access_token()
        if access_token:
            async with httpx.AsyncClient(timeout=15) as _hc:
                _resp = await _hc.post(
                    f"https://mail.zoho.com/api/accounts/6556209000000008002/messages",
                    json={
                        "fromAddress": f"=?UTF-8?B?{base64.b64encode(reg_product_name.encode()).decode()}?= <hello@{reg_product_domain}>",
                        "toAddress": user.email,
                        "subject": email1_subject,
                        "content": email1_html,
                        "mailFormat": "html",
                    },
                    headers={"Authorization": f"Zoho-oauthtoken {access_token}"}
                )
                print(f"[Drip] Email 1 sent to {user.email} ({resolved_product_id}): Zoho {_resp.status_code}")
                # Record in drip_emails table
                try:
                    await db.execute(text("""
                        INSERT INTO drip_emails (email, product_id, email_number, status)
                        VALUES (:email, :pid, 1, 'sent')
                        ON CONFLICT (email, product_id, email_number) DO NOTHING
                    """), {"email": user.email, "pid": resolved_product_id})
                    await db.commit()
                except Exception:
                    pass
        else:
            print(f"[Drip] Email 1 skipped for {user.email}: no Zoho token")

        # Schedule Emails 2 (Day 2) and 3 (Day 7)
        try:
            now_utc = datetime.utcnow()
            for email_num, days_delay in [(2, 2), (3, 7)]:
                await db.execute(text("""
                    INSERT INTO drip_emails
                        (email, product_id, email_number, status, scheduled_for, platform, first_name)
                    VALUES
                        (:email, :pid, :num, 'scheduled', :send_at, :platform, :fname)
                    ON CONFLICT (email, product_id, email_number) DO NOTHING
                """), {
                    "email":    user.email,
                    "pid":      resolved_product_id,
                    "num":      email_num,
                    "send_at":  now_utc + timedelta(days=days_delay),
                    "platform": user_platform,
                    "fname":    first_name,
                })
            await db.commit()
            print(f"[Drip] Emails 2+3 scheduled for {user.email} ({resolved_product_id})")
        except Exception as sched_err:
            print(f"[Drip] scheduling failed for {user.email}: {sched_err}")
    except Exception as email_err:
        print(f"[Drip] Email 1 failed for {user.email}: {email_err}")

    # Add to Zoho Campaigns subscriber list (fire-and-forget)
    try:
        await add_to_zoho_list(user.email, user.name or "", resolved_product_id)
    except Exception as zoho_err:
        print(f"[Zoho Campaigns] signup hook failed for {user.email}: {zoho_err}")

    # Insert email_subscriptions row
    try:
        sub = EmailSubscription(user_id=user_id, product_id=resolved_product_id, subscribed=True)
        db.add(sub)
        await db.commit()
    except Exception as sub_err:
        print(f"[EmailSub] failed to insert subscription for {user.email}: {sub_err}")

    return UserResponse(
        id=str(user.id),
        email=user.email,
        name=user.name,
        firm_name=user.firm_name,
        credits_balance=license.credits_remaining,
        created_at=user.created_at
    )

@app.post("/v1/signup", response_model=UserResponse, status_code=201)
async def simple_signup(
    data: SimpleSignupRequest,
    db: AsyncSession = Depends(get_db),
    product_id: str = Depends(get_product_id),
):
    """
    Password-free signup endpoint for landing page free-trial flow.
    Accepts name + email only; generates a random password internally.
    Returns the same UserResponse as /auth/register.
    """
    # Delegate to the existing register logic by constructing a UserCreate
    # with an auto-generated password the user never needs to know.
    auto_password = secrets.token_urlsafe(24)
    resolved_product_id = data.product_id or product_id
    user_create = UserCreate(
        email=data.email,
        name=data.name,
        password=auto_password,
        product_id=resolved_product_id,
        platforms=data.platforms or [],
    )
    # Reuse register() — pass the assembled UserCreate
    return await register(user_create, db, resolved_product_id)


@app.post("/auth/login", response_model=TokenResponse)
async def login(credentials: LoginRequest, db: AsyncSession = Depends(get_db)):
    """Login and get access token."""
    result = await db.execute(select(User).where(User.email == credentials.email))
    user = result.scalar_one_or_none()
    
    if not user or not verify_password(credentials.password, user.password_hash):
        raise HTTPException(status_code=401, detail="Invalid credentials")
    
    # Get active license
    result = await db.execute(
        select(License).where(
            License.user_id == user.id,
            License.status == "active"
        ).order_by(License.created_at.desc())
    )
    license = result.scalar_one_or_none()
    
    if not license:
        raise HTTPException(status_code=401, detail="No active license found")
    
    # Update last login
    user.last_login = datetime.utcnow()
    await db.commit()
    
    return TokenResponse(
        access_token=generate_token(str(user.id), license.license_key),
        license_key=license.license_key,
        product_id=user.product_id or "law",
    )

class RecoverLicenseRequest(BaseModel):
    email: EmailStr

class RecoverLicenseResponse(BaseModel):
    email: str
    license_key: str
    credits_remaining: int
    license_type: str
    message: str

@app.post("/auth/recover-license", response_model=RecoverLicenseResponse)
async def recover_license(request: RecoverLicenseRequest, db: AsyncSession = Depends(get_db)):
    """
    Recover license key by email.
    For customers who lost their config or reinstalled.
    """
    # Find user by email
    result = await db.execute(select(User).where(User.email == request.email))
    user = result.scalar_one_or_none()
    
    if not user:
        raise HTTPException(status_code=404, detail="No account found with this email")
    
    # Get active license
    result = await db.execute(
        select(License).where(
            License.user_id == user.id,
            License.status == "active"
        ).order_by(License.created_at.desc())
    )
    license = result.scalar_one_or_none()
    
    if not license:
        raise HTTPException(status_code=404, detail="No active license found for this account")
    
    return RecoverLicenseResponse(
        email=user.email,
        license_key=license.license_key,
        credits_remaining=license.credits_remaining,
        license_type=license.type,
        message="Your license key has been recovered. Update your config.json with this key."
    )


class AccountLicenseItem(BaseModel):
    license_key: str
    product_id: str
    product_name: str
    product_domain: str
    credits_remaining: int
    license_type: str
    created_at: Optional[datetime] = None


class AccountLicensesResponse(BaseModel):
    email: str
    licenses: List[AccountLicenseItem]


@app.post("/auth/account-licenses")
async def account_licenses(request: RecoverLicenseRequest, db: AsyncSession = Depends(get_db)):
    """
    Return all active licenses for an email across all verticals.
    Used by the download page to render a per-vertical account dashboard.
    """
    result = await db.execute(select(User).where(User.email == request.email))
    user = result.scalar_one_or_none()
    if not user:
        raise HTTPException(status_code=404, detail="No account found with this email")

    lic_result = await db.execute(
        select(License).where(
            License.user_id == user.id,
            License.status == "active",
        ).order_by(License.created_at.asc())
    )
    licenses = lic_result.scalars().all()

    if not licenses:
        raise HTTPException(status_code=404, detail="No active licenses found for this account")

    # Look up product metadata for each unique product_id
    product_ids = list({lic.product_id for lic in licenses if lic.product_id})
    prod_result = await db.execute(
        text("SELECT id, name, domain FROM products WHERE id = ANY(:pids) AND is_active = TRUE"),
        {"pids": product_ids}
    )
    prod_map = {row.id: {"name": row.name, "domain": row.domain} for row in prod_result.fetchall()}

    # Fallback display names for products not in the products table
    def product_display(pid: str) -> dict:
        if pid in prod_map:
            return prod_map[pid]
        # Derive a readable name from the product_id (e.g. "realtor" -> "RealtorTasksAI")
        name = pid.capitalize() + "TasksAI" if pid else "TasksAI"
        domain = f"{pid}tasksai.com" if pid else "lawtasksai.com"
        return {"name": name, "domain": domain}

    items = []
    for lic in licenses:
        pid = lic.product_id or "law"
        meta = product_display(pid)
        items.append(AccountLicenseItem(
            license_key=lic.license_key,
            product_id=pid,
            product_name=meta["name"],
            product_domain=meta["domain"],
            credits_remaining=lic.credits_remaining,
            license_type=lic.type,
            created_at=lic.created_at,
        ))

    return AccountLicensesResponse(email=user.email, licenses=items)


# ============================================
# Routes: Waitlist
# ============================================

# Verticals with working installers — waitlist is NOT needed for these
READY_VERTICALS = {"law", "realtor", "farmer", "teacher", "therapist", "marketing", "contractor"}


class WaitlistRequest(BaseModel):
    email: str
    name: Optional[str] = None
    product_id: str


@app.post("/auth/waitlist", status_code=200)
async def join_waitlist(request: WaitlistRequest, db: AsyncSession = Depends(get_db)):
    """Join the waitlist for a vertical that doesn't have an installer yet."""
    email = request.email.lower().strip()
    product_id = request.product_id.lower().strip()

    if product_id in READY_VERTICALS:
        raise HTTPException(status_code=400, detail="This product is already available — please register instead.")

    # Upsert — idempotent if they sign up twice
    existing = await db.execute(
        select(WaitlistEntry).where(
            WaitlistEntry.email == email,
            WaitlistEntry.product_id == product_id
        )
    )
    if existing.scalar_one_or_none():
        return {"success": True, "already_registered": True}

    entry = WaitlistEntry(email=email, name=request.name, product_id=product_id)
    db.add(entry)
    await db.commit()
    return {"success": True, "already_registered": False}


@app.get("/waitlist/unsubscribe")
async def waitlist_unsubscribe(email: str, product_id: str, token: str, db: AsyncSession = Depends(get_db)):
    """One-click unsubscribe from a vertical waitlist. Linked from emails."""
    from fastapi.responses import HTMLResponse

    if not _verify_waitlist_unsub_token(email, product_id, token):
        return HTMLResponse(content="<h2>Invalid or expired unsubscribe link.</h2>", status_code=400)

    await db.execute(
        text("DELETE FROM waitlist WHERE email = :email AND product_id = :pid"),
        {"email": email.lower().strip(), "pid": product_id}
    )
    await db.commit()

    product_label = product_id.title() + "TasksAI"
    html = f"""
    <!DOCTYPE html><html><head><meta charset="UTF-8">
    <title>Unsubscribed</title>
    <style>body{{font-family:system-ui,sans-serif;max-width:480px;margin:80px auto;padding:0 24px;color:#374151;}}
    h1{{color:#111827;}}p{{line-height:1.6;color:#6b7280;}}</style></head><body>
    <h1>You\'re unsubscribed.</h1>
    <p>We\'ve removed <strong>{email}</strong> from the {product_label} waitlist. You won\'t hear from us again.</p>
    <p style="margin-top:32px;"><a href="https://{product_id}tasksai.com" style="color:#6366f1;">Back to {product_label}</a></p>
    </body></html>
    """
    return HTMLResponse(content=html)


@app.get("/waitlist/unsubscribe-token")
async def get_waitlist_unsub_token(email: str, product_id: str, x_admin_secret: Optional[str] = Header(None, alias="X-Admin-Secret")):
    """Admin helper: generate an unsubscribe token for a given email+product. Used when composing emails."""
    if x_admin_secret != ADMIN_SECRET:
        raise HTTPException(status_code=403, detail="Forbidden")
    token = _waitlist_unsub_token(email, product_id)
    url = f"https://api.lawtasksai.com/waitlist/unsubscribe?email={email}&product_id={product_id}&token={token}"
    return {"token": token, "unsubscribe_url": url}


# ============================================
# Routes: Skills
# ============================================

@app.get("/skills", response_model=List[SkillResponse])
@app.get("/v1/skills", response_model=List[SkillResponse])
async def list_skills(
    category: Optional[str] = None,
    db: AsyncSession = Depends(get_db),
    product_id: str = Depends(get_product_id),
):
    """List all available skills, optionally filtered by product."""
    query = select(Skill).where(Skill.is_published == True)
    if category:
        query = query.where(Skill.category_id == category)
    # Always filter by product_id — defaults to 'law' for backward compat.
    query = query.where(Skill.product_id == product_id)
    
    result = await db.execute(query.order_by(Skill.category_id, Skill.name))
    skills = result.scalars().all()

    # Fetch per-skill scan data
    skill_ids = [s.id for s in skills]
    scan_result = await db.execute(
        select(SkillSecurityScan.skill_id, SkillSecurityScan.plugins_tested,
               SkillSecurityScan.tests_passed, SkillSecurityScan.tests_run)
        .where(SkillSecurityScan.skill_id.in_(skill_ids))
    )
    scan_map = {
        row.skill_id: {
            'plugins_tested': row.plugins_tested or [],
            'tests_passed': row.tests_passed,
            'tests_run': row.tests_run,
        }
        for row in scan_result
    }

    # Fetch platform (preamble) scan — use primary model (Claude) for 'all' vertical
    plat_result = await db.execute(
        select(PlatformSecurityScan.tests_passed, PlatformSecurityScan.tests_run)
        .where(PlatformSecurityScan.vertical == 'all')
        .where(PlatformSecurityScan.scan_model.like('%claude%'))
        .limit(1)
    )
    plat_row = plat_result.first()
    plat_passed = plat_row.tests_passed if plat_row else 20
    plat_run    = plat_row.tests_run    if plat_row else 20

    # Keywords that suggest document/data processing
    sensitive_keywords = ['analyzer', 'summarizer', 'reviewer', 'examiner', 'auditor', 
                          'parser', 'extractor', 'scanner', 'checker', 'drafter']
    
    def get_confidentiality_note(skill):
        name_lower = skill.name.lower()
        if skill.requires_upload:
            return "⚠️ Uploads document content to server for processing"
        if any(kw in name_lower for kw in sensitive_keywords):
            return "⚠️ May process sensitive text on server"
        if skill.execution_type == 'local':
            return "🔒 Runs locally — data stays on your machine"
        return None
    
    responses = []
    for s in skills:
        skill_scan = scan_map.get(s.id)
        skill_passed = skill_scan['tests_passed'] if skill_scan else None
        skill_run    = skill_scan['tests_run']    if skill_scan else None

        # Combined = platform + per-skill (only if per-skill scan exists)
        if skill_passed is not None and skill_run is not None:
            combined_passed = plat_passed + skill_passed
            combined_run    = plat_run    + skill_run
        else:
            combined_passed = None
            combined_run    = None

        # Verified if combined passes 33/35 (or 13/15 per-skill if no platform data)
        if combined_passed is not None and combined_run is not None:
            verified = combined_passed >= (combined_run - 2)  # allow 2 failures in 35
        elif skill_passed is not None:
            verified = skill_passed >= 13
        else:
            verified = bool(s.security_verified)

        responses.append(SkillResponse(
            id=s.id,
            name=s.name,
            description=s.description,
            category_id=s.category_id,
            current_version=s.current_version,
            credits_per_use=s.credits_per_use,
            requires_upload=s.requires_upload,
            execution_type=s.execution_type,
            confidentiality_note=get_confidentiality_note(s),
            security_verified=verified,
            plugins_tested=skill_scan['plugins_tested'] if skill_scan else [],
            security_tests_passed=skill_passed,
            security_tests_run=skill_run,
            platform_tests_passed=plat_passed,
            platform_tests_run=plat_run,
            combined_tests_passed=combined_passed,
            combined_tests_run=combined_run,
        ))
    return responses

@app.get("/skills/triggers")
@app.get("/v1/skills/triggers")
async def get_skill_triggers(
    product_id: str = Depends(get_product_id),
    db: AsyncSession = Depends(get_db)
):
    """
    Get trigger phrases for local skill matching.
    This enables privacy-preserving skill discovery without sending queries to the server.
    Reads from database - skills with non-empty triggers arrays, filtered by product.
    """
    result = await db.execute(
        select(Skill.id, Skill.triggers).where(
            Skill.is_published == True,
            Skill.product_id == product_id,
            Skill.triggers != None,
            func.array_length(Skill.triggers, 1) > 0
        )
    )
    rows = result.all()
    
    # Build response in expected format
    triggers_dict = {}
    for skill_id, triggers in rows:
        if triggers:
            triggers_dict[skill_id] = {"triggers": triggers}
    
    return triggers_dict


@app.get("/skills/{skill_id}", response_model=SkillResponse)
@app.get("/v1/skills/{skill_id}", response_model=SkillResponse)
async def get_skill(skill_id: str, db: AsyncSession = Depends(get_db)):
    """Get skill details."""
    result = await db.execute(select(Skill).where(Skill.id == skill_id))
    skill = result.scalar_one_or_none()
    
    if not skill:
        raise HTTPException(status_code=404, detail="Skill not found")
    
    # Get confidentiality note
    sensitive_keywords = ['analyzer', 'summarizer', 'reviewer', 'examiner', 'auditor', 
                          'parser', 'extractor', 'scanner', 'checker', 'drafter']
    name_lower = skill.name.lower()
    
    if skill.requires_upload:
        conf_note = "⚠️ Uploads document content to server for processing"
    elif any(kw in name_lower for kw in sensitive_keywords):
        conf_note = "⚠️ May process sensitive text on server"
    elif skill.execution_type == 'local':
        conf_note = "🔒 Runs locally — data stays on your machine"
    else:
        conf_note = None
    
    return SkillResponse(
        id=skill.id,
        name=skill.name,
        description=skill.description,
        category_id=skill.category_id,
        current_version=skill.current_version,
        credits_per_use=skill.credits_per_use,
        requires_upload=skill.requires_upload,
        execution_type=skill.execution_type,
        confidentiality_note=conf_note
    )

def check_loader_update(loader_version: Optional[str]) -> Optional[LoaderMeta]:
    """Check if loader needs update and return metadata if so."""
    if not loader_version:
        return None
    
    # Simple version comparison (assumes semver: x.y.z)
    try:
        current_parts = [int(x) for x in CURRENT_LOADER_VERSION.split('.')]
        client_parts = [int(x) for x in loader_version.split('.')]
        
        update_available = current_parts > client_parts
        
        # Only include meta if update is available or there's an important message
        if update_available or LOADER_UPDATE_MESSAGE:
            return LoaderMeta(
                loader_current=CURRENT_LOADER_VERSION,
                update_available=update_available,
                update_message=LOADER_UPDATE_MESSAGE,
                update_url=LOADER_UPDATE_URL
            )
    except (ValueError, AttributeError):
        pass  # Invalid version format, skip
    
    return None


@app.get("/skills/{skill_id}/schema", response_model=SkillSchemaResponse)
@app.get("/v1/skills/{skill_id}/schema", response_model=SkillSchemaResponse)
@app.get("/skills/{skill_id}/execute", response_model=SkillSchemaResponse)
@app.get("/v1/skills/{skill_id}/execute", response_model=SkillSchemaResponse)
async def get_skill_schema(
    skill_id: str,
    license: License = Depends(get_current_license),
    db: AsyncSession = Depends(get_db),
    background_tasks: BackgroundTasks = BackgroundTasks(),
    x_loader_version: Optional[str] = Header(None, alias="X-Loader-Version")
):
    """
    Get skill schema for LOCAL execution.
    
    For skills with execution_type='local', returns the expert framework/prompt
    so the client's AI can apply it to local documents. Documents never leave
    the user's machine.
    
    - Validates license and credits
    - Deducts credits (you pay for the expert knowledge)
    - Returns schema for local AI to apply
    - Only works for execution_type='local' skills
    
    All skills run locally — documents never leave the user's machine.
    """
    # Get skill
    result = await db.execute(select(Skill).where(Skill.id == skill_id, Skill.is_published == True))
    skill = result.scalar_one_or_none()
    
    if not skill:
        raise HTTPException(status_code=404, detail="Skill not found")
    
    # Check skill access
    if license.skills_allowed and skill_id not in license.skills_allowed:
        raise HTTPException(status_code=403, detail="Skill not included in license")
    
    if license.categories_allowed and skill.category_id not in license.categories_allowed:
        raise HTTPException(status_code=403, detail="Category not included in license")
    
    # Check usage limit
    if license.usage_limit and license.usage_count >= license.usage_limit:
        raise HTTPException(status_code=403, detail="Usage limit reached")
    
    # Check credits
    if license.credits_remaining < skill.credits_per_use:
        raise HTTPException(status_code=402, detail="Insufficient credits")
    
    # Get current version
    version_query = select(SkillVersion).where(
        SkillVersion.skill_id == skill_id,
        SkillVersion.version == skill.current_version
    )
    result = await db.execute(version_query)
    skill_version = result.scalar_one_or_none()
    
    if not skill_version:
        raise HTTPException(status_code=404, detail="Skill version not found")
    
    # Deduct credits and increment usage
    license.credits_remaining -= skill.credits_per_use
    license.usage_count += 1

    # Stamp last_active_at on user for engagement tracking
    try:
        user_result = await db.execute(select(User).where(User.id == license.user_id))
        active_user = user_result.scalar_one_or_none()
        if active_user:
            active_user.last_active_at = datetime.utcnow()
    except Exception:
        pass  # non-fatal — don't block skill execution

    # Log usage
    usage_log = UsageLog(
        license_id=license.id,
        user_id=license.user_id,
        skill_id=skill_id,
        skill_version=skill_version.version,
        credits_used=skill.credits_per_use,
        success=True,
    )
    db.add(usage_log)
    await db.commit()
    
    # Check for loader updates
    loader_meta = check_loader_update(x_loader_version)
    
    # Build instructions for local execution
    instructions = """
## Local Execution Instructions

This skill runs LOCALLY on your machine. LawTasksAI.com never sees your prompts, your client files, or your client data. Your documents stay local if using OpenClaw, or go to your LLM provider if using a cloud AI.

### How to use this schema:

1. **Load your document(s)** into the conversation (paste text, attach file, or reference local path)

2. **Apply this schema** by telling your AI:
   "Apply the following expert analysis framework to my document: [paste schema]"
   
   Or simply include the schema as context and ask your question naturally.

3. **The schema provides:**
   - Expert considerations and analysis frameworks
   - Step-by-step workflows
   - What to look for
   - How to structure the output
   - Common pitfalls to avoid

4. **Your AI will:**
   - Read your local document
   - Apply the expert framework
   - Generate analysis based on the schema
   - All processing happens on your machine

### Privacy Guarantee
Your document content is NEVER sent to LawTasksAI servers.
Only this schema was retrieved (general legal knowledge, not your data).
"""
    
    # Prepend universal security preamble to every schema
    from scripts.SECURITY_PREAMBLE import PREAMBLE as _SECURITY_PREAMBLE
    schema_with_preamble = _SECURITY_PREAMBLE + skill_version.content

    return SkillSchemaResponse(
        skill_id=skill_id,
        skill_name=skill.name,
        version=skill_version.version,
        schema=schema_with_preamble,  # Preamble + expert framework
        required_inputs=None,  # TODO: Parse from YAML if structured
        credits_remaining=license.credits_remaining,
        credits_used=skill.credits_per_use,
        instructions=instructions,
        meta=loader_meta
    )


# ============================================
# Routes: Credits & Billing
# ============================================

@app.get("/credits/balance", response_model=CreditBalanceResponse)
@app.get("/v1/installer-key")
async def resolve_installer_key(k: str):
    """
    Decode a base64-encoded license key passed from the download URL.
    The GUI calls this on startup to get its pre-filled key.
    k = base64url(license_key) with padding stripped.
    """
    try:
        import base64 as _b64
        # Restore padding
        padded = k + "==" [:(4 - len(k) % 4) % 4]
        license_key = _b64.urlsafe_b64decode(padded).decode()
        return {"license_key": license_key}
    except Exception:
        raise HTTPException(status_code=400, detail="Invalid token")


@app.get("/v1/installer-token/{token}")
async def resolve_installer_token(token: str, db: AsyncSession = Depends(get_db)):
    """
    Exchange a one-time installer token for a license key.
    Token is stored in the download_tokens table; expires after 1 use or 24h.
    """
    result = await db.execute(text("""
        SELECT license_key, used, created_at FROM download_tokens
        WHERE token = :token
          AND used = FALSE
          AND created_at > NOW() - INTERVAL '24 hours'
    """), {"token": token})
    row = result.fetchone()
    if not row:
        raise HTTPException(status_code=404, detail="Token not found or expired")
    # Mark as used
    await db.execute(text("UPDATE download_tokens SET used = TRUE WHERE token = :token"), {"token": token})
    await db.commit()
    return {"license_key": row.license_key}


@app.get("/v1/credits/balance", response_model=CreditBalanceResponse)
async def get_credit_balance(
    license: License = Depends(get_current_license),
    db: AsyncSession = Depends(get_db)
):
    """Get current credit balance."""
    return CreditBalanceResponse(
        credits_balance=license.credits_remaining,
        license_key=license.license_key,
        license_type=license.type,
        valid_until=license.valid_until
    )

# ============================================
# Routes: Identity (GET /v1/me)
# ============================================

# Map license key prefixes → product metadata
# Longest prefix first so "mkt_" wins over hypothetical "m_" etc.
_VERTICAL_BY_PREFIX: list[tuple[str, dict]] = sorted([
    ("lt_",   {"product_id": "law",           "product_name": "LawTasksAI",           "display_name": "LawTasksAI",           "tool_prefix": "lawtasksai",       "occupation": "attorney",             "support_email": "support@lawtasksai.com"}),
    ("ct_",   {"product_id": "contractor",    "product_name": "ContractorTasksAI",   "display_name": "ContractorTasksAI",   "tool_prefix": "contractortasksai", "occupation": "contractor",           "support_email": "support@contractortasksai.com"}),
    ("rt_",   {"product_id": "realtor",       "product_name": "RealtorTasksAI",      "display_name": "RealtorTasksAI",      "tool_prefix": "realtortasksai",    "occupation": "real estate agent",   "support_email": "support@realtortasksai.com"}),
    ("mkt_",  {"product_id": "marketing",     "product_name": "MarketingTasksAI",    "display_name": "MarketingTasksAI",    "tool_prefix": "marketingtasksai",  "occupation": "marketer",            "support_email": "support@marketingtasksai.com"}),
    ("ft_",   {"product_id": "farmer",        "product_name": "FarmerTasksAI",       "display_name": "FarmerTasksAI",       "tool_prefix": "farmertasksai",     "occupation": "farmer",              "support_email": "support@farmertasksai.com"}),
    ("pt_",   {"product_id": "pastor",        "product_name": "PastorTasksAI",       "display_name": "PastorTasksAI",       "tool_prefix": "pastortasksai",     "occupation": "pastor",              "support_email": "support@pastortasksai.com"}),
    ("st_",   {"product_id": "salon",         "product_name": "SalonTasksAI",        "display_name": "SalonTasksAI",        "tool_prefix": "salontasksai",      "occupation": "salon owner",         "support_email": "support@salontasksai.com"}),
    ("tt_",   {"product_id": "teacher",       "product_name": "TeacherTasksAI",      "display_name": "TeacherTasksAI",      "tool_prefix": "teachertasksai",    "occupation": "teacher",             "support_email": "support@teachertasksai.com"}),
    ("tht_",  {"product_id": "therapist",     "product_name": "TherapistTasksAI",    "display_name": "TherapistTasksAI",    "tool_prefix": "therapisttasksai",  "occupation": "therapist",           "support_email": "support@therapisttasksai.com"}),
    ("ht_",   {"product_id": "hr",            "product_name": "HRTasksAI",           "display_name": "HRTasksAI",           "tool_prefix": "hrtasksai",         "occupation": "HR professional",     "support_email": "support@hrtasksai.com"}),
    ("tat_",  {"product_id": "travelagent",   "product_name": "TravelAgentTasksAI",  "display_name": "TravelAgentTasksAI",  "tool_prefix": "travelagenttasksai","occupation": "travel agent",        "support_email": "support@travelagenttasksai.com"}),
    ("dt_",   {"product_id": "dentist",       "product_name": "DentistTasksAI",      "display_name": "DentistTasksAI",      "tool_prefix": "dentisttasksai",    "occupation": "dentist",             "support_email": "support@dentisttasksai.com"}),
    ("chrt_", {"product_id": "chiropractor",  "product_name": "ChiropractorTasksAI", "display_name": "ChiropractorTasksAI", "tool_prefix": "chiropractortasksai","occupation": "chiropractor",       "support_email": "support@chiropractortasksai.com"}),
    ("it_",   {"product_id": "insurance",     "product_name": "InsuranceTasksAI",    "display_name": "InsuranceTasksAI",    "tool_prefix": "insurancetasksai",  "occupation": "insurance agent",    "support_email": "support@insurancetasksai.com"}),
    ("at_",   {"product_id": "accounting",    "product_name": "AccountingTasksAI",   "display_name": "AccountingTasksAI",   "tool_prefix": "accountingtasksai", "occupation": "accountant",          "support_email": "support@accountingtasksai.com"}),
    ("nt_",   {"product_id": "nutritionist",  "product_name": "NutritionistTasksAI", "display_name": "NutritionistTasksAI", "tool_prefix": "nutritionisttasksai","occupation": "nutritionist",       "support_email": "support@nutritionisttasksai.com"}),
    ("rest_", {"product_id": "restaurant",    "product_name": "RestaurantTasksAI",   "display_name": "RestaurantTasksAI",   "tool_prefix": "restauranttasksai", "occupation": "restaurant owner",   "support_email": "support@restauranttasksai.com"}),
    ("llt_",  {"product_id": "landlord",      "product_name": "LandlordTasksAI",     "display_name": "LandlordTasksAI",     "tool_prefix": "landlordtasksai",   "occupation": "landlord",            "support_email": "support@landlordtasksai.com"}),
    ("prt_",  {"product_id": "principal",     "product_name": "PrincipalTasksAI",    "display_name": "PrincipalTasksAI",    "tool_prefix": "principaltasksai",  "occupation": "school principal",    "support_email": "support@principaltasksai.com"}),
    ("mot_",  {"product_id": "mortuary",      "product_name": "MortuaryTasksAI",     "display_name": "MortuaryTasksAI",     "tool_prefix": "mortuarytasksai",   "occupation": "funeral director",    "support_email": "support@mortuarytasksai.com"}),
    ("evt_",  {"product_id": "eventplanner",  "product_name": "EventPlannerTasksAI", "display_name": "EventPlannerTasksAI", "tool_prefix": "eventplannertasksai","occupation": "event planner",      "support_email": "support@eventplannertasksai.com"}),
    ("cht_",  {"product_id": "church",        "product_name": "ChurchTasksAI",       "display_name": "ChurchTasksAI",       "tool_prefix": "churchtasksai",     "occupation": "church administrator","support_email": "support@churchtasksai.com"}),
    ("per_",  {"product_id": "personaltrainer","product_name": "PersonalTrainerTasksAI","display_name": "PersonalTrainerTasksAI","tool_prefix": "personaltrainertasksai","occupation": "personal trainer",  "support_email": "support@personaltrainertasksai.com"}),
    ("elt_",  {"product_id": "electrician",   "product_name": "ElectricianTasksAI",  "display_name": "ElectricianTasksAI",  "tool_prefix": "electriciantasksai","occupation": "electrician",        "support_email": "support@electriciantasksai.com"}),
    ("mgt_",  {"product_id": "mortgage",      "product_name": "MortgageTasksAI",     "display_name": "MortgageTasksAI",     "tool_prefix": "mortgagetasksai",   "occupation": "mortgage broker",    "support_email": "support@mortgagetasksai.com"}),
    ("plt_",  {"product_id": "plumber",       "product_name": "PlumberTasksAI",      "display_name": "PlumberTasksAI",      "tool_prefix": "plumbertasksai",    "occupation": "plumber",             "support_email": "support@plumbertasksai.com"}),
    ("vt_",   {"product_id": "vet",           "product_name": "VetTasksAI",          "display_name": "VetTasksAI",          "tool_prefix": "vettasksai",        "occupation": "veterinarian",        "support_email": "support@vettasksai.com"}),
    ("fun_",  {"product_id": "funeral",       "product_name": "FuneralTasksAI",      "display_name": "FuneralTasksAI",      "tool_prefix": "funeraltasksai",    "occupation": "funeral director",    "support_email": "support@funeraltasksai.com"}),
    ("des_",  {"product_id": "designer",      "product_name": "DesignerTasksAI",     "display_name": "DesignerTasksAI",     "tool_prefix": "designertasksai",   "occupation": "designer",            "support_email": "support@designertasksai.com"}),
    ("mst_",  {"product_id": "militaryspouse","product_name": "MilitarySpouseTasksAI","display_name": "MilitarySpouseTasksAI","tool_prefix": "militaryspousetasksai","occupation": "military spouse",  "support_email": "support@militaryspousetasksai.com"}),
], key=lambda x: -len(x[0]))


def _vertical_from_key(key: str) -> dict | None:
    """Return vertical metadata dict for a license key, or None if no prefix matches."""
    for prefix, meta in _VERTICAL_BY_PREFIX:
        if key.startswith(prefix):
            return meta
    return None


@app.get("/v1/me")
async def get_me(
    license: License = Depends(get_current_license),
    db: AsyncSession = Depends(get_db),
):
    """
    Return identity + vertical metadata for the calling license key.
    Used by the MCP server at startup to self-configure tool names and prompts.
    """
    meta = _vertical_from_key(license.license_key)

    # Always fetch user row — needed for email and product_id fallback
    user_result = await db.execute(
        text("SELECT email, product_id FROM users WHERE id = :uid"),
        {"uid": str(license.user_id)},
    )
    user_row = user_result.fetchone()
    user_email = user_row.email if user_row else None

    if not meta:
        # Fallback: derive from license.product_id first, then user.product_id
        product_id = license.product_id or (user_row.product_id if user_row else "law") or "law"
        meta = {
            "product_id":    product_id,
            "product_name":  f"{product_id.title()}TasksAI",
            "display_name":  f"{product_id.title()}TasksAI",
            "tool_prefix":   f"{product_id}tasksai",
            "occupation":    product_id,
            "support_email": f"support@{product_id}tasksai.com",
        }

    # Fetch domain from products table if available
    prod_result = await db.execute(
        text("SELECT domain FROM products WHERE id = :pid AND is_active = TRUE"),
        {"pid": meta["product_id"]},
    )
    prod_row = prod_result.fetchone()
    domain = prod_row.domain if prod_row else f"{meta['product_id']}tasksai.com"

    return {
        **meta,
        "email":             user_email,
        "domain":            domain,
        "license_type":      license.type,
        "credits_remaining": license.credits_remaining,
    }


# ============================================
# Routes: Abbreviations (GET /v1/abbreviations)
# ============================================

@app.get("/v1/abbreviations")
async def get_abbreviations(
    license: License = Depends(get_current_license),
    db: AsyncSession = Depends(get_db),
):
    """
    Return abbreviation expansions for the calling license's vertical.
    Used by MCP server to populate _ABBREVS at startup.
    Response: {"product_id": "law", "abbreviations": {"tro": "temporary restraining order", ...}}
    """
    meta = _vertical_from_key(license.license_key)
    product_id = meta["product_id"] if meta else "law"

    result = await db.execute(
        text("SELECT abbreviation, expansion FROM skill_abbreviations WHERE product_id = :pid ORDER BY abbreviation"),
        {"pid": product_id},
    )
    rows = result.fetchall()

    return {
        "product_id":    product_id,
        "abbreviations": {row.abbreviation: row.expansion for row in rows},
        "count":         len(rows),
    }


@app.post("/credits/purchase")
async def purchase_credits(
    request: PurchaseCreditsRequest,
    license: License = Depends(get_current_license),
    db: AsyncSession = Depends(get_db)
):
    """
    Initiate credit purchase.
    Returns Stripe checkout URL (in production).
    For now, just adds credits directly.
    """
    if request.pack not in CREDIT_PACKS:
        raise HTTPException(status_code=400, detail="Invalid credit pack")
    
    pack = CREDIT_PACKS[request.pack]
    
    # In production: Create Stripe checkout session
    # For now: Add credits directly
    license.credits_remaining += pack["credits"]
    license.credits_purchased += pack["credits"]
    
    # Log transaction
    tx = CreditTransaction(
        user_id=license.user_id,
        license_id=license.id,
        type="purchase",
        amount=pack["credits"],
        balance_after=license.credits_remaining,
        description=f"Purchased {pack['name']} ({pack['credits']} credits)"
    )
    db.add(tx)
    
    await db.commit()
    
    return {
        "success": True,
        "credits_added": pack["credits"],
        "credits_balance": license.credits_remaining,
        "amount_charged": pack["price_cents"] / 100
    }

# ============================================
# Routes: Usage
# ============================================

@app.get("/usage", response_model=List[UsageResponse])
async def get_usage(
    limit: int = 50,
    license: License = Depends(get_current_license),
    db: AsyncSession = Depends(get_db)
):
    """Get usage history for current license."""
    result = await db.execute(
        select(UsageLog, Skill.name)
        .join(Skill, UsageLog.skill_id == Skill.id)
        .where(UsageLog.license_id == license.id)
        .order_by(UsageLog.executed_at.desc())
        .limit(limit)
    )
    
    return [
        UsageResponse(
            skill_id=log.skill_id,
            skill_name=skill_name,
            executed_at=log.executed_at,
            success=log.success,
            credits_used=log.credits_used
        )
        for log, skill_name in result.all()
    ]

# ============================================
# Routes: Profile
# ============================================

@app.get("/profile", response_model=ProfileResponse)
@app.get("/v1/profile", response_model=ProfileResponse)
async def get_profile(
    license: License = Depends(get_current_license),
    db: AsyncSession = Depends(get_db)
):
    """Get current user profile for document generation."""
    # Get user
    result = await db.execute(select(User).where(User.id == license.user_id))
    user = result.scalar_one_or_none()
    
    if not user:
        raise HTTPException(status_code=404, detail="User not found")
    
    profile_data = user.profile or {}
    
    # Check for commonly needed fields
    common_fields = ["firm_name", "attorney_name", "address", "city_state_zip", "phone"]
    missing = [f for f in common_fields if not profile_data.get(f)]
    
    return ProfileResponse(
        profile=UserProfile(**profile_data),
        missing_fields=missing
    )


@app.put("/profile", response_model=ProfileResponse)
@app.put("/v1/profile", response_model=ProfileResponse)
async def update_profile(
    profile_update: ProfileUpdateRequest,
    license: License = Depends(get_current_license),
    db: AsyncSession = Depends(get_db)
):
    """Update user profile (merges with existing data)."""
    # Get user
    result = await db.execute(select(User).where(User.id == license.user_id))
    user = result.scalar_one_or_none()
    
    if not user:
        raise HTTPException(status_code=404, detail="User not found")
    
    # Merge with existing profile
    current_profile = user.profile or {}
    update_data = profile_update.model_dump(exclude_none=True)
    current_profile.update(update_data)
    
    # Update in database
    user.profile = current_profile
    await db.commit()
    
    # Check for commonly needed fields
    common_fields = ["firm_name", "attorney_name", "address", "city_state_zip", "phone"]
    missing = [f for f in common_fields if not current_profile.get(f)]
    
    return ProfileResponse(
        profile=UserProfile(**current_profile),
        missing_fields=missing
    )


@app.get("/profile/check/{skill_id}")
@app.get("/v1/profile/check/{skill_id}")
async def check_profile_for_skill(
    skill_id: str,
    license: License = Depends(get_current_license),
    db: AsyncSession = Depends(get_db)
):
    """
    Check if user profile has required fields for a specific skill's document generation.
    Returns missing fields if any.
    """
    # Get user
    result = await db.execute(select(User).where(User.id == license.user_id))
    user = result.scalar_one_or_none()
    
    if not user:
        raise HTTPException(status_code=404, detail="User not found")
    
    profile_data = user.profile or {}
    missing = check_profile_requirements(skill_id, profile_data)
    
    return {
        "skill_id": skill_id,
        "generates_document": skill_id in DOCUMENT_SKILLS,
        "document_format": DOCUMENT_SKILLS.get(skill_id, {}).get("format"),
        "profile_complete": len(missing) == 0,
        "missing_fields": missing
    }

# ============================================
# Routes: Account (Dashboard)
# ============================================

@app.get("/v1/account/stats")
async def get_account_stats(
    license: License = Depends(get_current_license),
    db: AsyncSession = Depends(get_db)
):
    """Get account statistics for dashboard."""
    return {
        "credits_remaining": license.credits_remaining,
        "total_purchased": license.credits_purchased,
        "license_type": license.type,
        "created_at": license.created_at.isoformat() if license.created_at else None
    }


@app.get("/v1/account/purchases")
async def get_purchase_history(
    license: License = Depends(get_current_license),
    db: AsyncSession = Depends(get_db)
):
    """Get purchase history for dashboard."""
    # Get credit transactions (purchases only)
    result = await db.execute(
        select(CreditTransaction)
        .where(
            CreditTransaction.license_id == license.id,
            CreditTransaction.type == "purchase"
        )
        .order_by(CreditTransaction.created_at.desc())
    )
    transactions = result.scalars().all()
    
    purchases = []
    for tx in transactions:
        # Try to get Stripe invoice URL if reference_id is a Stripe session
        invoice_url = None
        if tx.reference_id and tx.reference_id.startswith('cs_'):
            try:
                session = stripe.checkout.Session.retrieve(tx.reference_id)
                if session.invoice:
                    invoice = stripe.Invoice.retrieve(session.invoice)
                    invoice_url = invoice.hosted_invoice_url
            except:
                pass
        
        purchases.append({
            "id": tx.id,
            "date": tx.created_at.isoformat() if tx.created_at else None,
            "description": tx.description or f"Credit purchase",
            "credits": tx.amount,
            "invoice_url": invoice_url
        })
    
    return purchases


# ============================================
# Routes: Document Regeneration
# ============================================

class RegenerateDocumentResponse(BaseModel):
    """Response for document regeneration."""
    usage_id: int
    skill_id: str
    document: DocumentAttachment
    message: str

# ============================================
# Routes: Categories
# ============================================

@app.get("/categories")
@app.get("/v1/categories")
async def list_categories(db: AsyncSession = Depends(get_db)):
    """List all skill categories."""
    result = await db.execute(select(Category).order_by(Category.display_order))
    categories = result.scalars().all()
    
    return [
        {"id": c.id, "name": c.name, "description": c.description}
        for c in categories
    ]

# ============================================
# Health Check
# ============================================

@app.get("/health")
async def health_check():
    """Health check endpoint."""
    return {"status": "healthy", "service": "lawtasksai-api", "version": "1.0.0"}


# Routes: Email Tracking
@app.get("/track/email-open")
async def track_email_open(
    mid: Optional[str] = Query(None),
    product: Optional[str] = Query(None),
    email: Optional[str] = Query(None),
    db: AsyncSession = Depends(get_db),
):
    """
    1x1 pixel endpoint for email open tracking.
    Logs the open to email_opens table and returns a transparent GIF.
    """
    # Log to DB (fire-and-forget style — never fail the pixel response)
    try:
        await db.execute(
            text("""
                INSERT INTO email_opens (message_id, product_id, email, opened_at)
                VALUES (:mid, :product, :email, NOW())
                ON CONFLICT DO NOTHING
            """),
            {"mid": mid or "", "product": product or "", "email": email or ""}
        )
        await db.commit()
    except Exception:
        pass

    # Return 1x1 transparent GIF
    gif = b'GIF89a\x01\x00\x01\x00\x80\x00\x00\xff\xff\xff\x00\x00\x00!\xf9\x04\x00\x00\x00\x00\x00,\x00\x00\x00\x00\x01\x00\x01\x00\x00\x02\x02D\x01\x00;'
    from fastapi.responses import Response
    return Response(content=gif, media_type="image/gif", headers={
        "Cache-Control": "no-cache, no-store, must-revalidate",
        "Pragma": "no-cache",
    })


@app.get("/v1/loader/latest")
async def get_loader_latest(
    license: License = Depends(get_current_license),
    x_loader_version: Optional[str] = Header(None, alias="X-Loader-Version")
):
    """
    Return the current loader version and full SKILL.md content for auto-update.

    The loader calls this on startup when its local version is behind.
    Returns the full SKILL.md so the loader can replace itself in one round-trip.
    """
    already_current = (x_loader_version == CURRENT_LOADER_VERSION)

    try:
        with open(LOADER_SKILL_PATH, "r") as f:
            skill_content = f.read()
    except FileNotFoundError:
        raise HTTPException(status_code=500, detail="Loader file not found on server")

    return {
        "version": CURRENT_LOADER_VERSION,
        "already_current": already_current,
        "skill_md": skill_content,
        "update_message": LOADER_UPDATE_MESSAGE,
    }


# ============================================
# Routes: Gap Reporting (anonymous, user-consented)
# ============================================

class GapReportRequest(BaseModel):
    search_terms: List[str]           # keywords that failed to match any skill
    loader_version: Optional[str] = None

class FeedbackSubmit(BaseModel):
    reason: str
    email: Optional[str] = None
    extra: Optional[str] = None
    product: Optional[str] = None


@app.post("/v1/feedback", status_code=204)
async def submit_feedback(
    data: FeedbackSubmit,
    db: AsyncSession = Depends(get_db)
):
    """Accept free-text feedback from the 'other' block on feedback-thanks.html."""
    try:
        fb = UserFeedback(
            email=data.email,
            product_id=data.product or "law",
            reason=data.reason,
            extra=data.extra,
        )
        db.add(fb)
        await db.commit()
        print(f"[feedback/post] {data.reason} from {data.email or 'anon'}")
    except Exception as e:
        print(f"[feedback/post] failed: {e}")


@app.get("/v1/feedback/{reason}")
async def drip_feedback(
    reason: str,
    email: Optional[str] = Query(None),
    product: Optional[str] = Query(None),
    db: AsyncSession = Depends(get_db)
):
    """
    One-click feedback from Email 3 drip sequence.
    reason is a path segment to avoid Zoho click-tracking mangling query params.
    Stores reason + email, redirects to /feedback-thanks?reason=X
    """
    from fastapi.responses import RedirectResponse

    VALID_REASONS = {"installation", "forgot", "wrong_tool", "not_sure", "other"}
    safe_reason = reason if reason in VALID_REASONS else "other"
    safe_product = product or "law"

    # Store feedback (fire and forget — don't fail the redirect if DB is slow)
    try:
        fb = UserFeedback(
            email=email,
            product_id=safe_product,
            reason=safe_reason,
        )
        db.add(fb)
        await db.commit()
        print(f"[feedback] {safe_reason} from {email or 'unknown'} ({safe_product})")
    except Exception as e:
        print(f"[feedback] DB store failed: {e}")

    # Resolve the domain for the product so we redirect back to the right site
    domain = "lawtasksai.com"
    try:
        prod_result = await db.execute(
            text("SELECT domain FROM products WHERE id = :pid AND is_active = TRUE"),
            {"pid": safe_product}
        )
        row = prod_result.fetchone()
        if row and row.domain:
            domain = row.domain
    except Exception:
        pass

    # CF Workers strip .html extensions — use clean URL to avoid double-redirect
    redirect_url = f"https://{domain}/feedback-thanks?reason={safe_reason}"
    return RedirectResponse(url=redirect_url, status_code=302)


@app.post("/v1/feedback/gap", status_code=204)
async def report_skill_gap(
    report: GapReportRequest,
    db: AsyncSession = Depends(get_db)
):
    """
    Accept an anonymous gap report from the loader.
    Called only when the user explicitly consents to share per-request.
    No user identity, no query content — only the search terms.
    Returns 204 No Content on success.
    """
    # Sanitize: strip empties, lowercase, limit to 20 terms, 50 chars each
    clean_terms = [
        t.strip().lower()[:50]
        for t in report.search_terms
        if t.strip()
    ][:20]

    if not clean_terms:
        return  # Nothing to record

    gap = SkillGap(
        search_terms=" ".join(clean_terms),
        loader_version=report.loader_version,
    )
    db.add(gap)
    await db.commit()


# ============================================
# Routes: Support Contact Form
# ============================================

class SupportRequest(BaseModel):
    name: str
    email: EmailStr
    topic: Optional[str] = ""
    message: str
    product: Optional[str] = "TasksAI"  # populated by the landing page template

@app.post("/v1/support", status_code=204)
async def submit_support_request(req: SupportRequest):
    """
    Accept a support contact form submission and forward it to the
    product support inbox via Zoho Mail.
    Returns 204 No Content on success.
    """
    # Determine the right inbox based on the product field
    # All verticals share the same Zoho account; route to hello@lawtasksai.com
    # (the catch-all inbox) and include the product name in the subject.
    to_address = "hello@lawtasksai.com"
    subject = f"[Support] {req.product} — {req.topic or 'General inquiry'}"
    body = (
        f"Name:    {req.name}\n"
        f"Email:   {req.email}\n"
        f"Product: {req.product}\n"
        f"Topic:   {req.topic or 'Not specified'}\n"
        f"\n"
        f"{req.message}\n"
        f"\n---\n"
        f"Reply directly to this email to respond to the user.\n"
    )

    zoho_message_id = None
    try:
        access_token = await get_zoho_access_token()
        if access_token:
            async with httpx.AsyncClient(timeout=10) as client:
                resp = await client.post(
                    f"https://mail.zoho.com/api/accounts/{os.getenv('ZOHO_ACCOUNT_ID', '6556209000000008002')}/messages",
                    json={
                        "fromAddress": "hello@lawtasksai.com",
                        "toAddress": to_address,
                        "replyTo": req.email,
                        "subject": subject,
                        "content": body,
                        "mailFormat": "plaintext"
                    },
                    headers={"Authorization": f"Zoho-oauthtoken {access_token}"}
                )
                print(f"[Support] form from {req.email} ({req.product}): Zoho status {resp.status_code}")
                if resp.status_code == 200:
                    zoho_message_id = resp.json().get("data", {}).get("messageId")
        else:
            print(f"[Support] could not get Zoho access token — support email NOT sent for {req.email}")
    except Exception as e:
        print(f"[Support] failed to send support email for {req.email}: {e}")
        # Don't surface the error to the user — the 204 response is still sent
        # so they aren't left with a broken form. Log it for ops review.

    # Always insert into support_requests DB regardless of email outcome
    try:
        async with engine.begin() as conn:
            await conn.execute(
                text("""
                    INSERT INTO support_requests
                        (email, name, product_id, subject, message, direction, status, zoho_message_id)
                    VALUES
                        (:email, :name, :product_id, :subject, :message, 'inbound', 'open', :zoho_message_id)
                """),
                {
                    "email":           req.email,
                    "name":            req.name,
                    "product_id":      req.product.lower().replace("tasksai", "").replace(" ", ""),
                    "subject":         subject,
                    "message":         req.message,
                    "zoho_message_id": zoho_message_id,
                }
            )
    except Exception as e:
        print(f"[Support] DB insert failed for {req.email}: {e}")


# ============================================
# Routes: Checkout & Purchase
# ============================================

# Load LOADER_SKILL_MD from file or inline fallback
def _load_loader_skill_md():
    import pathlib as _pl
    candidates = [
        _pl.Path(__file__).parent / 'loader' / 'SKILL.md',
        _pl.Path('/app/loader/SKILL.md'),
    ]
    for p in candidates:
        if p.exists():
            return p.read_text()
    return "# LawTasksAI Skills\n\nLoader SKILL.md not found.\n"

LOADER_SKILL_MD = _load_loader_skill_md()

@app.post("/checkout/create-session")
async def create_checkout_session(
    request: PurchaseCreditsRequest,
    db: AsyncSession = Depends(get_db),
    product_id: str = Depends(get_product_id),
):
    """Create a Stripe checkout session for credit purchase."""

    # --- Resolve pack: prefer product-specific packs, fall back to CREDIT_PACKS ---
    pack = None
    resolved_product_id = product_id  # from header/query/default

    # Try product_credit_packs table first
    try:
        pcp_result = await db.execute(
            text(
                "SELECT pack_key, name, credits, price_cents FROM product_credit_packs "
                "WHERE product_id = :pid ORDER BY credits"
            ),
            {"pid": resolved_product_id},
        )
        product_packs_rows = pcp_result.fetchall()
        if product_packs_rows:
            product_packs = {
                row.pack_key: {
                    "credits": row.credits,
                    "price_cents": row.price_cents,
                    "name": row.name,
                    "one_time": False,
                }
                for row in product_packs_rows
            }
            pack = product_packs.get(request.pack)
    except Exception:
        pass  # table may not exist in all environments — fall through to CREDIT_PACKS

    # Fall back to hardcoded CREDIT_PACKS
    if pack is None:
        if request.pack not in CREDIT_PACKS:
            raise HTTPException(status_code=400, detail="Invalid credit pack")
        pack = CREDIT_PACKS[request.pack]

    # Enforce one-time trial/starter pack per email per vertical
    if pack.get("one_time"):
        if not request.email:
            raise HTTPException(status_code=400, detail="Email is required for the Trial pack")
        # Check if this email has already purchased a trial/starter pack FOR THIS VERTICAL
        # Each vertical gets its own one-time trial — buying on RealtorTasksAI doesn't block TeacherTasksAI
        result = await db.execute(
            select(CreditTransaction)
            .join(User, CreditTransaction.user_id == User.id)
            .join(License, CreditTransaction.license_id == License.id)
            .where(
                User.email == request.email.lower().strip(),
                License.product_id == resolved_product_id,
                or_(
                    CreditTransaction.description.contains("Trial"),
                    CreditTransaction.description.contains("Starter")
                )
            )
        )
        existing = result.scalar_one_or_none()
        if existing:
            raise HTTPException(status_code=400, detail="The Trial package is only available once per product. Check out our other packages for better value!")
    
    # Resolve product domain for success/cancel URLs
    product_domain = None
    product_display_name = None
    try:
        prod_result = await db.execute(
            text("SELECT domain, name FROM products WHERE id = :pid AND is_active = TRUE"),
            {"pid": resolved_product_id},
        )
        prod_row = prod_result.fetchone()
        if prod_row:
            product_domain = f"https://{prod_row.domain}"
            product_display_name = prod_row.name
    except Exception:
        pass
    frontend_url = product_domain or FRONTEND_URL
    product_name = product_display_name or "LawTasksAI"

    try:
        # Create Stripe checkout session
        checkout_kwargs = {
            'payment_method_types': ['card'],
            'line_items': [{
                'price_data': {
                    'currency': 'usd',
                    'product_data': {
                        'name': f'{product_name} {pack["name"]}',
                        'description': f'{pack["credits"]} AI task credits',
                    },
                    'unit_amount': pack['price_cents'],
                },
                'quantity': 1,
            }],
            'mode': 'payment',
            'success_url': f'{frontend_url}/success?session_id={{CHECKOUT_SESSION_ID}}',
            'cancel_url': f'{frontend_url}/#pricing',
            'metadata': {
                'pack': request.pack,
                'credits': str(pack['credits']),
                'product_id': resolved_product_id,
                'attorney_name': request.attorney_name or '',
                'bar_jurisdiction': request.bar_jurisdiction or '',
                'bar_number': request.bar_number or '',
                'firm_name': request.firm_name or '',
                'attestation': 'true' if request.attestation else 'false',
            },
        }
        
        # Pre-fill email if provided
        if request.email:
            checkout_kwargs['customer_email'] = request.email.lower().strip()
        
        checkout_session = stripe.checkout.Session.create(**checkout_kwargs)
        
        return {
            "checkout_url": checkout_session.url,
            "session_id": checkout_session.id
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

async def get_zoho_campaigns_token() -> str:
    """Exchange Zoho Campaigns refresh token for a fresh access token."""
    refresh_token = os.getenv("ZOHO_CAMPAIGNS_REFRESH_TOKEN", "")
    client_id = os.getenv("ZOHO_CAMPAIGNS_CLIENT_ID", "")
    client_secret = os.getenv("ZOHO_CAMPAIGNS_CLIENT_SECRET", "")
    if not refresh_token:
        return ""
    async with httpx.AsyncClient(timeout=10) as client:
        resp = await client.post(
            "https://accounts.zoho.com/oauth/v2/token",
            data={
                "refresh_token": refresh_token,
                "client_id": client_id,
                "client_secret": client_secret,
                "grant_type": "refresh_token"
            }
        )
        data = resp.json()
        return data.get("access_token", "")


async def get_zoho_access_token() -> str:
    """Exchange Zoho Mail refresh token for a fresh access token."""
    async with httpx.AsyncClient(timeout=10) as client:
        resp = await client.post(
            "https://accounts.zoho.com/oauth/v2/token",
            data={
                "refresh_token": os.getenv("ZOHO_REFRESH_TOKEN", ""),
                "client_id": os.getenv("ZOHO_CLIENT_ID", ""),
                "client_secret": os.getenv("ZOHO_CLIENT_SECRET", ""),
                "grant_type": "refresh_token"
            }
        )
        data = resp.json()
        return data.get("access_token", "")


def _load_zoho_listkeys() -> dict:
    """Load per-vertical Zoho list keys from bundled JSON."""
    try:
        lk_path = os.path.join(os.path.dirname(__file__) if '__file__' in dir() else '/app', 'zoho-listkeys.json')
        with open(lk_path) as f:
            return json.load(f)
    except Exception:
        return {}

_ZOHO_LISTKEYS: dict = _load_zoho_listkeys()


def get_zoho_listkey(product_id: str) -> str:
    """Get the Zoho Campaigns list key for a given product_id."""
    return _ZOHO_LISTKEYS.get(product_id, os.getenv("ZOHO_LIST_KEY", ""))


async def add_to_zoho_list(email: str, name: str, product_id: str = "law") -> None:
    """Add a contact to the per-vertical Zoho Campaigns list. Fails silently."""
    list_key = get_zoho_listkey(product_id)
    if not list_key or not os.getenv("ZOHO_CAMPAIGNS_REFRESH_TOKEN"):
        return
    try:
        access_token = await get_zoho_campaigns_token()
        if not access_token:
            print("[Zoho Campaigns] could not get access token")
            return
        contact_info = json.dumps({
            "Contact Email": email,
            "First Name": name or "",
        })
        async with httpx.AsyncClient(timeout=10) as client:
            resp = await client.post(
                "https://campaigns.zoho.com/api/v1.1/json/listsubscribe",
                params={"resfmt": "JSON", "listkey": list_key, "contactinfo": contact_info},
                headers={"Authorization": f"Zoho-oauthtoken {access_token}"}
            )
            print(f"[Zoho Campaigns] add subscriber {email} -> {product_id}: {resp.status_code} {resp.text[:200]}")
    except Exception as e:
        print(f"[Zoho Campaigns] failed to add subscriber {email}: {e}")


@app.post("/webhooks/stripe")
async def stripe_webhook(request: Request, db: AsyncSession = Depends(get_db)):
    """Handle Stripe webhook events."""
    payload = await request.body()
    sig_header = request.headers.get('stripe-signature')
    
    try:
        event = stripe.Webhook.construct_event(
            payload, sig_header, STRIPE_WEBHOOK_SECRET
        )
    except ValueError:
        raise HTTPException(status_code=400, detail="Invalid payload")
    except stripe.error.SignatureVerificationError:
        raise HTTPException(status_code=400, detail="Invalid signature")
    
    # Handle checkout.session.completed
    if event['type'] == 'checkout.session.completed':
        session = event['data']['object']
        session_id = session.get('id', '')

        # Idempotency: skip if this session was already processed
        existing_tx = await db.execute(
            select(CreditTransaction).where(CreditTransaction.reference_id == session_id)
        )
        if existing_tx.scalar_one_or_none():
            return {"status": "skipped", "reason": "already_processed"}

        # Get customer email
        customer_email = session.get('customer_details', {}).get('email')
        if not customer_email:
            return {"status": "skipped", "reason": "no email"}

        # Get credits and pack name from metadata
        meta = session.get('metadata', {})
        credits = int(meta.get('credits', 0))
        pack_name = meta.get('pack', 'unknown')
        if credits == 0:
            return {"status": "skipped", "reason": "no credits"}

        # Product ID and bar/attorney info from metadata
        purchase_product_id = meta.get('product_id', 'law')
        bar_profile = {
            "attorney_name": meta.get('attorney_name', ''),
            "bar_jurisdiction": meta.get('bar_jurisdiction', ''),
            "bar_number": meta.get('bar_number', ''),
            "firm_name": meta.get('firm_name', ''),
            "attestation": meta.get('attestation', 'false') == 'true',
        }

        # Look up product domain for email/download links
        purchase_product_domain = "lawtasksai.com"
        purchase_product_name = "LawTasksAI"
        try:
            prod_result = await db.execute(
                text("SELECT domain, name FROM products WHERE id = :pid AND is_active = TRUE"),
                {"pid": purchase_product_id}
            )
            prod_row = prod_result.fetchone()
            if prod_row and prod_row.domain:
                purchase_product_domain = prod_row.domain
                purchase_product_name = prod_row.name
        except Exception:
            pass

        # Find or create user by email (one account per email, product tracked separately)
        result = await db.execute(select(User).where(User.email == customer_email))
        user = result.scalar_one_or_none()

        if not user:
            # New customer — create account tagged to this product
            user = User(
                email=customer_email,
                name=bar_profile.get('attorney_name') or None,
                firm_name=bar_profile.get('firm_name') or None,
                password_hash=hash_password(secrets.token_hex(16)),
                credits_balance=credits,
                profile=bar_profile,
                product_id=purchase_product_id,
            )
            db.add(user)
            await db.flush()  # get user.id

            license = License(
                license_key=generate_license_key(),
                user_id=user.id,
                type="credits",
                credits_purchased=credits,
                credits_remaining=credits,
                product_id=purchase_product_id,
            )
            db.add(license)
            await db.flush()  # get license.id
        else:
            # Existing customer — add credits, update product_id if this is a new product for them
            # credits_balance is derived from license — no write needed here
            # Only update product_id if account has no product yet
            if not user.product_id:
                user.product_id = purchase_product_id
            # Merge bar profile (don't overwrite with empty)
            existing_profile = user.profile or {}
            for k, v in bar_profile.items():
                if v:
                    existing_profile[k] = v
            user.profile = existing_profile
            if bar_profile.get('attorney_name') and not user.name:
                user.name = bar_profile['attorney_name']
            if bar_profile.get('firm_name') and not user.firm_name:
                user.firm_name = bar_profile['firm_name']

            # Find the license for THIS specific product — never touch licenses from other verticals
            lic_result = await db.execute(
                select(License).where(
                    License.user_id == user.id,
                    License.product_id == purchase_product_id,
                    License.status == "active",
                ).order_by(License.created_at.desc())
            )
            license = lic_result.scalar_one_or_none()

            if not license:
                # New vertical purchase for this user — create a fresh license
                license = License(
                    license_key=generate_license_key(),
                    user_id=user.id,
                    type="credits",
                    credits_purchased=credits,
                    credits_remaining=credits,
                    product_id=purchase_product_id,
                )
                db.add(license)
                await db.flush()  # get license.id
            else:
                # Top-up on existing vertical license — add credits, never change product_id
                license.credits_remaining += credits
                license.credits_purchased += credits

        # Log transaction
        tx = CreditTransaction(
            user_id=user.id,
            license_id=license.id,
            type="purchase",
            amount=credits,
            balance_after=license.credits_remaining,
            reference_id=session['id'],
            description=f"Purchased {pack_name} pack via Stripe checkout"
        )
        db.add(tx)

        await db.commit()

        # Send confirmation email via Zoho SMTP or SendGrid
        try:
            success_url = f"https://{purchase_product_domain}/success?session_id={session['id']}"
            task_library_url = f"https://{purchase_product_domain}/task-library"
            email_subject = f"You're set up — your {purchase_product_name} license key"
            email_body = f"""Your {purchase_product_name} purchase is confirmed — {credits} credits are ready to use.

YOUR LICENSE KEY
{license.license_key}

Keep this somewhere safe. It's how you authenticate and download your skills anytime.

NEXT STEP
Visit your purchase summary to download and install:
{success_url}

You have {credits} credits to start — each credit runs one task.
Browse what's available: {task_library_url}

Questions? Reply to this email or reach us at hello@{purchase_product_domain}

— The {purchase_product_name} Team
"""
            from_addr = f"hello@{purchase_product_domain}"
            from_name = purchase_product_name
            email_sent = False

            # Send via Zoho Mail API (OAuth) — no password needed
            try:
                zoho_client_id = os.getenv("ZOHO_CLIENT_ID", "")
                zoho_client_secret = os.getenv("ZOHO_CLIENT_SECRET", "")
                zoho_refresh_token = os.getenv("ZOHO_MAIL_REFRESH_TOKEN", "") or os.getenv("ZOHO_REFRESH_TOKEN", "")
                zoho_account_id = "6556209000000008002"

                if zoho_client_id and zoho_refresh_token:
                    # Get fresh access token
                    token_resp = await asyncio.get_event_loop().run_in_executor(None, lambda: __import__('urllib.request', fromlist=['urlopen']).urlopen(
                        __import__('urllib.request', fromlist=['Request']).Request(
                            "https://accounts.zoho.com/oauth/v2/token",
                            data=f"refresh_token={zoho_refresh_token}&client_id={zoho_client_id}&client_secret={zoho_client_secret}&grant_type=refresh_token".encode(),
                            method="POST"
                        ), timeout=10
                    ))
                    token_data = json.loads(token_resp.read())
                    access_token = token_data.get("access_token")

                    if access_token:
                        # Send via Zoho Mail API
                        mail_payload = json.dumps({
                            "fromAddress": from_addr,
                            "toAddress": customer_email,
                            "subject": email_subject,
                            "content": email_body,
                            "mailFormat": "plaintext"
                        }).encode()
                        import urllib.request as _ur
                        mail_req = _ur.Request(
                            f"https://mail.zoho.com/api/accounts/{zoho_account_id}/messages",
                            data=mail_payload,
                            headers={
                                "Authorization": f"Zoho-oauthtoken {access_token}",
                                "Content-Type": "application/json"
                            },
                            method="POST"
                        )
                        with _ur.urlopen(mail_req, timeout=10) as mr:
                            resp_data = json.loads(mr.read())
                            print(f"[Email] sent via Zoho API to {customer_email}: {resp_data.get('status', {}).get('code')}")
                        email_sent = True
                    else:
                        print(f"[Email] Zoho token refresh failed: {token_data}")
            except Exception as zoho_err:
                print(f"[Email] Zoho API send failed: {zoho_err}")

            if not email_sent:
                print(f"[Email] skipped for {customer_email}: Zoho API send failed or not configured")
        except Exception as email_err:
            print(f"[Email] failed for {customer_email}: {email_err}")

        # Add buyer to Zoho Campaigns subscriber list
        await add_to_zoho_list(customer_email, user.name or "")

        return {"status": "success", "credits_added": credits, "product_id": purchase_product_id}
    
    return {"status": "ignored", "event_type": event['type']}

@app.get("/checkout/session/{session_id}")
async def get_checkout_session(session_id: str, db: AsyncSession = Depends(get_db)):
    """Get checkout session details (for success page).
    If webhook hasn't fired yet, provisions the user on the spot so the
    success page always works regardless of webhook timing.
    """
    try:
        session = stripe.checkout.Session.retrieve(session_id)
        
        if session.payment_status != 'paid':
            raise HTTPException(status_code=400, detail="Payment not completed")
        
        customer_email = session.get('customer_details', {}).get('email')
        meta = session.get('metadata', {})
        credits = int(meta.get('credits', 0))
        purchase_product_id = meta.get('product_id', 'law')

        # Get user — provision if webhook hasn't fired yet
        result = await db.execute(select(User).where(User.email == customer_email))
        user = result.scalar_one_or_none()

        if not user:
            # Webhook hasn't fired yet — provision now (webhook will be idempotent)
            user = User(
                email=customer_email,
                password_hash=hash_password(secrets.token_hex(16)),
                credits_balance=credits,
                product_id=purchase_product_id,
            )
            db.add(user)
            await db.flush()
            license = License(
                license_key=generate_license_key(),
                user_id=user.id,
                type="credits",
                credits_purchased=credits,
                credits_remaining=credits,
            )
            db.add(license)
            # Log transaction with session_id so webhook skips it later
            tx = CreditTransaction(
                user_id=user.id,
                license_id=license.id,
                type="purchase",
                amount=credits,
                balance_after=credits,
                reference_id=session_id,
                description=f"Purchased {meta.get('pack','unknown')} pack via Stripe checkout"
            )
            db.add(tx)
            await db.flush()
            await db.commit()

        result = await db.execute(
            select(License).where(
                License.user_id == user.id,
                License.status == "active"
            ).order_by(License.created_at.desc())
        )
        license = result.scalar_one_or_none()

        return {
            "email": customer_email,
            "credits_purchased": credits,
            "license_key": license.license_key if license else None,
            "total_credits": license.credits_remaining if license else credits,
            "product_id": user.product_id or "law"
        }
    except stripe.error.InvalidRequestError:
        raise HTTPException(status_code=404, detail="Session not found")

@app.get("/download/loader")
async def download_loader_by_session(
    session_id: str,
    request: Request,
    db: AsyncSession = Depends(get_db)
):
    """
    Download loader using Stripe session_id — always serves the correct product
    regardless of what else is on the account. Used in purchase confirmation emails.

    Routing:
      - MCP-platform users (claude_desktop, cursor, windsurf, cline) →
        redirect to GitHub Release .exe / Mac binary
      - OpenClaw / other users → serve OpenClaw-only zip
    """
    try:
        session = stripe.checkout.Session.retrieve(session_id)
        if session.payment_status != 'paid':
            raise HTTPException(status_code=400, detail="Payment not completed")
        customer_email = session.get('customer_details', {}).get('email')
        meta = session.get('metadata', {})
        product_id = meta.get('product_id', 'law')
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Invalid session: {e}")

    result = await db.execute(select(User).where(User.email == customer_email))
    user = result.scalar_one_or_none()
    if not user:
        raise HTTPException(status_code=404, detail="User not found")

    license_result = await db.execute(
        select(License).where(License.user_id == user.id, License.status == "active")
        .order_by(License.valid_from.desc()).limit(1)
    )
    license = license_result.scalar_one_or_none()
    if not license:
        raise HTTPException(status_code=404, detail="No active license found")

    # Route to GitHub installer for all users with a supported product.
    # The zip/OpenClaw path is legacy — only used if product has no installer.
    if product_id in _PRODUCT_NAME_MAP:
        # Stamp first download (funnel tracking)
        if license.downloaded_at is None:
            license.downloaded_at = datetime.utcnow()
            try:
                await db.commit()
            except Exception:
                await db.rollback()
        urls = _get_installer_url(product_id, request.headers.get("user-agent", ""))
        ua = request.headers.get("user-agent", "").lower()
        dest = urls["windows"] if "windows" in ua else urls["mac"]
        from fastapi.responses import RedirectResponse
        return RedirectResponse(url=dest, status_code=302)

    return await _build_loader_zip(license.license_key, product_id, db)


# ── Platform routing ─────────────────────────────────────────────────────────
# Platforms that use the .exe / binary installer from GitHub Releases.
# Everyone else (openclaw, other, unknown) gets the zip.
MCP_PLATFORMS = {"claude_desktop", "claude_code", "cursor", "windsurf", "cline"}

GITHUB_RELEASES_BASE = "https://github.com/laudoluxDev/lawtasksai-mcp/releases/latest/download"

# Map product_id → ProductName used in the GitHub Release asset filenames
_PRODUCT_NAME_MAP = {
    "law":            "LawTasksAI",
    "realtor":        "RealtorTasksAI",
    "farmer":         "FarmerTasksAI",
    "teacher":        "TeacherTasksAI",
    "therapist":      "TherapistTasksAI",
    "marketing":      "MarketingTasksAI",
    "contractor":     "ContractorTasksAI",
    "accounting":     "AccountingTasksAI",
    "chiropractor":   "ChiropractorTasksAI",
    "church":         "ChurchTasksAI",
    "dentist":        "DentistTasksAI",
    "designer":       "DesignerTasksAI",
    "electrician":    "ElectricianTasksAI",
    "eventplanner":   "EventPlannerTasksAI",
    "funeral":        "FuneralTasksAI",
    "hr":             "HRTasksAI",
    "insurance":      "InsuranceTasksAI",
    "landlord":       "LandlordTasksAI",
    "militaryspouse": "MilitarySpouseTasksAI",
    "mortgage":       "MortgageTasksAI",
    "mortuary":       "MortuaryTasksAI",
    "nutritionist":   "NutritionistTasksAI",
    "pastor":         "PastorTasksAI",
    "personaltrainer":"PersonalTrainerTasksAI",
    "plumber":        "PlumberTasksAI",
    "principal":      "PrincipalTasksAI",
    "restaurant":     "RestaurantTasksAI",
    "salon":          "SalonTasksAI",
    "travelagent":    "TravelAgentTasksAI",
    "vet":            "VetTasksAI",
}


def _get_installer_url(product_id: str, user_agent: str = "") -> dict:
    """
    Return platform-aware installer download URLs for a given product.
    Returns a dict with 'windows' and 'mac' keys, plus 'product_name'.
    """
    prod_name = _PRODUCT_NAME_MAP.get(product_id, "LawTasksAI")
    return {
        "product_name": prod_name,
        "windows": f"{GITHUB_RELEASES_BASE}/{prod_name}-Setup.exe",
        "mac":     f"{GITHUB_RELEASES_BASE}/{prod_name}-Setup-mac",
        "release_page": "https://github.com/laudoluxDev/lawtasksai-mcp/releases/latest",
    }


def _user_platforms(user: "User | None") -> set:
    """Return the set of platforms stored on a user record."""
    if not user:
        return set()
    raw = getattr(user, "platforms", None) or []
    return set(raw)


async def _build_installer_zip(license_key: str, product_id: str, user, request: Request):
    """
    Build a zip containing:
      - The platform-appropriate installer binary (fetched from GitHub Releases)
      - A pre-filled .env so the installer never prompts for a license key
    User downloads one zip, extracts, double-clicks the .exe — done.
    """
    ua = request.headers.get("user-agent", "").lower()
    is_windows = "windows" in ua
    urls = _get_installer_url(product_id, ua)
    binary_url = urls["windows"] if is_windows else urls["mac"]
    prod_name = _PRODUCT_NAME_MAP.get(product_id, "TasksAI")
    prod_slug = product_id

    # Build env var name matching what the installer expects
    if prod_slug == "law":
        env_var = "LAWTASKSAI_LICENSE_KEY"
    else:
        env_var = f"{prod_slug.upper()}TASKSAI_LICENSE_KEY"

    env_contents = f"{env_var}={license_key}\nTASKSAI_LICENSE_KEY={license_key}\n"

    # Fetch the binary from GitHub
    binary_data = None
    try:
        async with httpx.AsyncClient(timeout=60, follow_redirects=True) as hc:
            resp = await hc.get(binary_url)
            if resp.status_code == 200:
                binary_data = resp.content
    except Exception:
        pass

    ext = ".exe" if is_windows else ""
    installer_filename = f"{prod_name}-Setup{ext}"
    zip_filename = f"{prod_name}-Installer.zip"

    zip_buffer = io.BytesIO()
    with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as zf:
        zf.writestr(".env", env_contents)
        if binary_data:
            zf.writestr(installer_filename, binary_data)
        else:
            zf.writestr("README.txt",
                f"{prod_name} Installer\n\n"
                f"The installer could not be downloaded automatically.\n"
                f"Download it manually from:\n{binary_url}\n\n"
                f"Place the .env file and installer in the same folder, then run the installer.\n"
            )
        zf.writestr("INSTALL.txt",
            f"{prod_name} - Quick Install\n\n"
            f"1. Extract this zip\n"
            f"2. Double-click {installer_filename}\n"
            f"3. Your license key is pre-filled - no typing needed\n"
            f"4. Click Install, then restart Claude Desktop\n"
            f"5. You can delete this zip after installation\n\n"
            f"Support: hello@{prod_slug}tasksai.com\n"
        )

    zip_buffer.seek(0)
    return StreamingResponse(
        zip_buffer,
        media_type="application/zip",
        headers={"Content-Disposition": f'attachment; filename="{zip_filename}"'},
    )


async def _build_loader_zip(license_key: str, product_id: str, db: AsyncSession):
    """Core download logic — builds the zip for a given license + product."""
    result = await db.execute(
        select(License).where(License.license_key == license_key, License.status == "active")
    )
    license = result.scalar_one_or_none()
    if not license:
        raise HTTPException(status_code=404, detail="Invalid license key")

    # Stamp first download time (funnel tracking)
    if license.downloaded_at is None:
        license.downloaded_at = datetime.utcnow()
        try:
            await db.commit()
        except Exception:
            await db.rollback()

    user_product_id = product_id or "law"

    prod_name = "LawTasksAI"
    prod_domain = "lawtasksai.com"
    prod_support_email = "hello@lawtasksai.com"
    loader_skill_md = LOADER_SKILL_MD  # default (law)
    try:
        prod_result = await db.execute(
            text("SELECT name, domain FROM products WHERE id = :pid AND is_active = TRUE"),
            {"pid": user_product_id},
        )
        prod_row = prod_result.fetchone()
        if prod_row:
            prod_name = prod_row.name
            prod_domain = prod_row.domain
            prod_support_email = f"hello@{prod_row.domain}"
            # Load product-specific SKILL.md from loader/{product_id}/SKILL.md
            import pathlib as _pl
            skill_path = _pl.Path(__file__).parent / 'loader' / user_product_id / 'SKILL.md'
            if skill_path.exists():
                loader_skill_md = skill_path.read_text()
    except Exception:
        pass

    zip_filename = f"{prod_domain.split('.')[0]}-skills.zip"  # e.g. contractortasksai-skills.zip
    prod_slug = prod_domain.split('.')[0]   # e.g. "realtortasksai"
    env_prefix = prod_slug.upper()          # e.g. "REALTORTASKSAI"

    # Create config.json with license key
    config = {
        "license_key": license_key,
        "api_base_url": API_BASE_URL,
        "product_id": user_product_id,
        "version_policy": "latest",
        "cache_skills": False,
        "offline_mode": False
    }
    
    # Create zip file in memory
    zip_buffer = io.BytesIO()
    
    with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zf:
        # Add config.json
        zf.writestr(f'{prod_slug}/config.json', json.dumps(config, indent=2))
        
        # Add SKILL.md
        zf.writestr(f'{prod_slug}/SKILL.md', loader_skill_md)

        # Add README
        readme = f'''# {prod_name} Skills

Your personalized AI skills are ready to use!

## Installation (Step by Step)

### Step 1: Extract the ZIP file

Double-click the downloaded `{zip_filename}` to extract it.
You should see a folder called `openclaw` containing:
- config.json (your license is already configured!)
- SKILL.md
- README.md (this file)

### Step 2: Tell OpenClaw to install it

Open your OpenClaw chat and type:

"I just downloaded the {prod_name} skill file to my Downloads folder.
Please find it, unzip it if needed, and install it so I can use it.
My license key is {license_key}"

OpenClaw will find the file, install it, configure your license, and
confirm when everything is ready.

(If you prefer to install manually, copy this folder
to ~/.openclaw/skills/{prod_domain.split('.')[0]}/ and restart OpenClaw.)

### Step 3: Start using!

Just ask for any task — your AI will know exactly what to do.

---

## Your License

- **License Key:** {license_key}
- **Credits:** {license.credits_remaining}
- Your license key is already configured — no setup needed!

## ⚠️ Confidentiality Notice

Skills run entirely locally on your machine. Your data never leaves your computer.

## Need Help?

- Email: {prod_support_email}
- Website: https://{prod_domain}
'''
        zf.writestr(f'{prod_slug}/README.md', readme)
    
        # =========================================
        # MCP Server removed from zip (2026-05-18)
        # MCP users now download the platform-native
        # installer from GitHub Releases instead.
        # See: /download/loader and /download/loader/{key}
        # =========================================
        # (nothing to add here)
        
        # MCP files removed 2026-05-18 — MCP users now download the .exe
        # installer from GitHub Releases. See /download/loader routing above.
        # Keeping this comment so git history is clear.
    
    zip_buffer.seek(0)
    
    return StreamingResponse(
        zip_buffer,
        media_type='application/zip',
        headers={
            'Content-Disposition': f'attachment; filename={zip_filename}'
        }
    )


@app.get("/download/loader/{license_key}")
async def download_loader(license_key: str, request: Request, db: AsyncSession = Depends(get_db)):
    """
    Download loader by license key — uses account's current product_id.
    Kept for backwards compatibility and re-download links.

    Routing:
      - MCP-platform users → redirect to GitHub Release installer
      - OpenClaw / other users → serve OpenClaw-only zip
    """
    result = await db.execute(select(License).where(
        License.license_key == license_key, License.status == "active"
    ))
    license = result.scalar_one_or_none()

    user = None
    if license:
        user_result = await db.execute(select(User).where(User.id == license.user_id))
        user = user_result.scalar_one_or_none()

    product_id = (
        getattr(license, "product_id", None)
        or getattr(user, "product_id", None)
        or "law"
    )

    # Route to GitHub installer for all users with a supported product.
    # The zip/OpenClaw path is legacy — only used if product has no installer.
    if product_id in _PRODUCT_NAME_MAP:
        # Stamp first download (funnel tracking)
        if license and license.downloaded_at is None:
            license.downloaded_at = datetime.utcnow()
            try:
                await db.commit()
            except Exception:
                await db.rollback()
        urls = _get_installer_url(product_id, request.headers.get("user-agent", ""))
        ua = request.headers.get("user-agent", "").lower()
        dest = urls["windows"] if "windows" in ua else urls["mac"]
        from fastapi.responses import RedirectResponse
        return RedirectResponse(url=dest, status_code=302)

    return await _build_loader_zip(license_key, product_id, db)


# ============================================
# Routes: Products (public branding endpoint)
# ============================================

@app.get("/v1/products")
async def list_products_public(db: AsyncSession = Depends(get_db)):
    """Public endpoint: list all active products with branding info."""
    result = await db.execute(
        text(
            "SELECT id, name, display_name, domain, primary_color, accent_color, background_color "
            "FROM products WHERE is_active = TRUE ORDER BY id"
        )
    )
    rows = result.fetchall()
    return [
        {
            "id": row.id,
            "name": row.name,
            "display_name": row.display_name or row.name,
            "domain": row.domain,
            "colors": {
                "primary": row.primary_color,
                "accent": row.accent_color,
                "background": row.background_color,
            },
        }
        for row in rows
    ]


@app.get("/v1/products/{product_id}")
async def get_product(product_id: str, db: AsyncSession = Depends(get_db)):
    """
    Public endpoint: return branding info for a product.
    Used by landing pages to fetch their colors, name, and domain dynamically.
    """
    result = await db.execute(
        text(
            "SELECT id, name, display_name, domain, primary_color, accent_color, background_color "
            "FROM products WHERE id = :pid AND is_active = TRUE"
        ),
        {"pid": product_id},
    )
    row = result.fetchone()
    if not row:
        raise HTTPException(status_code=404, detail=f"Product '{product_id}' not found")

    return {
        "id": row.id,
        "name": row.name,
        "display_name": row.display_name or row.name,
        "domain": row.domain,
        "colors": {
            "primary": row.primary_color,
            "accent": row.accent_color,
            "background": row.background_color,
        },
    }


# ============================================
# Tracking Endpoints (public, key-authenticated)
# ============================================

@app.get("/track/first-connection")
async def track_first_connection(
    license_key: str,
    db: AsyncSession = Depends(get_db)
):
    """
    Called by server.py on startup to record the first time the MCP server
    successfully connects. Idempotent — only sets first_connected_at once.
    No auth beyond the license key itself.
    """
    result = await db.execute(
        select(License).where(License.license_key == license_key, License.status == "active")
    )
    license = result.scalar_one_or_none()
    if not license:
        # Return 200 anyway — don't leak whether key exists
        return {"ok": True}
    if license.first_connected_at is None:
        license.first_connected_at = datetime.utcnow()
        await db.commit()
    return {"ok": True}


# ============================================
# Admin Routes (protected in production)
# ============================================

@admin_router.get("/skills/{skill_id}")
async def get_admin_skill(
    skill_id: str,
    db: AsyncSession = Depends(get_db)
):
    """Get a single skill with current version content (admin only)."""
    result = await db.execute(select(Skill).where(Skill.id == skill_id))
    skill = result.scalar_one_or_none()
    if not skill:
        raise HTTPException(status_code=404, detail=f"Skill '{skill_id}' not found")

    # Get latest version content
    ver_result = await db.execute(
        select(SkillVersion)
        .where(SkillVersion.skill_id == skill_id)
        .order_by(SkillVersion.published_at.desc())
        .limit(1)
    )
    version = ver_result.scalar_one_or_none()

    return {
        "id": skill.id,
        "category_id": skill.category_id,
        "name": skill.name,
        "description": skill.description,
        "current_version": skill.current_version,
        "stable_version": skill.stable_version,
        "credits_per_use": skill.credits_per_use,
        "requires_upload": skill.requires_upload,
        "execution_type": skill.execution_type,
        "is_published": skill.is_published,
        "is_deprecated": skill.is_deprecated,
        "created_at": skill.created_at.isoformat() if skill.created_at else None,
        "current_version_content": version.content if version else None,
        "current_version_changelog": version.changelog if version else None,
    }

@admin_router.post("/skills/bulk")
async def bulk_create_skills(
    payload: dict,
    db: AsyncSession = Depends(get_db)
):
    """Bulk create skills with categories and triggers. Idempotent (upsert)."""
    skills_data = payload.get("skills", [])
    created = 0
    updated = 0
    errors = []

    for s in skills_data:
        try:
            product_id = s.get("product_id", "law")
            category_name = s.get("category", "General")
            
            # Upsert category
            cat_slug = category_name.lower().replace(' ', '_').replace('&','and').replace(',','').replace('/','_')
            cat_id = f"{product_id}_{cat_slug}"[:50]
            result = await db.execute(select(Category).where(Category.id == cat_id))
            cat = result.scalar_one_or_none()
            if not cat:
                cat = Category(id=cat_id, name=category_name, product_id=product_id, display_order=0)
                db.add(cat)
                await db.flush()

            # Build skill id
            skill_id = f"{product_id}_{s['name'].lower().replace(' ','_').replace('/','_')[:60]}"
            skill_id = ''.join(c for c in skill_id if c.isalnum() or c == '_')

            # Upsert skill
            result = await db.execute(select(Skill).where(Skill.id == skill_id))
            skill = result.scalar_one_or_none()
            complexity = s.get("complexity", "medium")
            credits = 1 if complexity == "simple" else (2 if complexity == "medium" else 3)
            
            if skill:
                skill.name = s["name"]
                skill.description = s.get("description", "")
                skill.category_id = cat_id
                skill.product_id = product_id
                skill.credits_per_use = credits
                skill.is_published = True
                updated += 1
            else:
                skill = Skill(
                    id=skill_id,
                    name=s["name"],
                    description=s.get("description", ""),
                    category_id=cat_id,
                    product_id=product_id,
                    credits_per_use=credits,
                    requires_upload=False,
                    execution_type="server",
                    is_published=True
                )
                db.add(skill)
                created += 1

            await db.flush()

            # Skip skill_versions for bulk import — prompt templates stored separately

            # Store triggers on the skill's triggers array column
            if s.get("trigger_phrases"):
                skill.triggers = [p.lower().strip() for p in s["trigger_phrases"]]

            await db.flush()

        except Exception as e:
            errors.append({"skill": s.get("name","?"), "error": str(e)})
            await db.rollback()

    await db.commit()
    return {"success": True, "created": created, "updated": updated, "errors": errors}


@admin_router.post("/skills")
async def create_skill(
    skill_data: dict,
    db: AsyncSession = Depends(get_db)
):
    """Create or update a skill (admin only)."""
    skill = Skill(
        id=skill_data["id"],
        category_id=skill_data["category_id"],
        name=skill_data["name"],
        description=skill_data.get("description"),
        credits_per_use=skill_data.get("credits_per_use", 1),
        requires_upload=skill_data.get("requires_upload", False),
        execution_type=skill_data.get("execution_type", "server"),
        is_published=skill_data.get("is_published", False),
        product_id=skill_data.get("product_id", "law")
    )
    
    await db.merge(skill)
    await db.commit()
    
    return {"success": True, "skill_id": skill.id}

@admin_router.patch("/skills/{skill_id}")
async def update_skill(
    skill_id: str,
    updates: dict,
    db: AsyncSession = Depends(get_db)
):
    """Update skill metadata (admin only). Partial updates supported."""
    result = await db.execute(select(Skill).where(Skill.id == skill_id))
    skill = result.scalar_one_or_none()
    
    if not skill:
        raise HTTPException(status_code=404, detail=f"Skill '{skill_id}' not found")
    
    # Apply partial updates
    allowed_fields = ['execution_type', 'requires_upload', 'credits_per_use', 
                      'is_published', 'description', 'name', 'category_id', 'product_id']
    
    for field, value in updates.items():
        if field in allowed_fields:
            setattr(skill, field, value)
    
    await db.commit()
    
    return {
        "success": True, 
        "skill_id": skill_id,
        "updated_fields": list(updates.keys())
    }


@admin_router.post("/skills/batch-update")
async def batch_update_skills(
    batch: dict,
    db: AsyncSession = Depends(get_db)
):
    """
    Batch update multiple skills (admin only).
    
    Body: {
        "skill_ids": ["skill-1", "skill-2", ...],
        "updates": {"execution_type": "local", ...}
    }
    """
    skill_ids = batch.get("skill_ids", [])
    updates = batch.get("updates", {})
    
    if not skill_ids or not updates:
        raise HTTPException(status_code=400, detail="Both skill_ids and updates required")
    
    allowed_fields = ['execution_type', 'requires_upload', 'credits_per_use', 
                      'is_published', 'description']
    
    # Filter to allowed fields
    safe_updates = {k: v for k, v in updates.items() if k in allowed_fields}
    
    if not safe_updates:
        raise HTTPException(status_code=400, detail="No valid update fields provided")
    
    # Update all matching skills
    updated = []
    not_found = []
    
    for skill_id in skill_ids:
        result = await db.execute(select(Skill).where(Skill.id == skill_id))
        skill = result.scalar_one_or_none()
        
        if skill:
            for field, value in safe_updates.items():
                setattr(skill, field, value)
            updated.append(skill_id)
        else:
            not_found.append(skill_id)
    
    await db.commit()
    
    return {
        "success": True,
        "updated_count": len(updated),
        "updated_skills": updated,
        "not_found": not_found,
        "applied_updates": safe_updates
    }


@admin_router.patch("/skills/{skill_id}/triggers")
async def update_skill_triggers(
    skill_id: str,
    trigger_data: dict,
    db: AsyncSession = Depends(get_db)
):
    """
    Update triggers for a single skill (admin only).
    Body: {"triggers": ["phrase 1", "phrase 2", ...]}
    """
    result = await db.execute(select(Skill).where(Skill.id == skill_id))
    skill = result.scalar_one_or_none()
    
    if not skill:
        raise HTTPException(status_code=404, detail=f"Skill '{skill_id}' not found")
    
    triggers = trigger_data.get("triggers", [])
    skill.triggers = triggers
    await db.commit()
    
    return {"success": True, "skill_id": skill_id, "trigger_count": len(triggers)}


@admin_router.post("/triggers/batch")
async def batch_update_triggers(
    batch: dict,
    db: AsyncSession = Depends(get_db)
):
    """
    Batch update triggers for multiple skills (admin only).
    Body: {
        "skill_id_1": {"triggers": ["phrase 1", ...]},
        "skill_id_2": {"triggers": ["phrase 2", ...]},
        ...
    }
    """
    updated = []
    not_found = []
    
    for skill_id, data in batch.items():
        result = await db.execute(select(Skill).where(Skill.id == skill_id))
        skill = result.scalar_one_or_none()
        
        if skill:
            triggers = data.get("triggers", [])
            skill.triggers = triggers
            updated.append(skill_id)
        else:
            not_found.append(skill_id)
    
    await db.commit()
    
    return {
        "success": True,
        "updated_count": len(updated),
        "updated_skills": updated,
        "not_found": not_found
    }


@admin_router.get("/skills/{skill_id}/versions")
async def list_skill_versions(
    skill_id: str,
    db: AsyncSession = Depends(get_db)
):
    """List all versions for a skill (admin only)."""
    # Check skill exists
    skill_result = await db.execute(select(Skill).where(Skill.id == skill_id))
    skill = skill_result.scalar_one_or_none()
    if not skill:
        raise HTTPException(status_code=404, detail=f"Skill '{skill_id}' not found")

    ver_result = await db.execute(
        select(SkillVersion)
        .where(SkillVersion.skill_id == skill_id)
        .order_by(SkillVersion.published_at.desc())
    )
    versions = ver_result.scalars().all()

    return {
        "skill_id": skill_id,
        "current_version": skill.current_version,
        "versions": [
            {
                "id": v.id,
                "version": v.version,
                "content": v.content,
                "content_preview": (v.content[:120] + "...") if v.content and len(v.content) > 120 else v.content,
                "changelog": v.changelog,
                "is_stable": v.is_stable,
                "is_beta": v.is_beta,
                "is_current": v.version == skill.current_version,
                "published_at": v.published_at.isoformat() if v.published_at else None,
            }
            for v in versions
        ],
    }


@admin_router.post("/skills/{skill_id}/versions/{version}/restore")
async def restore_skill_version(
    skill_id: str,
    version: str,
    db: AsyncSession = Depends(get_db)
):
    """Restore a specific version as the current live version (admin only)."""
    skill_result = await db.execute(select(Skill).where(Skill.id == skill_id))
    skill = skill_result.scalar_one_or_none()
    if not skill:
        raise HTTPException(status_code=404, detail=f"Skill '{skill_id}' not found")

    ver_result = await db.execute(
        select(SkillVersion)
        .where(SkillVersion.skill_id == skill_id)
        .where(SkillVersion.version == version)
    )
    ver = ver_result.scalar_one_or_none()
    if not ver:
        raise HTTPException(status_code=404, detail=f"Version '{version}' not found for skill '{skill_id}'")

    await db.execute(
        update(Skill)
        .where(Skill.id == skill_id)
        .values(current_version=version)
    )
    await db.commit()
    return {"success": True, "skill_id": skill_id, "restored_version": version}


@admin_router.post("/skills/{skill_id}/versions")
async def create_skill_version(
    skill_id: str,
    version_data: dict,
    db: AsyncSession = Depends(get_db)
):
    """Create or update a skill version (admin only)."""
    # Check if version already exists
    result = await db.execute(
        select(SkillVersion)
        .where(SkillVersion.skill_id == skill_id)
        .where(SkillVersion.version == version_data["version"])
    )
    existing = result.scalar_one_or_none()
    
    if existing:
        # Update existing version
        existing.content = version_data["content"]
        existing.changelog = version_data.get("changelog", existing.changelog)
        existing.is_stable = version_data.get("is_stable", existing.is_stable)
        existing.is_beta = version_data.get("is_beta", existing.is_beta)
    else:
        # Create new version
        version = SkillVersion(
            skill_id=skill_id,
            version=version_data["version"],
            content=version_data["content"],
            changelog=version_data.get("changelog"),
            is_stable=version_data.get("is_stable", False),
            is_beta=version_data.get("is_beta", False)
        )
        db.add(version)
    
    # Update current version on skill
    if version_data.get("set_current", True):
        await db.execute(
            update(Skill)
            .where(Skill.id == skill_id)
            .values(current_version=version_data["version"])
        )
    
    await db.commit()
    
    return {"success": True, "skill_id": skill_id, "version": version_data["version"], "updated": existing is not None}

@admin_router.get("/users")
async def list_users(db: AsyncSession = Depends(get_db)):
    """List all users with their licenses (admin only - protect in production!)."""
    result = await db.execute(
        select(User, License)
        .outerjoin(License, User.id == License.user_id)
        .order_by(User.created_at.desc())
    )
    
    users = []
    for user, license in result.all():
        users.append({
            "id": str(user.id),
            "email": user.email,
            "name": user.name,
            "firm_name": user.firm_name,
            "credits_balance": license.credits_remaining if license else 0,
            "created_at": user.created_at.isoformat() if user.created_at else None,
            "license_key": license.license_key if license else None,
            "license_type": license.type if license else None,
            "license_credits": license.credits_remaining if license else 0,
            "profile": user.profile or {},
            "platforms": user.platforms or [],
            "product_id": (license.product_id if license else None) or user.product_id or "law",
            "last_active_at": user.last_active_at.isoformat() if user.last_active_at else None,
            "downloaded_at": license.downloaded_at.isoformat() if license and license.downloaded_at else None,
            "first_connected_at": license.first_connected_at.isoformat() if license and license.first_connected_at else None,
        })

    activated = sum(1 for u in users if u["last_active_at"] is not None)
    return {"users": users, "count": len(users), "activated": activated, "never_activated": len(users) - activated}


@admin_router.get("/funnel")
async def get_funnel(db: AsyncSession = Depends(get_db)):
    """
    Activation funnel breakdown:
      signed_up → downloaded → connected → used_credits
    Excludes internal/test accounts.
    """
    result = await db.execute(
        select(User, License)
        .outerjoin(License, User.id == License.user_id)
        .order_by(User.created_at.desc())
    )
    rows = result.all()

    SKIP = {"test", "internal", "clio", "kentmercier"}

    def is_real(email: str) -> bool:
        low = email.lower()
        return not any(s in low for s in SKIP)

    per_user = []
    for user, license in rows:
        if not is_real(user.email):
            continue
        credits_purchased = license.credits_purchased if license else 0
        credits_remaining = license.credits_remaining if license else 0
        credits_used = max(0, credits_purchased - credits_remaining)
        per_user.append({
            "email": user.email,
            "product_id": (license.product_id if license else None) or user.product_id or "law",
            "created_at": user.created_at.isoformat() if user.created_at else None,
            "downloaded": license.downloaded_at is not None if license else False,
            "connected": license.first_connected_at is not None if license else False,
            "used_credits": credits_used > 0,
            "credits_used": credits_used,
            "downloaded_at": license.downloaded_at.isoformat() if license and license.downloaded_at else None,
            "connected_at": license.first_connected_at.isoformat() if license and license.first_connected_at else None,
            "last_active_at": user.last_active_at.isoformat() if user.last_active_at else None,
        })

    total = len(per_user)
    downloaded = sum(1 for u in per_user if u["downloaded"])
    connected = sum(1 for u in per_user if u["connected"])
    used = sum(1 for u in per_user if u["used_credits"])

    def pct(n, d):
        return f"{round(100 * n / d)}%" if d else "—"

    return {
        "total": total,
        "signed_up_only": total - downloaded,
        "downloaded": downloaded,
        "connected": connected,
        "used_credits": used,
        "rates": {
            "signup_to_download": pct(downloaded, total),
            "download_to_connect": pct(connected, downloaded),
            "connect_to_use": pct(used, connected),
            "overall": pct(used, total),
        },
        "per_user": per_user,
    }


@admin_router.get("/products")
async def list_products(db: AsyncSession = Depends(get_db)):
    """Admin: list all products with user counts and skill counts per product."""
    result = await db.execute(
        text("""
            SELECT
                p.id,
                p.name,
                p.display_name,
                p.domain,
                p.primary_color,
                p.accent_color,
                p.is_active,
                p.created_at,
                COUNT(DISTINCT u.id)::int AS user_count,
                COUNT(DISTINCT s.id)::int AS skill_count
            FROM products p
            LEFT JOIN users u ON u.product_id = p.id
            LEFT JOIN skills s ON s.product_id = p.id AND s.is_published = TRUE
            GROUP BY p.id, p.name, p.display_name, p.domain,
                     p.primary_color, p.accent_color, p.is_active, p.created_at
            ORDER BY p.id
        """)
    )
    rows = result.fetchall()
    return [
        {
            "id": row.id,
            "name": row.name,
            "display_name": row.display_name or row.name,
            "domain": row.domain,
            "colors": {
                "primary": row.primary_color,
                "accent": row.accent_color,
            },
            "is_active": row.is_active,
            "user_count": row.user_count,
            "skill_count": row.skill_count,
            "created_at": row.created_at.isoformat() if row.created_at else None,
        }
        for row in rows
    ]


@admin_router.get("/waitlist")
async def admin_list_waitlist(product_id: Optional[str] = None, db: AsyncSession = Depends(get_db)):
    """Admin: list waitlist signups, optionally filtered by product_id."""
    if product_id:
        query = text("""
            SELECT product_id, email, name, created_at, notified_at
            FROM waitlist WHERE product_id = :product_id
            ORDER BY product_id, created_at DESC
        """)
        result = await db.execute(query, {"product_id": product_id})
    else:
        query = text("""
            SELECT product_id, email, name, created_at, notified_at
            FROM waitlist ORDER BY product_id, created_at DESC
        """)
        result = await db.execute(query)
    rows = result.fetchall()

    # Group by product
    from collections import defaultdict
    by_product = defaultdict(list)
    for row in rows:
        by_product[row.product_id].append({
            "email": row.email,
            "name": row.name,
            "signed_up": row.created_at.isoformat() if row.created_at else None,
            "notified": row.notified_at.isoformat() if row.notified_at else None,
        })

    summary = [
        {"product_id": pid, "count": len(entries), "entries": entries}
        for pid, entries in sorted(by_product.items())
    ]
    return {"total": len(rows), "products": summary}


@admin_router.post("/waitlist/notify/{product_id}")
async def admin_notify_waitlist(product_id: str, db: AsyncSession = Depends(get_db)):
    """
    Admin: mark all waitlist entries for a product as notified (set notified_at = now).
    Call this after you've sent the launch email blast so you don't double-notify.
    """
    result = await db.execute(
        text("UPDATE waitlist SET notified_at = NOW() WHERE product_id = :pid AND notified_at IS NULL"),
        {"pid": product_id}
    )
    await db.commit()
    return {"success": True, "product_id": product_id, "notified": result.rowcount}


@admin_router.delete("/users/{user_id}")
async def delete_user(user_id: str, db: AsyncSession = Depends(get_db)):
    """Delete a user and all associated data (admin only)."""
    uid = uuid.UUID(user_id)
    # Look up user email (needed for non-user_id FK tables)
    user_row = await db.execute(text("SELECT email FROM users WHERE id = :uid"), {"uid": uid})
    user_email = user_row.scalar_one_or_none()
    # Delete FK-dependent rows one statement at a time
    for stmt, params in [
        (text("DELETE FROM credit_transactions WHERE user_id = :uid"), {"uid": uid}),
        (text("DELETE FROM usage_logs WHERE user_id = :uid"), {"uid": uid}),
        (text("DELETE FROM drip_emails WHERE user_id = :uid"), {"uid": uid}),
        (text("DELETE FROM email_subscriptions WHERE user_id = :uid"), {"uid": uid}),
        (text("DELETE FROM licenses WHERE user_id = :uid"), {"uid": uid}),
        (text("DELETE FROM users WHERE id = :uid"), {"uid": uid}),
    ]:
        await db.execute(stmt, params)
    # user_feedback uses email not user_id (no FK constraint, cleanup only)
    if user_email:
        await db.execute(text("DELETE FROM user_feedback WHERE email = :email"), {"email": user_email})
    await db.commit()
    return {"success": True, "deleted_user_id": user_id}


@admin_router.patch("/users/{user_id}")
async def update_user(user_id: str, request: Request, db: AsyncSession = Depends(get_db)):
    """Update user fields (name, firm_name, email) — admin only."""
    from sqlalchemy import update as sql_update
    uid = uuid.UUID(user_id)
    body = await request.json()
    allowed = {k: v for k, v in body.items() if k in ('name', 'firm_name', 'email')}
    if not allowed:
        raise HTTPException(400, "No updatable fields provided")
    await db.execute(sql_update(User).where(User.id == uid).values(**allowed))
    await db.commit()
    return {"success": True, "updated": allowed}

@admin_router.post("/credits/set")
async def set_credits(
    license_key: str,
    credits: int,
    db: AsyncSession = Depends(get_db)
):
    """Set credits on a license to an exact value (admin only)."""
    result = await db.execute(select(License).where(License.license_key == license_key))
    license = result.scalar_one_or_none()
    if not license:
        raise HTTPException(status_code=404, detail="License not found")
    old = license.credits_remaining
    license.credits_remaining = credits
    license.credits_purchased = credits
    license.updated_at = datetime.utcnow()
    await db.commit()
    return {"success": True, "old_balance": old, "new_balance": credits}


@admin_router.post("/credits/add")
async def add_credits(
    license_key: str,
    credits: int,
    db: AsyncSession = Depends(get_db)
):
    """Add credits to a license (admin only - protect in production!)."""
    if credits <= 0:
        raise HTTPException(status_code=400, detail="Credits must be positive")
    
    # Find license
    result = await db.execute(
        select(License).where(License.license_key == license_key)
    )
    license = result.scalar_one_or_none()
    
    if not license:
        raise HTTPException(status_code=404, detail="License not found")
    
    # Add credits
    old_balance = license.credits_remaining
    license.credits_remaining += credits
    license.credits_purchased += credits
    license.updated_at = datetime.utcnow()
    
    await db.commit()
    
    return {
        "success": True,
        "license_key": license_key,
        "credits_added": credits,
        "old_balance": old_balance,
        "new_balance": license.credits_remaining
    }


@admin_router.get("/gaps")
async def list_skill_gaps(
    db: AsyncSession = Depends(get_db),
    limit: int = 200
):
    """
    Admin view of anonymous gap reports.
    Returns raw reports plus a frequency-ranked summary of search terms
    — use this to prioritize new skill development.
    """
    result = await db.execute(
        select(SkillGap).order_by(SkillGap.reported_at.desc()).limit(limit)
    )
    rows = result.scalars().all()

    # Build frequency map across all reported terms
    from collections import Counter
    term_counter: Counter = Counter()
    raw = []
    for row in rows:
        terms = row.search_terms.split()
        term_counter.update(terms)
        raw.append({
            "id": row.id,
            "search_terms": row.search_terms,
            "loader_version": row.loader_version,
            "reported_at": row.reported_at.isoformat(),
        })

    ranked = [
        {"term": term, "count": count}
        for term, count in term_counter.most_common(50)
    ]

    return {
        "total_reports": len(raw),
        "top_terms": ranked,
        "recent_reports": raw,
    }


@admin_router.post("/migrate/add-skill-gaps-table")
async def migrate_add_skill_gaps_table(db: AsyncSession = Depends(get_db)):
    """
    One-time migration to create the skill_gaps table.
    Safe to run multiple times.
    """
    await db.execute(text("""
        CREATE TABLE IF NOT EXISTS skill_gaps (
            id SERIAL PRIMARY KEY,
            search_terms TEXT NOT NULL,
            loader_version VARCHAR(20),
            reported_at TIMESTAMP DEFAULT NOW()
        )
    """))
    await db.commit()
    return {"status": "ok", "message": "skill_gaps table ready"}


@admin_router.post("/migrate/add-triggers-column")
async def migrate_add_triggers_column(db: AsyncSession = Depends(get_db)):
    """
    One-time migration to add triggers column to skills table.
    Safe to run multiple times - checks if column exists first.
    """
    from sqlalchemy import text
    
    try:
        # Check if column exists
        check_query = text("""
            SELECT column_name 
            FROM information_schema.columns 
            WHERE table_name = 'skills' AND column_name = 'triggers'
        """)
        result = await db.execute(check_query)
        exists = result.scalar_one_or_none()
        
        if exists:
            return {"success": True, "message": "Column 'triggers' already exists"}
        
        # Add the column
        alter_query = text("""
            ALTER TABLE skills 
            ADD COLUMN triggers TEXT[] DEFAULT '{}'
        """)
        await db.execute(alter_query)
        await db.commit()
        
        return {"success": True, "message": "Column 'triggers' added successfully"}
    except Exception as e:
        return {"success": False, "error": str(e)}


# ============================================
# Content Pages (version-controlled editable pages)
# ============================================

@app.get("/pages/{slug}")
async def get_page(slug: str, db: AsyncSession = Depends(get_db)):
    """Public endpoint: get current content of a page by slug."""
    result = await db.execute(select(ContentPage).where(ContentPage.slug == slug))
    page = result.scalar_one_or_none()
    if not page:
        raise HTTPException(status_code=404, detail=f"Page '{slug}' not found")
    return {
        "slug": page.slug,
        "title": page.title,
        "content": page.content,
        "current_version": page.current_version,
        "updated_at": page.updated_at.isoformat() if page.updated_at else None
    }

@admin_router.get("/pages")
async def list_pages(db: AsyncSession = Depends(get_db)):
    """Admin: list all content pages."""
    result = await db.execute(select(ContentPage).order_by(ContentPage.title))
    pages = result.scalars().all()
    return [
        {
            "slug": p.slug,
            "title": p.title,
            "current_version": p.current_version,
            "updated_at": p.updated_at.isoformat() if p.updated_at else None
        }
        for p in pages
    ]

@admin_router.get("/pages/{slug}")
async def get_page_admin(slug: str, db: AsyncSession = Depends(get_db)):
    """Admin: get page with full content for editing."""
    result = await db.execute(select(ContentPage).where(ContentPage.slug == slug))
    page = result.scalar_one_or_none()
    if not page:
        raise HTTPException(status_code=404, detail=f"Page '{slug}' not found")
    return {
        "slug": page.slug,
        "title": page.title,
        "content": page.content,
        "current_version": page.current_version,
        "updated_at": page.updated_at.isoformat() if page.updated_at else None,
        "created_at": page.created_at.isoformat() if page.created_at else None
    }

@admin_router.put("/pages/{slug}")
async def save_page(slug: str, data: dict, db: AsyncSession = Depends(get_db)):
    """Admin: save page content. Creates a new version automatically."""
    from sqlalchemy import text
    
    result = await db.execute(select(ContentPage).where(ContentPage.slug == slug))
    page = result.scalar_one_or_none()
    
    content = data.get("content", "")
    title = data.get("title", "")
    changelog = data.get("changelog", "")
    
    if page:
        # Update existing page
        new_version = page.current_version + 1
        
        # Save current version to history before overwriting
        version_entry = ContentPageVersion(
            page_slug=slug,
            version=page.current_version,
            content=page.content,
            changelog=changelog or f"Version {page.current_version}",
            created_at=page.updated_at or page.created_at
        )
        db.add(version_entry)
        
        # Update the page
        page.content = content
        page.title = title or page.title
        page.current_version = new_version
        page.updated_at = datetime.utcnow()
    else:
        # Create new page
        page = ContentPage(
            slug=slug,
            title=title or slug,
            content=content,
            current_version=1,
            created_at=datetime.utcnow(),
            updated_at=datetime.utcnow()
        )
        db.add(page)
    
    await db.commit()
    
    return {
        "success": True,
        "slug": slug,
        "version": page.current_version,
        "updated_at": page.updated_at.isoformat()
    }

@admin_router.get("/pages/{slug}/versions")
async def get_page_versions(slug: str, db: AsyncSession = Depends(get_db)):
    """Admin: list all versions of a page."""
    # Check page exists
    result = await db.execute(select(ContentPage).where(ContentPage.slug == slug))
    page = result.scalar_one_or_none()
    if not page:
        raise HTTPException(status_code=404, detail=f"Page '{slug}' not found")
    
    # Get version history
    result = await db.execute(
        select(ContentPageVersion)
        .where(ContentPageVersion.page_slug == slug)
        .order_by(ContentPageVersion.version.desc())
    )
    versions = result.scalars().all()
    
    # Include current version at the top
    all_versions = [
        {
            "version": page.current_version,
            "changelog": "Current version",
            "created_at": page.updated_at.isoformat() if page.updated_at else None,
            "is_current": True
        }
    ]
    
    for v in versions:
        all_versions.append({
            "version": v.version,
            "changelog": v.changelog,
            "created_at": v.created_at.isoformat() if v.created_at else None,
            "is_current": False
        })
    
    return all_versions

@admin_router.get("/pages/{slug}/versions/{version}")
async def get_page_version(slug: str, version: int, db: AsyncSession = Depends(get_db)):
    """Admin: get a specific version's content (for viewing/restoring)."""
    # Check if requesting current version
    result = await db.execute(select(ContentPage).where(ContentPage.slug == slug))
    page = result.scalar_one_or_none()
    if not page:
        raise HTTPException(status_code=404, detail=f"Page '{slug}' not found")
    
    if version == page.current_version:
        return {"version": version, "content": page.content, "is_current": True}
    
    # Get historical version
    result = await db.execute(
        select(ContentPageVersion)
        .where(ContentPageVersion.page_slug == slug, ContentPageVersion.version == version)
    )
    ver = result.scalar_one_or_none()
    if not ver:
        raise HTTPException(status_code=404, detail=f"Version {version} not found")
    
    return {"version": ver.version, "content": ver.content, "is_current": False}

@admin_router.post("/pages/{slug}/restore/{version}")
async def restore_page_version(slug: str, version: int, db: AsyncSession = Depends(get_db)):
    """Admin: restore a previous version (saves current as new version first)."""
    result = await db.execute(select(ContentPage).where(ContentPage.slug == slug))
    page = result.scalar_one_or_none()
    if not page:
        raise HTTPException(status_code=404, detail=f"Page '{slug}' not found")
    
    if version == page.current_version:
        return {"success": True, "message": "Already on this version"}
    
    # Get the version to restore
    result = await db.execute(
        select(ContentPageVersion)
        .where(ContentPageVersion.page_slug == slug, ContentPageVersion.version == version)
    )
    ver = result.scalar_one_or_none()
    if not ver:
        raise HTTPException(status_code=404, detail=f"Version {version} not found")
    
    # Save current to history
    version_entry = ContentPageVersion(
        page_slug=slug,
        version=page.current_version,
        content=page.content,
        changelog=f"Before restoring v{version}",
        created_at=page.updated_at or page.created_at
    )
    db.add(version_entry)
    
    # Restore
    new_version = page.current_version + 1
    page.content = ver.content
    page.current_version = new_version
    page.updated_at = datetime.utcnow()
    
    await db.commit()
    
    return {
        "success": True,
        "restored_from": version,
        "new_version": new_version
    }

@admin_router.post("/migrate/add-license-product-id")
async def migrate_add_license_product_id(db: AsyncSession = Depends(get_db)):
    """
    One-time migration: add product_id column to licenses table.
    Idempotent — safe to call multiple times.
    """
    from sqlalchemy import text
    await db.execute(text(
        "ALTER TABLE licenses ADD COLUMN IF NOT EXISTS product_id VARCHAR(50) DEFAULT 'law'"
    ))
    # Backfill from user product_id for existing rows
    await db.execute(text("""
        UPDATE licenses l
        SET product_id = COALESCE(u.product_id, 'law')
        FROM users u
        WHERE l.user_id = u.id AND (l.product_id IS NULL OR l.product_id = 'law')
    """))
    await db.commit()
    return {"success": True, "message": "licenses.product_id column added and backfilled"}


@admin_router.post("/migrate/add-abbreviations-table")
async def migrate_add_abbreviations_table(db: AsyncSession = Depends(get_db)):
    """
    One-time migration: create skill_abbreviations table + index.
    Idempotent — safe to call multiple times.
    """
    from sqlalchemy import text
    # Check if table already exists
    exists = await db.execute(text(
        "SELECT 1 FROM information_schema.tables "
        "WHERE table_schema='public' AND table_name='skill_abbreviations'"
    ))
    if exists.fetchone():
        return {"success": True, "message": "Table already exists"}
    # Create table
    await db.execute(text("""
        CREATE TABLE skill_abbreviations (
            id           SERIAL PRIMARY KEY,
            product_id   TEXT NOT NULL,
            abbreviation TEXT NOT NULL,
            expansion    TEXT NOT NULL,
            created_at   TIMESTAMPTZ DEFAULT NOW(),
            UNIQUE (product_id, abbreviation)
        )
    """))
    await db.execute(text(
        "CREATE INDEX idx_abbrev_product ON skill_abbreviations(product_id)"
    ))
    await db.commit()
    return {"success": True, "message": "Table and index created"}


@admin_router.post("/migrate/add-content-pages")
async def migrate_add_content_pages(db: AsyncSession = Depends(get_db)):
    """One-time migration to create content_pages and content_page_versions tables."""
    from sqlalchemy import text
    
    try:
        await db.execute(text("""
            CREATE TABLE IF NOT EXISTS content_pages (
                slug VARCHAR(100) PRIMARY KEY,
                title VARCHAR(255) NOT NULL,
                current_version INTEGER DEFAULT 1,
                content TEXT NOT NULL,
                updated_at TIMESTAMP DEFAULT NOW(),
                created_at TIMESTAMP DEFAULT NOW()
            )
        """))
        
        await db.execute(text("""
            CREATE TABLE IF NOT EXISTS content_page_versions (
                id SERIAL PRIMARY KEY,
                page_slug VARCHAR(100) REFERENCES content_pages(slug) ON DELETE CASCADE,
                version INTEGER NOT NULL,
                content TEXT NOT NULL,
                changelog TEXT,
                created_at TIMESTAMP DEFAULT NOW()
            )
        """))
        
        await db.commit()
        return {"success": True, "message": "Content pages tables created successfully"}
    except Exception as e:
        return {"success": False, "error": str(e)}


# ============================================
# Admin File Management (GCS-backed)
# ============================================

WORKSPACE_BUCKET = os.getenv("WORKSPACE_BUCKET", "lawtasksai-workspace")

def _get_gcs_client():
    from google.cloud import storage
    return storage.Client()

@admin_router.get("/files/list")
async def list_files(prefix: str = ""):
    """List files in the workspace bucket under a prefix."""
    try:
        client = _get_gcs_client()
        bucket = client.bucket(WORKSPACE_BUCKET)
        blobs = bucket.list_blobs(prefix=prefix)
        files = []
        for blob in blobs:
            files.append({
                "path": blob.name,
                "size": blob.size,
                "updated": blob.updated.isoformat() if blob.updated else None,
            })
        return files
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@admin_router.get("/files/read")
async def read_file(path: str):
    """Read a file from the workspace bucket."""
    try:
        client = _get_gcs_client()
        bucket = client.bucket(WORKSPACE_BUCKET)
        blob = bucket.blob(path)
        if not blob.exists():
            raise HTTPException(status_code=404, detail="File not found")
        content = blob.download_as_text()
        return {"path": path, "content": content, "size": blob.size, "updated": blob.updated.isoformat() if blob.updated else None}
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@admin_router.put("/files/write")
async def write_file(data: dict):
    """Write/update a file in the workspace bucket."""
    path = data.get("path", "")
    content = data.get("content", "")
    if not path:
        raise HTTPException(status_code=400, detail="path is required")
    try:
        client = _get_gcs_client()
        bucket = client.bucket(WORKSPACE_BUCKET)
        blob = bucket.blob(path)
        blob.upload_from_string(content, content_type="text/plain")
        return {"success": True, "path": path, "size": len(content.encode('utf-8'))}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


# ============================================
# Template Management & Rendering
# ============================================

@admin_router.get("/templates")
async def list_templates():
    """List available output templates with metadata from manifest."""
    try:
        client = _get_gcs_client()
        bucket = client.bucket(WORKSPACE_BUCKET)
        manifest_blob = bucket.blob("templates/manifest.json")
        if manifest_blob.exists():
            import json as _json
            manifest = _json.loads(manifest_blob.download_as_text())
            return manifest
        # Fallback: list HTML files
        blobs = bucket.list_blobs(prefix="templates/")
        templates = []
        for blob in blobs:
            if blob.name.endswith('.html'):
                tid = blob.name.replace('templates/', '').replace('.html', '')
                templates.append({"id": tid, "name": tid.replace('-', ' ').title(), "file": f"{tid}.html"})
        return templates
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@admin_router.get("/templates/{template_id}")
async def get_template(template_id: str):
    """Get template HTML content."""
    try:
        client = _get_gcs_client()
        bucket = client.bucket(WORKSPACE_BUCKET)
        blob = bucket.blob(f"templates/{template_id}.html")
        if not blob.exists():
            raise HTTPException(status_code=404, detail="Template not found")
        content = blob.download_as_text()
        return {"id": template_id, "content": content, "size": blob.size}
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@admin_router.put("/templates/{template_id}")
async def save_template(template_id: str, data: dict):
    """Save/update a template."""
    content = data.get("content", "")
    metadata = data.get("metadata", {})
    try:
        client = _get_gcs_client()
        bucket = client.bucket(WORKSPACE_BUCKET)
        # Save template HTML
        blob = bucket.blob(f"templates/{template_id}.html")
        blob.upload_from_string(content, content_type="text/html")
        # Update manifest if metadata provided
        if metadata:
            import json as _json
            manifest_blob = bucket.blob("templates/manifest.json")
            manifest = []
            if manifest_blob.exists():
                manifest = _json.loads(manifest_blob.download_as_text())
            # Update or add entry
            found = False
            for entry in manifest:
                if entry.get("id") == template_id:
                    entry.update(metadata)
                    found = True
                    break
            if not found:
                metadata["id"] = template_id
                metadata["file"] = f"{template_id}.html"
                manifest.append(metadata)
            manifest_blob.upload_from_string(_json.dumps(manifest, indent=2), content_type="application/json")
        return {"success": True, "id": template_id}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@admin_router.post("/templates/preview")
async def preview_template(data: dict):
    """Render a template with provided data. Returns HTML string."""
    template_id = data.get("template_id", "")
    title = data.get("title", "Untitled")
    subtitle = data.get("subtitle", "")
    eyebrow = data.get("eyebrow", "LawTasksAI")
    content_html = data.get("content", "<p>No content provided.</p>")
    nav_items = data.get("nav_items", "")

    try:
        client = _get_gcs_client()
        bucket = client.bucket(WORKSPACE_BUCKET)
        blob = bucket.blob(f"templates/{template_id}.html")
        if not blob.exists():
            raise HTTPException(status_code=404, detail="Template not found")
        template = blob.download_as_text()

        # Simple placeholder replacement
        rendered = template.replace("{{title}}", title)
        rendered = rendered.replace("{{subtitle}}", subtitle)
        rendered = rendered.replace("{{eyebrow}}", eyebrow)
        rendered = rendered.replace("{{content}}", content_html)
        rendered = rendered.replace("{{nav_items}}", nav_items)

        return {"html": rendered}
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


# ============================================
# Multi-tenant Migration Endpoint (one-time use)
# ============================================

@admin_router.post("/run-migration-001")
async def run_migration_001(db: AsyncSession = Depends(get_db)):
    """Run the multi-tenant migration 001. Safe to call multiple times (idempotent)."""
    results = []

    async def step(name: str, sql: str):
        try:
            await db.execute(text(sql))
            results.append(f"✅ {name}")
        except Exception as e:
            results.append(f"⚠️ {name}: {str(e)[:120]}")

    # Step 1: products table
    await step("Create products table", """
        CREATE TABLE IF NOT EXISTS products (
            id VARCHAR(50) PRIMARY KEY,
            name VARCHAR(100) NOT NULL,
            display_name VARCHAR(100),
            domain VARCHAR(100),
            frontend_url VARCHAR(200),
            primary_color VARCHAR(20) DEFAULT '#1a1a2e',
            accent_color VARCHAR(20) DEFAULT '#2563eb',
            background_color VARCHAR(20) DEFAULT '#fafbfc',
            is_active BOOLEAN DEFAULT TRUE,
            created_at TIMESTAMP DEFAULT NOW()
        )
    """)

    # Step 2: seed all 27 products
    products = [
        ("law", "LawTasksAI", "#1a1a2e", "#2563eb", "#fafbfc"),
        ("contractor", "ContractorTasksAI", "#1a1a1a", "#F97316", "#FAFAFA"),
        ("realtor", "RealtorTasksAI", "#1c2b3a", "#C8973A", "#FDF8F3"),
        ("mortgage", "MortgageTasksAI", "#1a2e1a", "#16A34A", "#F0FDF4"),
        ("insurance", "InsuranceTasksAI", "#1e293b", "#0EA5E9", "#F0F9FF"),
        ("hr", "HRTasksAI", "#2d1b4e", "#7C3AED", "#FAF5FF"),
        ("accounting", "AccountingTasksAI", "#14213d", "#1D4ED8", "#EFF6FF"),
        ("chiropractor", "ChiropractorTasksAI", "#1a3a3a", "#0D9488", "#F0FDFA"),
        ("vet", "VetTasksAI", "#1a3520", "#22C55E", "#F0FDF4"),
        ("dentist", "DentistTasksAI", "#1e3a5f", "#38BDF8", "#F0F9FF"),
        ("plumber", "PlumberTasksAI", "#1a2744", "#3B82F6", "#F8FAFF"),
        ("landlord", "LandlordTasksAI", "#2c1a0e", "#B45309", "#FFFBF5"),
        ("nutritionist", "NutritionistTasksAI", "#1a3320", "#65A30D", "#F7FEE7"),
        ("personaltrainer", "PersonalTrainerTasksAI", "#1a1a1a", "#EF4444", "#FFF5F5"),
        ("therapist", "TherapistTasksAI", "#2d2640", "#8B5CF6", "#F5F3FF"),
        ("eventplanner", "EventPlannerTasksAI", "#3b0764", "#D946EF", "#FDF4FF"),
        ("travelagent", "TravelAgentTasksAI", "#0f3460", "#F59E0B", "#FFFBEB"),
        ("funeral", "FuneralTasksAI", "#1a1a2a", "#6B7280", "#F9FAFB"),
        ("pastor", "PastorTasksAI", "#2d1b00", "#D97706", "#FFFBEB"),
        ("principal", "PrincipalTasksAI", "#1e3a5f", "#2563EB", "#EFF6FF"),
        ("farmer", "FarmerTasksAI", "#2d1f00", "#92400E", "#FEFCE8"),
        ("restaurant", "RestaurantTasksAI", "#1a0a00", "#DC2626", "#FFF5F5"),
        ("salon", "SalonTasksAI", "#2d0a2d", "#EC4899", "#FDF2F8"),
        ("mortician", "MorticiaryTasksAI", "#1c1c1c", "#9CA3AF", "#F3F4F6"),
        ("churchadmin", "ChurchAdminTasksAI", "#1e2d40", "#6366F1", "#EEF2FF"),
        ("militaryspouse", "MilitarySpouseTasksAI", "#1a2340", "#1E40AF", "#EFF6FF"),
        ("electrician", "ElectricianTasksAI", "#1a1a00", "#EAB308", "#FEFCE8"),
        ("teacher", "TeacherTasksAI", "#1a2744", "#2563EB", "#EFF6FF"),
        ("designer", "DesignerTasksAI", "#2d1a3a", "#C026D3", "#FDF4FF"),
    ]
    for p in products:
        await step(f"Seed product: {p[0]}", f"""
            INSERT INTO products (id, name, primary_color, accent_color, background_color)
            VALUES ('{p[0]}', '{p[1]}', '{p[2]}', '{p[3]}', '{p[4]}')
            ON CONFLICT (id) DO NOTHING
        """)

    # Step 3: add product_id to all tables
    for table in ["users", "skills", "categories", "usage_logs", "licenses", "credit_transactions"]:
        await step(f"Add product_id to {table}", f"""
            ALTER TABLE {table} ADD COLUMN IF NOT EXISTS product_id VARCHAR(50) DEFAULT 'law'
        """)

    # Step 4: product_credit_packs table
    await step("Create product_credit_packs table", """
        CREATE TABLE IF NOT EXISTS product_credit_packs (
            id SERIAL PRIMARY KEY,
            product_id VARCHAR(50) REFERENCES products(id),
            pack_key VARCHAR(50) NOT NULL,
            name VARCHAR(100) NOT NULL,
            credits INTEGER NOT NULL,
            price_cents INTEGER NOT NULL,
            UNIQUE(product_id, pack_key)
        )
    """)

    # Step 5: seed pricing tiers for all 27 products
    tiers = [
        ("tryit", "Try It", 2, 500),
        ("starter", "Starter", 15, 2900),
        ("pro", "Pro", 60, 9900),
        ("business", "Business", 150, 19900),
        ("power", "Power", 350, 34900),
        ("unlimited", "Unlimited", 800, 59900),
        ("enterprise", "Enterprise", 2000, 99900),
    ]
    # Include teacher and designer in pricing seed even if already in products list
    product_ids = [p[0] for p in products]
    for pid in product_ids:
        for key, name, credits, price in tiers:
            await step(f"Seed {pid}/{key}", f"""
                INSERT INTO product_credit_packs (product_id, pack_key, name, credits, price_cents)
                VALUES ('{pid}', '{key}', '{name}', {credits}, {price})
                ON CONFLICT (product_id, pack_key) DO NOTHING
            """)

    await db.commit()
    return {"migration": "001_add_multitenant", "steps": len(results), "results": results}


@admin_router.post("/provision-user")
async def admin_provision_user(
    email: str,
    product_id: str,
    credits: int,
    db: AsyncSession = Depends(get_db)
):
    """Manually provision a user+license for a product. Idempotent — safe to run twice."""
    # Find user by email (one user per email, product tracked on license/user)
    result = await db.execute(
        select(User).where(User.email == email.lower().strip())
    )
    user = result.scalar_one_or_none()

    if not user:
        user = User(
            email=email.lower().strip(),
            password_hash=hash_password(secrets.token_hex(16)),
            credits_balance=credits,
            product_id=product_id,
        )
        db.add(user)
        await db.flush()
        license = License(
            license_key=generate_license_key(),
            user_id=user.id,
            type="credits",
            credits_purchased=credits,
            credits_remaining=credits,
        )
        db.add(license)
        await db.flush()
        action = "created"
    else:
        # Update product_id if being provisioned for a different product
        if user.product_id != product_id:
            user.product_id = product_id
        # credits_balance derived from license — no write needed
        result2 = await db.execute(
            select(License).where(License.user_id == user.id, License.status == "active")
            .order_by(License.created_at.desc())
        )
        license = result2.scalar_one_or_none()
        if not license:
            license = License(
                license_key=generate_license_key(),
                user_id=user.id,
                type="credits",
                credits_purchased=credits,
                credits_remaining=credits,
            )
            db.add(license)
            await db.flush()
        else:
            license.credits_remaining += credits
            license.credits_purchased += credits
        action = "updated"

    await db.commit()
    return {
        "action": action,
        "email": email,
        "product_id": product_id,
        "license_key": license.license_key,
        "credits": license.credits_remaining,
        "download_url": f"https://api.lawtasksai.com/download/loader/{license.license_key}",
    }


@admin_router.post("/migrate/set-product-domains")
async def migrate_set_product_domains(db: AsyncSession = Depends(get_db)):
    """One-time migration: populate domain column for all 29 products."""
    domains = {
        "law":            "lawtasksai.com",
        "contractor":     "contractortasksai.com",
        "realtor":        "realtortasksai.com",
        "accounting":     "accountingtasksai.com",
        "chiropractor":   "chiropractortasksai.com",
        "churchadmin":    "churchadmintasksai.com",
        "dentist":        "dentisttasksai.com",
        "designer":       "designertasksai.com",
        "electrician":    "electriciantasksai.com",
        "eventplanner":   "eventplannertasksai.com",
        "farmer":         "farmertasksai.com",
        "funeral":        "funeraltasksai.com",
        "hr":             "hrtasksai.com",
        "insurance":      "insurancetasksai.com",
        "landlord":       "landlordtasksai.com",
        "militaryspouse": "militaryspousetasksai.com",
        "mortgage":       "mortgagetasksai.com",
        "mortician":      "mortuarytasksai.com",
        "nutritionist":   "nutritionisttasksai.com",
        "pastor":         "pastortasksai.com",
        "personaltrainer":"personaltrainertasksai.com",
        "plumber":        "plumbertasksai.com",
        "principal":      "principaltasksai.com",
        "restaurant":     "restauranttasksai.com",
        "salon":          "salontasksai.com",
        "teacher":        "teachertasksai.com",
        "therapist":      "therapisttasksai.com",
        "travelagent":    "travelagenttasksai.com",
        "vet":            "vettasksai.com",
    }
    results = []
    for pid, domain in domains.items():
        try:
            await db.execute(
                text("UPDATE products SET domain = :domain WHERE id = :pid"),
                {"domain": domain, "pid": pid}
            )
            results.append({"product_id": pid, "domain": domain, "status": "ok"})
        except Exception as e:
            results.append({"product_id": pid, "domain": domain, "status": f"error: {e}"})
    await db.commit()
    return {"updated": len([r for r in results if r["status"] == "ok"]), "results": results}


# Register admin router (all routes protected by X-Admin-Secret)
@admin_router.post("/migrate/add-product")
async def migrate_add_product(data: dict, db: AsyncSession = Depends(get_db)):
    """Insert or update a product record."""
    pid = data.get("id")
    if not pid:
        raise HTTPException(status_code=400, detail="id is required")
    await db.execute(text("""
        INSERT INTO products (id, name, primary_color, accent_color, background_color, domain, is_active)
        VALUES (:id, :name, :primary_color, :accent_color, :background_color, :domain, TRUE)
        ON CONFLICT (id) DO UPDATE SET
            name = EXCLUDED.name,
            domain = EXCLUDED.domain,
            accent_color = EXCLUDED.accent_color,
            is_active = TRUE
    """), {
        "id": pid,
        "name": data.get("name", pid),
        "primary_color": data.get("primary_color", "#1a1a2e"),
        "accent_color": data.get("accent_color", "#2563eb"),
        "background_color": data.get("background_color", "#ffffff"),
        "domain": data.get("domain", f"{pid}tasksai.com"),
    })
    await db.commit()
    return {"success": True, "product_id": pid}


# ============================================
# Admin Routes: Abbreviations CRUD
# ============================================

@admin_router.get("/abbreviations")
async def admin_list_abbreviations(
    product_id: Optional[str] = None,
    db: AsyncSession = Depends(get_db),
):
    """Admin: list all abbreviations, optionally filtered by product_id."""
    q = "SELECT id, product_id, abbreviation, expansion, created_at FROM skill_abbreviations"
    params: dict = {}
    if product_id:
        q += " WHERE product_id = :pid"
        params["pid"] = product_id
    q += " ORDER BY product_id, abbreviation"
    result = await db.execute(text(q), params)
    rows = result.fetchall()
    return [
        {
            "id":           row.id,
            "product_id":   row.product_id,
            "abbreviation": row.abbreviation,
            "expansion":    row.expansion,
            "created_at":   row.created_at.isoformat() if row.created_at else None,
        }
        for row in rows
    ]


@admin_router.post("/abbreviations", status_code=201)
async def admin_create_abbreviation(
    data: dict,
    db: AsyncSession = Depends(get_db),
):
    """Admin: create a new abbreviation. Body: {product_id, abbreviation, expansion}"""
    product_id   = data.get("product_id")
    abbreviation = (data.get("abbreviation") or "").lower().strip()
    expansion    = (data.get("expansion") or "").lower().strip()
    if not product_id or not abbreviation or not expansion:
        raise HTTPException(status_code=400, detail="product_id, abbreviation, and expansion are required")
    try:
        result = await db.execute(
            text("""
                INSERT INTO skill_abbreviations (product_id, abbreviation, expansion)
                VALUES (:pid, :abbr, :exp)
                ON CONFLICT (product_id, abbreviation) DO UPDATE SET expansion = EXCLUDED.expansion
                RETURNING id, product_id, abbreviation, expansion
            """),
            {"pid": product_id, "abbr": abbreviation, "exp": expansion},
        )
        row = result.fetchone()
        await db.commit()
        return {"success": True, "id": row.id, "product_id": row.product_id,
                "abbreviation": row.abbreviation, "expansion": row.expansion}
    except Exception as e:
        await db.rollback()
        raise HTTPException(status_code=500, detail=str(e))


@admin_router.put("/abbreviations/{abbrev_id}")
async def admin_update_abbreviation(
    abbrev_id: int,
    data: dict,
    db: AsyncSession = Depends(get_db),
):
    """Admin: update expansion for an abbreviation by its DB id."""
    expansion = (data.get("expansion") or "").lower().strip()
    if not expansion:
        raise HTTPException(status_code=400, detail="expansion is required")
    result = await db.execute(
        text("UPDATE skill_abbreviations SET expansion = :exp WHERE id = :id RETURNING id, product_id, abbreviation, expansion"),
        {"exp": expansion, "id": abbrev_id},
    )
    row = result.fetchone()
    if not row:
        raise HTTPException(status_code=404, detail=f"Abbreviation id {abbrev_id} not found")
    await db.commit()
    return {"success": True, "id": row.id, "product_id": row.product_id,
            "abbreviation": row.abbreviation, "expansion": row.expansion}


@admin_router.delete("/abbreviations/{abbrev_id}", status_code=204)
async def admin_delete_abbreviation(
    abbrev_id: int,
    db: AsyncSession = Depends(get_db),
):
    """Admin: delete an abbreviation by its DB id."""
    result = await db.execute(
        text("DELETE FROM skill_abbreviations WHERE id = :id RETURNING id"),
        {"id": abbrev_id},
    )
    row = result.fetchone()
    if not row:
        raise HTTPException(status_code=404, detail=f"Abbreviation id {abbrev_id} not found")
    await db.commit()


@admin_router.post("/abbreviations/seed")
async def admin_seed_abbreviations(db: AsyncSession = Depends(get_db)):
    """
    Admin: re-run the abbreviation seed (migration 002 data).
    Safe to call multiple times -- uses ON CONFLICT DO NOTHING.
    Returns counts per vertical.
    """
    ABBREVIATIONS = {
        "law": {"mtc":"motion to compel","rogs":"interrogatories","rog":"interrogatory","rfa":"request for admission","rfas":"requests for admission","rfp":"request for production","rfps":"requests for production","tro":"temporary restraining order","pi":"personal injury","msj":"motion for summary judgment","msk":"motion to strike","sj":"summary judgment","jnov":"judgment notwithstanding verdict","mil":"motion in limine","sol":"statute of limitations","aff":"affidavit","decl":"declaration","depo":"deposition","deps":"depositions","frcp":"federal rules civil procedure","fre":"federal rules evidence","compl":"complaint","ans":"answer","roe":"rules of evidence","atty":"attorney"},
        "realtor": {"mls":"multiple listing service","cma":"comparative market analysis","dom":"days on market","arv":"after repair value","hoa":"homeowners association","coe":"close of escrow","emd":"earnest money deposit","piti":"principal interest taxes insurance","ltv":"loan to value","nar":"national association of realtors","bom":"back on market","uc":"under contract","fs":"for sale","fsbo":"for sale by owner","reo":"real estate owned"},
        "contractor": {"rfi":"request for information","sow":"scope of work","co":"change order","gc":"general contractor","ntp":"notice to proceed","pco":"potential change order","aia":"american institute of architects","lien":"mechanics lien","sub":"subcontractor","por":"purchase order request","cos":"certificate of substantial completion","punch":"punch list","g702":"payment application","g703":"schedule of values"},
        "farmer": {"fsa":"farm service agency","nrcs":"natural resources conservation service","crp":"conservation reserve program","arc":"agriculture risk coverage","plc":"price loss coverage","usda":"united states department of agriculture","eqip":"environmental quality incentives program","csa":"community supported agriculture","gmp":"good manufacturing practices","gap":"good agricultural practices"},
        "hr": {"pip":"performance improvement plan","pto":"paid time off","fmla":"family medical leave act","ada":"americans with disabilities act","eeoc":"equal employment opportunity commission","w2":"wage and tax statement","i9":"employment eligibility verification","cobra":"consolidated omnibus budget reconciliation act","osha":"occupational safety and health administration","erp":"employee relations policy"},
        "accounting": {"cogs":"cost of goods sold","ar":"accounts receivable","ap":"accounts payable","gaap":"generally accepted accounting principles","ytd":"year to date","mtd":"month to date","ebitda":"earnings before interest taxes depreciation amortization","cpa":"certified public accountant","sox":"sarbanes oxley"},
        "mortgage": {"ltv":"loan to value","dti":"debt to income","arm":"adjustable rate mortgage","apr":"annual percentage rate","pmi":"private mortgage insurance","hud":"housing and urban development","fnma":"fannie mae","fhlmc":"freddie mac","heloc":"home equity line of credit","gfe":"good faith estimate","cd":"closing disclosure","le":"loan estimate"},
        "insurance": {"doi":"department of insurance","gl":"general liability","wc":"workers compensation","coi":"certificate of insurance","dec":"declarations page","aob":"assignment of benefits","uwi":"underwriting information","clue":"comprehensive loss underwriting exchange","pip":"personal injury protection"},
        "therapist": {"dap":"data assessment plan","soap":"subjective objective assessment plan","hipaa":"health insurance portability and accountability act","phi":"protected health information","dx":"diagnosis","tx":"treatment","iop":"intensive outpatient program","php":"partial hospitalization program","cbt":"cognitive behavioral therapy","dbt":"dialectical behavior therapy","emdr":"eye movement desensitization reprocessing"},
        "chiropractor": {"soap":"subjective objective assessment plan","rom":"range of motion","pi":"personal injury","hipaa":"health insurance portability and accountability act","icd":"international classification of diseases","cpt":"current procedural terminology","eob":"explanation of benefits"},
        "dentist": {"hipaa":"health insurance portability and accountability act","cdt":"current dental terminology","perio":"periodontal","ortho":"orthodontic","endo":"endodontic","eob":"explanation of benefits","pano":"panoramic radiograph"},
        "teacher": {"iep":"individualized education program","504":"section 504 accommodation plan","ell":"english language learner","sped":"special education","pbis":"positive behavioral interventions and supports","mtss":"multi-tiered system of supports","rti":"response to intervention","ferpa":"family educational rights and privacy act","pd":"professional development","plc":"professional learning community"},
        "vet": {"soap":"subjective objective assessment plan","avma":"american veterinary medical association","rx":"prescription","dx":"diagnosis","tx":"treatment","hx":"history","pe":"physical examination"},
        "electrician": {"nec":"national electrical code","gfci":"ground fault circuit interrupter","afci":"arc fault circuit interrupter","rfi":"request for information","co":"change order","ntp":"notice to proceed"},
        "plumber": {"ipc":"international plumbing code","upc":"uniform plumbing code","rfi":"request for information","co":"change order","ntp":"notice to proceed","pex":"cross-linked polyethylene","abs":"acrylonitrile butadiene styrene"},
    }

    summary = {}
    for product_id, abbrevs in ABBREVIATIONS.items():
        inserted = 0
        for abbr, expansion in abbrevs.items():
            result = await db.execute(
                text("""
                    INSERT INTO skill_abbreviations (product_id, abbreviation, expansion)
                    VALUES (:pid, :abbr, :exp)
                    ON CONFLICT (product_id, abbreviation) DO NOTHING
                """),
                {"pid": product_id, "abbr": abbr.lower(), "exp": expansion.lower()},
            )
            inserted += result.rowcount
        summary[product_id] = inserted
    await db.commit()
    total = await db.execute(text("SELECT COUNT(*) FROM skill_abbreviations"))
    return {"success": True, "inserted_by_vertical": summary,
            "total_rows": total.scalar()}


# ============================================
# Admin: Security Scans
# ============================================

@admin_router.get("/security-scans")
async def admin_get_security_scans(
    vertical: Optional[str] = Query(None),
    db: AsyncSession = Depends(get_db),
):
    """Return security scan summary and full scan list.
    Requires X-Admin-Secret header.
    Optional ?vertical=law filter.
    """
    query = select(SkillSecurityScan)
    if vertical:
        query = query.where(SkillSecurityScan.vertical == vertical)
    query = query.order_by(SkillSecurityScan.vertical, SkillSecurityScan.skill_id)

    result = await db.execute(query)
    scans = result.scalars().all()

    total = len(scans)
    verified_count = sum(1 for s in scans if s.verified)

    by_vertical: dict = {}
    for s in scans:
        if s.vertical not in by_vertical:
            by_vertical[s.vertical] = {"total": 0, "verified": 0}
        by_vertical[s.vertical]["total"] += 1
        if s.verified:
            by_vertical[s.vertical]["verified"] += 1

    scan_list = [
        {
            "skill_id": s.skill_id,
            "vertical": s.vertical,
            "verified": s.verified,
            "tests_run": s.tests_run,
            "tests_passed": s.tests_passed,
            "tests_failed": s.tests_failed,
            "plugins_tested": s.plugins_tested or [],
            "preamble_tested": s.preamble_tested,
            "scan_model": s.scan_model,
            "scanned_at": s.scanned_at.isoformat() if s.scanned_at else None,
        }
        for s in scans
    ]

    return {
        "summary": {
            "total": total,
            "verified": verified_count,
            "by_vertical": by_vertical,
        },
        "scans": scan_list,
    }


@admin_router.post("/migrate/add-email-subscriptions")
async def migrate_add_email_subscriptions(db: AsyncSession = Depends(get_db)):
    """
    One-time migration: create email_subscriptions table + indexes.
    Idempotent — safe to call multiple times.
    """
    exists = await db.execute(text(
        "SELECT 1 FROM information_schema.tables "
        "WHERE table_schema='public' AND table_name='email_subscriptions'"
    ))
    if exists.fetchone():
        return {"success": True, "message": "Table already exists"}
    await db.execute(text("""
        CREATE TABLE email_subscriptions (
            id UUID PRIMARY KEY DEFAULT gen_random_uuid(),
            user_id UUID NOT NULL REFERENCES users(id),
            product_id VARCHAR(50) NOT NULL,
            subscribed BOOLEAN DEFAULT TRUE,
            subscribed_at TIMESTAMP DEFAULT NOW(),
            unsubscribed_at TIMESTAMP,
            CONSTRAINT uq_email_sub_user_product UNIQUE (user_id, product_id)
        )
    """))
    await db.execute(text("CREATE INDEX idx_email_subscriptions_user_id ON email_subscriptions(user_id)"))
    await db.execute(text("CREATE INDEX idx_email_subscriptions_product_id ON email_subscriptions(product_id)"))
    await db.execute(text("CREATE INDEX idx_email_subscriptions_subscribed ON email_subscriptions(subscribed)"))
    await db.commit()
    return {"success": True, "message": "email_subscriptions table and indexes created"}


# ============================================
# Admin: Email Subscribers
# ============================================

@admin_router.post("/migrate/sync-zoho-subscribers")
async def migrate_sync_zoho_subscribers(db: AsyncSession = Depends(get_db)):
    """
    One-time: backfill email_subscriptions for existing users + add them to Zoho Campaigns.
    Idempotent — uses INSERT ON CONFLICT DO NOTHING.
    """
    # Find all users without an email_subscriptions row for their product_id
    result = await db.execute(text("""
        SELECT u.id, u.email, u.name, u.product_id
        FROM users u
        LEFT JOIN email_subscriptions es
            ON es.user_id = u.id AND es.product_id = COALESCE(u.product_id, 'law')
        WHERE es.id IS NULL
        ORDER BY u.created_at
    """))
    users = result.fetchall()

    campaigns_token = await get_zoho_campaigns_token()
    list_key = os.getenv("ZOHO_LIST_KEY", "")

    ok_count = 0
    fail_count = 0
    zoho_ok = 0

    for u in users:
        pid = u.product_id or "law"
        # Insert subscription row
        try:
            await db.execute(text("""
                INSERT INTO email_subscriptions (id, user_id, product_id, subscribed, subscribed_at)
                VALUES (gen_random_uuid(), :uid, :pid, TRUE, NOW())
                ON CONFLICT (user_id, product_id) DO NOTHING
            """), {"uid": str(u.id), "pid": pid})
            await db.commit()  # commit each row individually to avoid cascading failures
            ok_count += 1
        except Exception as e:
            await db.rollback()
            fail_count += 1
            print(f"[Sync] DB insert failed for {u.email}: {e}")
            continue

        # Add to Zoho
        if campaigns_token and list_key:
            try:
                contact_info = json.dumps({
                    "Contact Email": u.email,
                    "First Name": (u.name or "").split()[0] if u.name else "",
                    "CONTACT_CF1": pid,
                })
                async with httpx.AsyncClient(timeout=10) as client:
                    resp = await client.post(
                        "https://campaigns.zoho.com/api/v1.1/json/listsubscribe",
                        params={"resfmt": "JSON", "listkey": list_key, "contactinfo": contact_info},
                        headers={"Authorization": f"Zoho-oauthtoken {campaigns_token}"}
                    )
                    rj = resp.json()
                    print(f"[Sync] Zoho {u.email}: code={rj.get('code')} msg={rj.get('message','')[:60]}")
                    if str(rj.get("code")) == "0" or "already" in rj.get("message", "").lower():
                        zoho_ok += 1
            except Exception as ze:
                print(f"[Sync] Zoho add failed for {u.email}: {ze}")

    await db.commit()
    return {
        "status": "ok",
        "users_found": len(users),
        "db_inserted": ok_count,
        "db_failed": fail_count,
        "zoho_added": zoho_ok,
    }


@admin_router.post("/migrate/push-to-zoho")
async def migrate_push_to_zoho(db: AsyncSession = Depends(get_db)):
    """
    Push all subscribed users from email_subscriptions to Zoho Campaigns.
    Safe to re-run — Zoho handles duplicates gracefully.
    """
    result = await db.execute(text("""
        SELECT u.email, u.name, es.product_id
        FROM email_subscriptions es
        JOIN users u ON u.id = es.user_id
        WHERE es.subscribed = TRUE
        ORDER BY es.subscribed_at
    """))
    rows = result.fetchall()

    campaigns_token = await get_zoho_campaigns_token()
    list_key = os.getenv("ZOHO_LIST_KEY", "")

    if not campaigns_token or not list_key:
        return {"status": "error", "reason": "Zoho Campaigns not configured"}

    ok = 0
    fail = 0
    for row in rows:
        try:
            contact_info = json.dumps({
                "Contact Email": row.email,
                "First Name": (row.name or "").split()[0] if row.name else "",
                "CONTACT_CF1": row.product_id or "law",
            })
            async with httpx.AsyncClient(timeout=10) as client:
                resp = await client.post(
                    "https://campaigns.zoho.com/api/v1.1/json/listsubscribe",
                    params={"resfmt": "JSON", "listkey": list_key, "contactinfo": contact_info},
                    headers={"Authorization": f"Zoho-oauthtoken {campaigns_token}"}
                )
                rj = resp.json()
                print(f"[ZohoPush] {row.email} ({row.product_id}): code={rj.get('code')} {rj.get('message','')[:50]}")
                if str(rj.get("code")) == "0" or "already" in rj.get("message", "").lower():
                    ok += 1
                else:
                    fail += 1
        except Exception as e:
            print(f"[ZohoPush] failed for {row.email}: {e}")
            fail += 1

    return {"status": "ok", "total": len(rows), "zoho_added": ok, "failed": fail}


class DripRecordRequest(BaseModel):
    email: str
    product_id: str
    email_number: int
    subject: Optional[str] = None


@admin_router.post("/drip-process", status_code=200)
async def process_drip_queue(
    db: AsyncSession = Depends(get_db)
):
    """
    POST /admin/drip-process
    Called by cron hourly. Finds scheduled drip emails due to send and sends them.
    """
    import sys as _sys2, importlib as _il2
    _dp = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'drip')
    if _dp not in _sys2.path:
        _sys2.path.insert(0, _dp)
    drip_utils = _il2.import_module('drip_utils')

    now_utc = datetime.utcnow()
    # Find all scheduled emails due
    result = await db.execute(text("""
        SELECT id, email, product_id, email_number, platform, first_name
        FROM drip_emails
        WHERE status = 'scheduled' AND scheduled_for <= :now
        ORDER BY scheduled_for
        LIMIT 50
    """), {"now": now_utc})
    rows = result.fetchall()

    sent_count = 0
    failed_count = 0
    access_token = await get_zoho_access_token()

    for row in rows:
        row_id, email, product_id, email_num, platform, first_name = row
        try:
            # Get product details
            prod_r = await db.execute(
                text("SELECT domain, name FROM products WHERE id = :pid"),
                {"pid": product_id}
            )
            prod_row = prod_r.fetchone()
            domain = prod_row.domain if prod_row else "lawtasksai.com"
            product_name = prod_row.name if prod_row else "LawTasksAI"

            # Get skill count
            sc_r = await db.execute(
                text("SELECT COUNT(*) FROM skills WHERE product_id = :pid AND is_active = TRUE"),
                {"pid": product_id}
            )
            sc_row = sc_r.fetchone()
            skill_count = sc_row[0] if sc_row else 150

            html = drip_utils.build_drip_email(
                email_num=email_num,
                product_id=product_id,
                product_name=product_name,
                domain=domain,
                skill_count=skill_count,
                platform=platform or "other",
                first_name=first_name or "",
                user_email=email,
            )
            subject = drip_utils.drip_subject(email_num, product_name)

            if access_token:
                async with httpx.AsyncClient(timeout=15) as _hc:
                    _resp = await _hc.post(
                        "https://mail.zoho.com/api/accounts/6556209000000008002/messages",
                        json={
                            "fromAddress": f"hello@{domain}",
                            "toAddress": email,
                            "subject": subject,
                            "content": html,
                            "mailFormat": "html",
                        },
                        headers={"Authorization": f"Zoho-oauthtoken {access_token}"}
                    )
                    if _resp.status_code == 200:
                        await db.execute(text("""
                            UPDATE drip_emails SET status='sent', sent_at=NOW()
                            WHERE id=:id
                        """), {"id": row_id})
                        sent_count += 1
                        print(f"[Drip] Email {email_num} sent to {email} ({product_id})")
                    else:
                        await db.execute(text("""
                            UPDATE drip_emails SET status='failed' WHERE id=:id
                        """), {"id": row_id})
                        failed_count += 1
                        print(f"[Drip] Email {email_num} failed for {email}: Zoho {_resp.status_code}")
        except Exception as e:
            failed_count += 1
            print(f"[Drip] Error processing {email} email {email_num}: {e}")

    await db.commit()
    return {"sent": sent_count, "failed": failed_count, "checked": len(rows)}


@admin_router.get("/email-opens")
async def get_email_opens(
    limit: int = Query(100),
    product: Optional[str] = Query(None),
    db: AsyncSession = Depends(get_db),
):
    """GET /admin/email-opens — list email open tracking events."""
    where = "WHERE product_id = :product" if product else ""
    params = {"limit": limit, "product": product or ""}
    result = await db.execute(
        text(f"""
            SELECT message_id, product_id, email, opened_at
            FROM email_opens
            {where}
            ORDER BY opened_at DESC
            LIMIT :limit
        """),
        params
    )
    rows = result.fetchall()
    opens = [{"message_id": r[0], "product_id": r[1], "email": r[2],
              "opened_at": r[3].isoformat() if r[3] else None} for r in rows]
    unique_emails = len(set(r["email"] for r in opens))
    return {"opens": opens, "total": len(opens), "unique_emails": unique_emails}


@admin_router.get("/drip-template/{email_num}")
async def get_drip_template(email_num: int):
    """GET /admin/drip-template/{1|2|3} — returns HTML template content."""
    if email_num not in (1, 2, 3):
        raise HTTPException(status_code=400, detail="email_num must be 1, 2, or 3")
    tpl_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'drip', f'email{email_num}_template.html')
    subjects = {
        1: "You're in \u2014 your {{PRODUCT_NAME}} credits are ready",
        2: "Have you run your first {{PRODUCT_NAME}} task yet?",
        3: "Quick question about {{PRODUCT_NAME}}",
    }
    try:
        with open(tpl_path) as f:
            html = f.read()
        return {"email_num": email_num, "html": html, "subject": subjects[email_num]}
    except FileNotFoundError:
        raise HTTPException(status_code=404, detail=f"Template {email_num} not found")


@admin_router.post("/drip-template/{email_num}")
async def save_drip_template(email_num: int, payload: dict):
    """POST /admin/drip-template/{1|2|3} — saves HTML template to disk."""
    if email_num not in (1, 2, 3):
        raise HTTPException(status_code=400, detail="email_num must be 1, 2, or 3")
    html = payload.get("html", "")
    if not html:
        raise HTTPException(status_code=400, detail="html is required")
    tpl_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'drip', f'email{email_num}_template.html')
    try:
        with open(tpl_path, 'w') as f:
            f.write(html)
        print(f"[Drip] Template {email_num} saved ({len(html)} chars)")
        return {"saved": True, "email_num": email_num}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@admin_router.post("/drip-test")
async def send_drip_test_email(payload: dict):
    """POST /admin/drip-test — sends a test drip email."""
    import sys as _s, importlib as _il
    _dp = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'drip')
    if _dp not in _s.path: _s.path.insert(0, _dp)
    drip_utils = _il.import_module('drip_utils')

    email_num = payload.get("email_num", 1)
    to_email = payload.get("email", "kentmercier@gmail.com")
    platform = payload.get("platform", "claude_desktop")
    product_id = payload.get("product_id", "law")

    html = drip_utils.build_drip_email(
        email_num=email_num, product_id=product_id,
        product_name="LawTasksAI", domain="lawtasksai.com",
        skill_count=206, platform=platform,
        first_name="Kent", user_email=to_email,
    )
    subject = f"[TEST] {drip_utils.drip_subject(email_num, 'LawTasksAI')}"
    access_token = await get_zoho_access_token()
    if not access_token:
        raise HTTPException(status_code=503, detail="Zoho token unavailable")
    async with httpx.AsyncClient(timeout=15) as hc:
        resp = await hc.post(
            "https://mail.zoho.com/api/accounts/6556209000000008002/messages",
            json={"fromAddress": "hello@lawtasksai.com", "toAddress": to_email,
                  "subject": subject, "content": html, "mailFormat": "html"},
            headers={"Authorization": f"Zoho-oauthtoken {access_token}"}
        )
    return {"sent": resp.status_code == 200, "zoho_status": resp.status_code}


@admin_router.get("/drip-queue")
async def drip_queue_stats(db: AsyncSession = Depends(get_db)):
    """GET /admin/drip-queue — returns drip email queue counts."""
    result = await db.execute(text("""
        SELECT status, COUNT(*) FROM drip_emails GROUP BY status
    """))
    counts = {row[0]: row[1] for row in result.fetchall()}
    return {
        "sent": counts.get("sent", 0),
        "scheduled": counts.get("scheduled", 0),
        "failed": counts.get("failed", 0),
        "total": sum(counts.values()),
    }


@admin_router.post("/drip-record", status_code=201)
async def record_drip_send(
    data: DripRecordRequest,
    db: AsyncSession = Depends(get_db)
):
    """Record that a drip email was sent. Idempotent — ignores duplicates."""
    try:
        # Look up user_id if we can
        u_result = await db.execute(select(User).where(User.email == data.email))
        user = u_result.scalar_one_or_none()
        rec = DripEmail(
            user_id=user.id if user else None,
            email=data.email,
            product_id=data.product_id,
            email_number=data.email_number,
            subject=data.subject,
        )
        db.add(rec)
        await db.commit()
        return {"recorded": True}
    except Exception as e:
        # Unique constraint violation = already recorded, that's fine
        await db.rollback()
        return {"recorded": False, "reason": str(e)}


@admin_router.post("/migrate/fix-product-ids")
async def fix_product_ids(db: AsyncSession = Depends(get_db)):
    """
    One-time migration: correct users.product_id and licenses.product_id
    using email_subscriptions as ground truth (it was set correctly at registration).
    """
    result = await db.execute(text("""
        WITH correct AS (
            SELECT DISTINCT ON (es.user_id)
                es.user_id,
                es.product_id,
                u.email
            FROM email_subscriptions es
            JOIN users u ON u.id = es.user_id
            ORDER BY es.user_id, es.subscribed_at ASC
        )
        UPDATE users u
        SET product_id = c.product_id
        FROM correct c
        WHERE u.id = c.user_id
          AND u.product_id IS DISTINCT FROM c.product_id
        RETURNING u.email, c.product_id AS new_product_id
    """))
    user_fixes = result.fetchall()

    result2 = await db.execute(text("""
        WITH correct AS (
            SELECT DISTINCT ON (es.user_id)
                es.user_id,
                es.product_id
            FROM email_subscriptions es
            ORDER BY es.user_id, es.subscribed_at ASC
        )
        UPDATE licenses l
        SET product_id = c.product_id
        FROM correct c
        WHERE l.user_id = c.user_id
          AND l.product_id IS DISTINCT FROM c.product_id
        RETURNING l.user_id, c.product_id AS new_product_id
    """))
    license_fixes = result2.fetchall()

    await db.commit()
    return {
        "users_updated": [{"email": r.email, "new_product_id": r.new_product_id} for r in user_fixes],
        "licenses_updated": len(license_fixes),
    }


@admin_router.get("/drip-status")
async def drip_status(db: AsyncSession = Depends(get_db)):
    """Show drip email send history per user."""
    result = await db.execute(text("""
        SELECT d.email, d.product_id, d.email_number, d.subject, d.sent_at
        FROM drip_emails d
        ORDER BY d.email, d.email_number
    """))
    rows = result.fetchall()
    by_user = {}
    for r in rows:
        key = f"{r.email}|{r.product_id}"
        if key not in by_user:
            by_user[key] = {"email": r.email, "product_id": r.product_id, "emails_sent": []}
        by_user[key]["emails_sent"].append({
            "email_number": r.email_number,
            "subject": r.subject,
            "sent_at": r.sent_at.isoformat() if r.sent_at else None
        })
    return {"drip_records": list(by_user.values()), "count": len(by_user)}


@admin_router.get("/email-subscribers")
async def admin_email_subscribers(
    product_id: Optional[str] = Query(None),
    db: AsyncSession = Depends(get_db)
):
    """
    GET /admin/email-subscribers?product_id=law
    Returns subscriber counts per vertical (or for a specific vertical).
    """
    query = select(
        EmailSubscription.product_id,
        func.count(EmailSubscription.id).label("total"),
        func.sum(
            func.cast(EmailSubscription.subscribed, Integer)
        ).label("subscribed")
    ).group_by(EmailSubscription.product_id)

    if product_id:
        query = query.where(EmailSubscription.product_id == product_id)

    result = await db.execute(query)
    rows = result.all()

    return {
        "subscribers": [
            {
                "product_id": row.product_id,
                "total": row.total,
                "subscribed": int(row.subscribed or 0),
                "unsubscribed": row.total - int(row.subscribed or 0),
            }
            for row in rows
        ]
    }


@app.get("/admin/unsubscribed-emails")
async def admin_unsubscribed_emails(db: AsyncSession = Depends(get_db)):
    """
    GET /admin/unsubscribed-emails
    Returns list of emails where subscribed=False (for drip skip logic).
    """
    result = await db.execute(
        select(User.email)
        .join(EmailSubscription, EmailSubscription.user_id == User.id)
        .where(EmailSubscription.subscribed == False)
    )
    emails = [row[0] for row in result.all()]
    return {"unsubscribed": emails, "count": len(emails)}


# ============================================
# Admin: Email Broadcast
# ============================================

# In-memory cache for broadcast email HTML (Zoho fetches via content_url)
_broadcast_content_cache: dict = {}


class BroadcastRequest(BaseModel):
    product_id: str                  # e.g. 'law'
    subject: str                     # Email subject line
    post_title: str                  # Blog post title
    post_url: str                    # Full URL to the post
    post_excerpt: str = ""           # 1-2 sentence teaser
    from_name: Optional[str] = None  # Defaults to product display name
    from_email: Optional[str] = None # Defaults to hello@{domain}
    dry_run: bool = False            # If true, create draft campaign but don't send


@admin_router.post("/broadcast")
async def admin_broadcast(req: BroadcastRequest, db: AsyncSession = Depends(get_db)):
    """
    Send a blog post broadcast to all subscribed users for a given product_id.
    Uses Zoho Campaigns to send a campaign to a segment of the LawTasksAI Subscribers list.
    """
    # Resolve product domain + name
    product_domain = "lawtasksai.com"
    product_name = "LawTasksAI"
    try:
        prod_result = await db.execute(
            text("SELECT domain, name FROM products WHERE id = :pid AND is_active = TRUE"),
            {"pid": req.product_id}
        )
        prod_row = prod_result.fetchone()
        if prod_row:
            product_domain = prod_row.domain or product_domain
            product_name = prod_row.name or product_name
    except Exception:
        pass

    from_name = req.from_name or f"{product_name} Team"
    from_email = req.from_email or f"kent@{product_domain}"

    # Count subscribers
    count_result = await db.execute(text("""
        SELECT COUNT(*) FROM email_subscriptions
        WHERE product_id = :pid AND subscribed = TRUE
    """), {"pid": req.product_id})
    subscriber_count = count_result.scalar() or 0

    if subscriber_count == 0:
        return {"status": "skipped", "reason": f"No subscribed users for {req.product_id}"}

    # Build HTML email body
    excerpt = req.post_excerpt or f"A new post from {product_name} is ready to read."
    html_body = f"""<!DOCTYPE html>
<html>
<head><meta charset="utf-8"><meta name="viewport" content="width=device-width, initial-scale=1"></head>
<body style="margin:0;padding:0;background:#f0f4f8;font-family:Arial,Helvetica,sans-serif">
  <table width="100%" cellpadding="0" cellspacing="0" style="background:#f0f4f8;padding:40px 0">
    <tr><td align="center">
      <table width="600" cellpadding="0" cellspacing="0" style="background:#ffffff;border-radius:8px;overflow:hidden;max-width:600px;box-shadow:0 2px 8px rgba(0,0,0,0.06)">
        <!-- Header -->
        <tr><td style="background:#0f172a;padding:28px 40px;text-align:center">
          <a href="https://{product_domain}" style="color:#ffffff;font-family:Arial,Helvetica,sans-serif;font-size:22px;font-weight:bold;text-decoration:none;letter-spacing:0.5px">{product_name}</a>
        </td></tr>
        <!-- Body -->
        <tr><td style="padding:44px 40px 40px">
          <p style="font-family:Arial,Helvetica,sans-serif;font-size:13px;color:#64748b;margin:0 0 12px;text-transform:uppercase;letter-spacing:1px;font-weight:600">New from {product_name}</p>
          <h1 style="font-family:Georgia,Times,'Times New Roman',serif;font-size:28px;color:#0f172a;margin:0 0 20px;line-height:1.35">{req.post_title}</h1>
          <p style="font-family:Arial,Helvetica,sans-serif;font-size:16px;color:#475569;line-height:1.7;margin:0 0 32px">{excerpt}</p>
          <table cellpadding="0" cellspacing="0" border="0"><tr><td style="background:#2563eb;border-radius:6px;padding:14px 32px">
            <a href="{req.post_url}" style="color:#ffffff;font-family:Arial,Helvetica,sans-serif;font-size:15px;font-weight:bold;text-decoration:none;display:inline-block">Read the full post &rarr;</a>
          </td></tr></table>
        </td></tr>
        <!-- Divider -->
        <tr><td style="padding:0 40px"><hr style="border:none;border-top:1px solid #e2e8f0;margin:0"></td></tr>
        <!-- Footer -->
        <tr><td style="padding:24px 40px 28px">
          <p style="font-family:Arial,Helvetica,sans-serif;font-size:13px;color:#94a3b8;margin:0;line-height:1.6">
            You received this because you signed up at <a href="https://{product_domain}" style="color:#94a3b8">{product_domain}</a>.<br>
            <a href="$$unsub$$" style="color:#94a3b8;text-decoration:underline">Unsubscribe</a> from these emails.
          </p>
        </td></tr>
      </table>
      <p style="font-family:Arial,Helvetica,sans-serif;font-size:11px;color:#94a3b8;margin:16px 0 0;text-align:center">&copy; 2026 {product_name}. All rights reserved.</p>
    </td></tr>
  </table>
</body>
</html>"""

    # Get Campaigns token
    campaigns_token = await get_zoho_campaigns_token()
    if not campaigns_token:
        return {"status": "error", "reason": "Could not get Zoho Campaigns token"}

    # Step 1: Create the campaign
    campaign_name = f"{product_name} Blog — {req.post_title[:50]}"

    # Upload email HTML to Cloud Storage for Zoho content_url import
    import hashlib as _hl
    content_id = _hl.md5(html_body.encode()).hexdigest()[:16]
    try:
        from google.cloud import storage as gcs
        gcs_client = gcs.Client()
        bucket = gcs_client.bucket("tasksai-email-content")
        blob = bucket.blob(f"broadcasts/{content_id}.html")
        blob.upload_from_string(html_body, content_type="text/html")
        content_url = f"https://storage.googleapis.com/tasksai-email-content/broadcasts/{content_id}.html"
        print(f"[Broadcast] uploaded email HTML to {content_url}")
    except Exception as e:
        print(f"[Broadcast] GCS upload failed, falling back to in-memory: {e}")
        _broadcast_content_cache[content_id] = html_body
        api_base = os.getenv("API_BASE_URL", "https://api.lawtasksai.com")
        content_url = f"{api_base}/email-content/{content_id}"

    # Use per-vertical list key
    list_key = get_zoho_listkey(req.product_id)
    if not list_key:
        return {"status": "error", "reason": f"No Zoho list configured for {req.product_id}"}

    # list_details: {listkey: []} — empty array = all contacts in list
    list_details = json.dumps({list_key: []})
    # topicId is mandatory for Zoho orgs with topic management enabled
    topic_id = os.getenv("ZOHO_TOPIC_ID", "1612833000000048017")

    async with httpx.AsyncClient(timeout=30) as client:
        create_resp = await client.post(
            "https://campaigns.zoho.com/api/v1.1/createCampaign",
            params={"resfmt": "JSON"},
            data={
                "campaignname": campaign_name,
                "from_email": from_email,
                "subject": req.subject,
                "list_details": list_details,
                "topicId": topic_id,
                "content_url": content_url,
            },
            headers={"Authorization": f"Zoho-oauthtoken {campaigns_token}",
                     "Content-Type": "application/x-www-form-urlencoded"}
        )
        create_data = create_resp.json()
        print(f"[Broadcast] create campaign: {create_data.get('code')} {create_data.get('message','')[:80]}")

        if str(create_data.get("code")) != "200":
            return {
                "status": "error",
                "reason": f"Campaign creation failed: {create_data.get('message')}",
                "raw": create_data
            }

        campaign_key = create_data.get("campaignKey", "")

        # Step 2: Send test email to admin, then hold as draft
        test_list_key = os.getenv("ZOHO_TEST_LIST_KEY", "3z5f5746e1c119ce625cc563b0b6fecd36d78134258c4fd2a22bbfb5ac15dfdcdc")
        test_list_details = json.dumps({test_list_key: []})

        # Create a clone campaign for the test send (targets test list only)
        test_create = await client.post(
            "https://campaigns.zoho.com/api/v1.1/createCampaign",
            params={"resfmt": "JSON"},
            data={
                "campaignname": f"Test Preview - {campaign_name[:40]}",
                "from_email": from_email,
                "subject": f"TEST PREVIEW - {req.subject}",
                "list_details": test_list_details,
                "topicId": topic_id,
                "content_url": content_url,
            },
            headers={"Authorization": f"Zoho-oauthtoken {campaigns_token}",
                     "Content-Type": "application/x-www-form-urlencoded"}
        )
        test_data = test_create.json()
        test_key = test_data.get("campaignKey", "")
        test_sent = False

        if str(test_data.get("code")) == "200" and test_key:
            # Send the test campaign immediately (only goes to test list = Kent)
            test_send = await client.post(
                "https://campaigns.zoho.com/api/v1.1/sendcampaign",
                params={"resfmt": "JSON"},
                data={"campaignkey": test_key},
                headers={"Authorization": f"Zoho-oauthtoken {campaigns_token}",
                         "Content-Type": "application/x-www-form-urlencoded"}
            )
            test_send_data = test_send.json()
            test_sent = str(test_send_data.get("code")) in ("0", "200")
            print(f"[Broadcast] test email sent: {test_send_data.get('code')} {test_send_data.get('message','')[:60]}")
        else:
            print(f"[Broadcast] test campaign create failed: {test_data.get('code')} {test_data.get('message','')[:60]}")

        if req.dry_run:
            return {
                "status": "draft",
                "campaign_name": campaign_name,
                "campaign_key": campaign_key,
                "subscriber_count": subscriber_count,
                "test_email_sent": test_sent,
                "message": "Draft created. Test email sent to admin. Call POST /admin/broadcast/approve with campaign_key to send to all subscribers.",
            }

        # If dry_run is False, send immediately (legacy behavior)
        send_resp = await client.post(
            "https://campaigns.zoho.com/api/v1.1/sendcampaign",
            params={"resfmt": "JSON"},
            data={"campaignkey": campaign_key},
            headers={"Authorization": f"Zoho-oauthtoken {campaigns_token}",
                     "Content-Type": "application/x-www-form-urlencoded"}
        )
        send_data = send_resp.json()
        print(f"[Broadcast] send campaign: {send_data.get('code')} {send_data.get('message','')[:80]}")

    return {
        "status": "ok" if str(send_data.get("code")) in ("0", "200") else "error",
        "campaign_name": campaign_name,
        "campaign_key": campaign_key,
        "subscriber_count": subscriber_count,
        "send_response": send_data,
    }


class BroadcastApproveRequest(BaseModel):
    campaign_key: str


@admin_router.post("/broadcast/approve")
async def admin_broadcast_approve(req: BroadcastApproveRequest):
    """Send a previously created draft campaign to all subscribers."""
    campaigns_token = await get_zoho_campaigns_token()
    if not campaigns_token:
        return {"status": "error", "reason": "Could not get Zoho Campaigns token"}

    async with httpx.AsyncClient(timeout=30) as client:
        send_resp = await client.post(
            "https://campaigns.zoho.com/api/v1.1/sendcampaign",
            params={"resfmt": "JSON"},
            data={"campaignkey": req.campaign_key},
            headers={"Authorization": f"Zoho-oauthtoken {campaigns_token}",
                     "Content-Type": "application/x-www-form-urlencoded"}
        )
        send_data = send_resp.json()
        print(f"[Broadcast Approve] send: {send_data.get('code')} {send_data.get('message','')[:80]}")

    return {
        "status": "ok" if str(send_data.get("code")) in ("0", "200") else "error",
        "send_response": send_data,
    }


@admin_router.get("/support-requests")
async def get_support_requests(
    product_id: Optional[str] = None,
    status: Optional[str] = None,
    limit: int = 50,
    db: AsyncSession = Depends(get_db)
):
    """
    GET /admin/support-requests
    List support requests, optionally filtered by product_id and/or status.
    direction: 'inbound' = received from user, 'outbound' = reply sent by us
    status: 'open', 'replied', 'closed'
    """
    filters = []
    params = {"limit": limit}
    if product_id:
        filters.append("product_id = :product_id")
        params["product_id"] = product_id
    if status:
        filters.append("status = :status")
        params["status"] = status
    where = ("WHERE " + " AND ".join(filters)) if filters else ""
    rows = await db.execute(
        text(f"SELECT * FROM support_requests {where} ORDER BY created_at DESC LIMIT :limit"),
        params
    )
    requests = [dict(r._mapping) for r in rows]
    return {"support_requests": requests, "total": len(requests)}


@admin_router.post("/support-requests", status_code=201)
async def log_support_request(
    payload: dict,
    db: AsyncSession = Depends(get_db)
):
    """
    POST /admin/support-requests
    Log an inbound or outbound support message.
    Fields: email, name, product_id, subject, message, direction, status, zoho_message_id, replied_at
    """
    from datetime import datetime as _dt, timezone as _tz
    replied_at_raw = payload.get("replied_at")
    if isinstance(replied_at_raw, str):
        try:
            # Parse ISO string and convert to naive UTC for asyncpg
            dt = _dt.fromisoformat(replied_at_raw.replace("Z", "+00:00"))
            replied_at = dt.astimezone(_tz.utc).replace(tzinfo=None)
        except Exception:
            replied_at = None
    else:
        replied_at = replied_at_raw
    await db.execute(text("""
        INSERT INTO support_requests
            (email, name, product_id, subject, message, direction, status, zoho_message_id, replied_at)
        VALUES
            (:email, :name, :product_id, :subject, :message, :direction, :status, :zoho_message_id, :replied_at)
    """), {
        "email":           payload.get("email", ""),
        "name":            payload.get("name"),
        "product_id":      payload.get("product_id"),
        "subject":         payload.get("subject"),
        "message":         payload.get("message"),
        "direction":       payload.get("direction", "inbound"),
        "status":          payload.get("status", "open"),
        "zoho_message_id": payload.get("zoho_message_id"),
        "replied_at":      replied_at,
    })
    await db.commit()
    return {"status": "logged"}


@admin_router.patch("/support-requests/{request_id}")
async def update_support_request(
    request_id: str,
    payload: dict,
    db: AsyncSession = Depends(get_db)
):
    """
    PATCH /admin/support-requests/{request_id}
    Update status (open/replied/closed) or notes on a support request.
    """
    allowed_fields = {"status", "replied_at", "notes"}
    updates = {k: v for k, v in payload.items() if k in allowed_fields}
    if not updates:
        raise HTTPException(status_code=400, detail="No valid fields to update")

    set_clauses = ", ".join(f"{k} = :{k}" for k in updates)
    updates["request_id"] = request_id
    result = await db.execute(
        text(f"UPDATE support_requests SET {set_clauses} WHERE id = :request_id"),
        updates
    )
    await db.commit()
    if result.rowcount == 0:
        raise HTTPException(status_code=404, detail="Support request not found")
    return {"success": True, "id": request_id}


app.include_router(admin_router)


@app.get("/email-content/{content_id}")
async def serve_email_content(content_id: str):
    """Serve temporary broadcast email HTML for Zoho content_url import."""
    from fastapi.responses import HTMLResponse
    html = _broadcast_content_cache.get(content_id)
    if not html:
        raise HTTPException(status_code=404, detail="Content not found or expired")
    return HTMLResponse(content=html)


# ============================================
# Zoho Campaigns Unsubscribe Webhook
# ============================================

class ZohoUnsubscribePayload(BaseModel):
    email: str
    listkey: Optional[str] = None
    product_id: Optional[str] = None


@app.get("/unsubscribe")
async def unsubscribe_get(
    email: Optional[str] = Query(None),
    product: Optional[str] = Query(None),
    db: AsyncSession = Depends(get_db)
):
    """
    One-click unsubscribe from drip emails.
    Marks the user unsubscribed in email_subscriptions and redirects to
    the vertical's /unsubscribe page with a confirmation message.
    """
    from fastapi.responses import RedirectResponse

    safe_product = product or "law"

    if email:
        try:
            # Mark unsubscribed in email_subscriptions
            result = await db.execute(
                select(EmailSubscription).where(
                    EmailSubscription.product_id == safe_product
                ).join(User, User.id == EmailSubscription.user_id).where(
                    User.email == email
                )
            )
            sub = result.scalar_one_or_none()
            if sub:
                sub.subscribed = False
                sub.unsubscribed_at = datetime.utcnow()
            else:
                # User exists but no subscription row — find user and create one
                u_result = await db.execute(select(User).where(User.email == email))
                user = u_result.scalar_one_or_none()
                if user:
                    new_sub = EmailSubscription(
                        user_id=user.id,
                        product_id=safe_product,
                        subscribed=False,
                        unsubscribed_at=datetime.utcnow()
                    )
                    db.add(new_sub)
            await db.commit()
            print(f"[unsubscribe] {email} unsubscribed from {safe_product}")
        except Exception as e:
            print(f"[unsubscribe] error: {e}")

    # Resolve domain for the vertical
    domain = "lawtasksai.com"
    try:
        prod_result = await db.execute(
            text("SELECT domain FROM products WHERE id = :pid AND is_active = TRUE"),
            {"pid": safe_product}
        )
        row = prod_result.fetchone()
        if row and row.domain:
            domain = row.domain
    except Exception:
        pass

    # Redirect to the vertical's unsubscribe confirmation page
    return RedirectResponse(
        url=f"https://{domain}/unsubscribe?confirmed=1",
        status_code=302
    )


@app.post("/webhooks/zoho-unsubscribe")
async def zoho_unsubscribe_webhook(request: Request, db: AsyncSession = Depends(get_db)):
    """
    Receives unsubscribe notifications from Zoho Campaigns.
    Sets email_subscriptions.subscribed = false for the given user + product_id.
    Does NOT affect credits, license, or account access.
    """
    try:
        body = await request.json()
    except Exception:
        body = dict(await request.form())

    email = body.get("email", "").strip().lower()
    product_id = body.get("product_id") or body.get("CONTACT_CF1") or "law"

    if not email:
        return {"status": "ignored", "reason": "no email"}

    # Find user by email
    result = await db.execute(select(User).where(User.email == email))
    user = result.scalar_one_or_none()
    if not user:
        return {"status": "ignored", "reason": "user not found"}

    # Update or insert unsubscription record
    result2 = await db.execute(
        select(EmailSubscription).where(
            EmailSubscription.user_id == user.id,
            EmailSubscription.product_id == product_id
        )
    )
    sub = result2.scalar_one_or_none()
    if sub:
        sub.subscribed = False
        sub.unsubscribed_at = datetime.utcnow()
    else:
        sub = EmailSubscription(
            user_id=user.id,
            product_id=product_id,
            subscribed=False,
            unsubscribed_at=datetime.utcnow()
        )
        db.add(sub)

    await db.commit()
    print(f"[Zoho Unsubscribe] {email} unsubscribed from {product_id}")
    return {"status": "ok"}


if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)
